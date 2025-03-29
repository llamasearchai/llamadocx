"""
Classes for working with tables in Word documents.

This module provides classes for working with tables, rows, columns, and cells
in Word documents. These classes wrap python-docx's table classes with enhanced
functionality.
"""

from typing import Optional, Union, List, Iterator, Tuple

from docx.table import Table as _Table, _Cell, _Row, _Column
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT

from .style import Style
from .paragraph import Paragraph


class Cell:
    """
    A cell in a table.

    This class wraps python-docx's _Cell class and provides additional
    functionality for cell formatting and content manipulation.

    Args:
        cell (_Cell): The underlying python-docx _Cell object

    Attributes:
        cell (_Cell): The underlying python-docx _Cell object
    """

    def __init__(self, cell: _Cell) -> None:
        self.cell = cell

    @property
    def text(self) -> str:
        """Get/set the cell's text content."""
        return self.cell.text

    @text.setter
    def text(self, value: str) -> None:
        self.cell.text = value

    @property
    def paragraphs(self) -> List[Paragraph]:
        """Get all paragraphs in the cell."""
        return [Paragraph(p) for p in self.cell.paragraphs]

    def add_paragraph(
        self,
        text: str = "",
        style: Optional[Union[str, Style]] = None
    ) -> Paragraph:
        """
        Add a paragraph to the cell.

        Args:
            text (str): The paragraph text
            style (Union[str, Style, None]): The paragraph style

        Returns:
            Paragraph: The created paragraph
        """
        style_name = style.name if isinstance(style, Style) else style
        para = self.cell.add_paragraph(text, style_name)
        return Paragraph(para)

    @property
    def width(self) -> Optional[float]:
        """Get/set the cell width in inches."""
        if self.cell.width is None:
            return None
        return self.cell.width.inches

    @width.setter
    def width(self, value: Optional[Union[int, float]]) -> None:
        if value is None:
            self.cell.width = None
        else:
            self.cell.width = Inches(value)

    @property
    def vertical_alignment(self) -> Optional[int]:
        """Get/set the cell's vertical alignment."""
        return self.cell.vertical_alignment

    @vertical_alignment.setter
    def vertical_alignment(self, value: Optional[int]) -> None:
        self.cell.vertical_alignment = value

    def set_vertical_alignment(self, alignment: str) -> None:
        """
        Set the cell's vertical alignment.

        Args:
            alignment (str): Alignment type ("top", "center", "bottom")
        """
        alignment_map = {
            "top": WD_ALIGN_VERTICAL.TOP,
            "center": WD_ALIGN_VERTICAL.CENTER,
            "bottom": WD_ALIGN_VERTICAL.BOTTOM,
        }
        if alignment.lower() not in alignment_map:
            raise ValueError(f"Invalid vertical alignment: {alignment}")
        self.cell.vertical_alignment = alignment_map[alignment.lower()]

    def merge(self, other: "Cell") -> None:
        """
        Merge this cell with another cell.

        Args:
            other (Cell): The cell to merge with
        """
        self.cell.merge(other.cell)

    def split(self) -> None:
        """Split a previously merged cell."""
        self.cell.split()

    def clear(self) -> None:
        """Clear the cell's content."""
        self.cell._tc.clear_content()


class Row:
    """
    A row in a table.

    This class wraps python-docx's _Row class and provides additional
    functionality for row manipulation.

    Args:
        row (_Row): The underlying python-docx _Row object

    Attributes:
        row (_Row): The underlying python-docx _Row object
    """

    def __init__(self, row: _Row) -> None:
        self.row = row

    @property
    def cells(self) -> List[Cell]:
        """Get all cells in the row."""
        return [Cell(cell) for cell in self.row.cells]

    def __getitem__(self, idx: int) -> Cell:
        """Get a cell by index."""
        return Cell(self.row.cells[idx])

    def __iter__(self) -> Iterator[Cell]:
        """Iterate over cells in the row."""
        return iter(self.cells)

    def __len__(self) -> int:
        """Get the number of cells in the row."""
        return len(self.row.cells)

    @property
    def height(self) -> Optional[float]:
        """Get/set the row height in inches."""
        if self.row.height is None:
            return None
        return self.row.height.inches

    @height.setter
    def height(self, value: Optional[Union[int, float]]) -> None:
        if value is None:
            self.row.height = None
        else:
            self.row.height = Inches(value)


class Column:
    """
    A column in a table.

    This class wraps python-docx's _Column class and provides additional
    functionality for column manipulation.

    Args:
        column (_Column): The underlying python-docx _Column object

    Attributes:
        column (_Column): The underlying python-docx _Column object
    """

    def __init__(self, column: _Column) -> None:
        self.column = column

    @property
    def cells(self) -> List[Cell]:
        """Get all cells in the column."""
        return [Cell(cell) for cell in self.column.cells]

    def __getitem__(self, idx: int) -> Cell:
        """Get a cell by index."""
        return Cell(self.column.cells[idx])

    def __iter__(self) -> Iterator[Cell]:
        """Iterate over cells in the column."""
        return iter(self.cells)

    def __len__(self) -> int:
        """Get the number of cells in the column."""
        return len(self.column.cells)

    @property
    def width(self) -> Optional[float]:
        """Get/set the column width in inches."""
        if self.column.width is None:
            return None
        return self.column.width.inches

    @width.setter
    def width(self, value: Optional[Union[int, float]]) -> None:
        if value is None:
            self.column.width = None
        else:
            self.column.width = Inches(value)


class Table:
    """
    A table in a Word document.

    This class wraps python-docx's Table class and provides additional
    functionality for table manipulation.

    Args:
        table (_Table): The underlying python-docx Table object

    Attributes:
        table (_Table): The underlying python-docx Table object
    """

    def __init__(self, table: _Table) -> None:
        self.table = table

    @property
    def rows(self) -> List[Row]:
        """Get all rows in the table."""
        return [Row(row) for row in self.table.rows]

    @property
    def columns(self) -> List[Column]:
        """Get all columns in the table."""
        return [Column(column) for column in self.table.columns]

    def __getitem__(self, idx: Union[int, Tuple[int, int]]) -> Union[Row, Cell]:
        """
        Get a row by index or cell by row and column indices.

        Args:
            idx (Union[int, Tuple[int, int]]): Row index or (row, col) indices

        Returns:
            Union[Row, Cell]: The requested row or cell
        """
        if isinstance(idx, tuple):
            row_idx, col_idx = idx
            return Cell(self.table.cell(row_idx, col_idx))
        return Row(self.table.rows[idx])

    def add_row(self) -> Row:
        """
        Add a row to the table.

        Returns:
            Row: The added row
        """
        return Row(self.table.add_row())

    def add_column(self) -> Column:
        """
        Add a column to the table.

        Returns:
            Column: The added column
        """
        return Column(self.table.add_column())

    @property
    def style(self) -> Optional[Style]:
        """Get/set the table's style."""
        if self.table.style is None:
            return None
        return Style(self.table.style)

    @style.setter
    def style(self, value: Optional[Union[str, Style]]) -> None:
        if isinstance(value, Style):
            self.table.style = value.style
        else:
            self.table.style = value

    @property
    def alignment(self) -> Optional[int]:
        """Get/set the table's alignment."""
        return self.table.alignment

    @alignment.setter
    def alignment(self, value: Optional[int]) -> None:
        self.table.alignment = value

    def set_alignment(self, alignment: str) -> None:
        """
        Set the table's alignment.

        Args:
            alignment (str): Alignment type ("left", "center", "right")
        """
        alignment_map = {
            "left": WD_TABLE_ALIGNMENT.LEFT,
            "center": WD_TABLE_ALIGNMENT.CENTER,
            "right": WD_TABLE_ALIGNMENT.RIGHT,
        }
        if alignment.lower() not in alignment_map:
            raise ValueError(f"Invalid table alignment: {alignment}")
        self.table.alignment = alignment_map[alignment.lower()]

    def autofit(self) -> None:
        """Autofit the table to its contents."""
        self.table.autofit = True

    def allow_autofit(self, value: bool = True) -> None:
        """
        Set whether the table can autofit.

        Args:
            value (bool): Whether to allow autofitting
        """
        self.table.allow_autofit = value

    def clear(self) -> None:
        """Clear all content from the table."""
        for row in self.rows:
            for cell in row:
                cell.clear()

    def __iter__(self) -> Iterator[Row]:
        """Iterate over rows in the table."""
        return iter(self.rows)

    def __len__(self) -> int:
        """Get the number of rows in the table."""
        return len(self.table.rows) 