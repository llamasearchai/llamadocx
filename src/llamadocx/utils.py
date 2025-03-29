"""
Utility functions for LlamaDocx.

This module provides various utility functions used throughout the LlamaDocx package.
"""

import os
import re
import logging
import hashlib
import tempfile
import warnings
from pathlib import Path
from typing import List, Dict, Any, Optional, Union, Tuple

# Set up module-level logger
logger = logging.getLogger(__name__)


def ensure_directory(directory: Union[str, Path]) -> Path:
    """Ensure a directory exists, creating it if necessary.

    Args:
        directory: The directory path to ensure exists.

    Returns:
        The Path object for the directory.
    """
    directory = Path(directory)
    if directory.exists() and not directory.is_dir():
        raise ValueError(f"Path exists but is not a directory: {directory}")

    if not directory.exists():
        directory.mkdir(parents=True, exist_ok=True)
        logger.debug(f"Created directory: {directory}")

    return directory


def sanitize_filename(filename: str) -> str:
    """Sanitize a filename to ensure it's valid across operating systems.

    Args:
        filename: The filename to sanitize.

    Returns:
        A sanitized filename string.
    """
    # Replace characters that are invalid in filenames
    invalid_chars = r'[<>:"/\\|?*]'
    sanitized = re.sub(invalid_chars, '_', filename)

    # Trim leading/trailing whitespace and periods
    sanitized = sanitized.strip(' .')

    # Ensure filename is not empty
    if not sanitized:
        sanitized = "unnamed"

    # Limit filename length (common limit is 255, but be conservative)
    if len(sanitized) > 200:
        # If too long, hash part of it and keep some of the original
        name_parts = os.path.splitext(sanitized)
        base_name = name_parts[0][:150]  # Keep first 150 chars
        extension = name_parts[1] if len(name_parts) > 1 else ""
        hash_part = hashlib.md5(sanitized.encode('utf-8')).hexdigest()[:10]
        sanitized = f"{base_name}_{hash_part}{extension}"

    return sanitized


def get_temp_file(suffix: str = '.docx') -> Tuple[str, Path]:
    """Get a temporary file path.

    Args:
        suffix: File extension for the temporary file.

    Returns:
        A tuple containing (file_id, file_path).
    """
    temp_dir = Path(tempfile.gettempdir()) / 'llamadocx'
    ensure_directory(temp_dir)

    # Generate a unique ID
    file_id = hashlib.md5(os.urandom(16)).hexdigest()[:10]
    temp_path = temp_dir / f"temp_{file_id}{suffix}"

    return file_id, temp_path


def cleanup_temp_files(file_id: Optional[str] = None) -> int:
    """Clean up temporary files.

    Args:
        file_id: Optional specific file ID to clean up.
            If None, all temporary files older than 1 day will be removed.

    Returns:
        Number of files removed.
    """
    import time
    from datetime import datetime, timedelta

    temp_dir = Path(tempfile.gettempdir()) / 'llamadocx'
    if not temp_dir.exists():
        return 0

    count = 0
    now = datetime.now()
    one_day_ago = now - timedelta(days=1)

    for temp_file in temp_dir.glob('temp_*'):
        # If file_id is specified, only remove matching files
        if file_id and file_id not in temp_file.name:
            continue

        # Check file age
        if not file_id:
            mtime = datetime.fromtimestamp(temp_file.stat().st_mtime)
            if mtime > one_day_ago:
                continue

        try:
            temp_file.unlink()
            count += 1
        except (PermissionError, OSError) as e:
            logger.warning(f"Could not remove temporary file {temp_file}: {e}")

    return count


def open_document(path: Union[str, Path], create_if_not_exists: bool = False):
    """Open a Word document, with proper error handling.

    Args:
        path: Path to the document.
        create_if_not_exists: If True, create a new document if the file doesn't exist.

    Returns:
        A Document object.

    Raises:
        FileNotFoundError: If the file doesn't exist and create_if_not_exists is False.
        ValueError: If the file exists but is not a valid Word document.
    """
    from docx import Document

    path = Path(path)

    if not path.exists():
        if create_if_not_exists:
            logger.info(f"Creating new document: {path}")
            doc = Document()
            return doc
        else:
            raise FileNotFoundError(f"Document not found: {path}")

    try:
        doc = Document(path)
        return doc
    except Exception as e:
        raise ValueError(f"Could not open document {path}: {e}")


def compare_documents(doc1_path: Union[str, Path], doc2_path: Union[str, Path]) -> Dict[str, Any]:
    """Compare two Word documents and return differences.

    Args:
        doc1_path: Path to the first document.
        doc2_path: Path to the second document.

    Returns:
        A dictionary containing differences between the documents.
    """
    from docx import Document

    doc1 = Document(doc1_path)
    doc2 = Document(doc2_path)

    differences = {
        'paragraphs': {
            'added': [],
            'removed': [],
            'modified': []
        },
        'tables': {
            'added': 0,
            'removed': 0,
            'modified': []
        }
    }

    # Compare paragraphs
    paragraphs1 = [p.text for p in doc1.paragraphs if p.text.strip()]
    paragraphs2 = [p.text for p in doc2.paragraphs if p.text.strip()]

    # Find added and removed paragraphs
    for i, p in enumerate(paragraphs1):
        if p not in paragraphs2:
            differences['paragraphs']['removed'].append({
                'index': i,
                'text': p
            })

    for i, p in enumerate(paragraphs2):
        if p not in paragraphs1:
            differences['paragraphs']['added'].append({
                'index': i,
                'text': p
            })

    # Compare tables (simple count for now)
    table_diff = len(doc2.tables) - len(doc1.tables)
    if table_diff > 0:
        differences['tables']['added'] = table_diff
    elif table_diff < 0:
        differences['tables']['removed'] = abs(table_diff)

    return differences


def extract_text_with_formatting(doc_path: Union[str, Path]) -> List[Dict[str, Any]]:
    """Extract text with formatting information from a Word document.

    Args:
        doc_path: Path to the document.

    Returns:
        A list of dictionaries containing text and formatting information.
    """
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document(doc_path)
    formatted_content = []

    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue

        # Get paragraph style and alignment
        style_name = paragraph.style.name if paragraph.style else "Normal"
        alignment = paragraph.alignment
        alignment_name = "LEFT"
        if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            alignment_name = "CENTER"
        elif alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            alignment_name = "RIGHT"
        elif alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
            alignment_name = "JUSTIFY"

        # Extract text and formatting for each run
        runs_info = []
        for run in paragraph.runs:
            if not run.text.strip():
                continue

            run_info = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font': run.font.name,
                'size': run.font.size
            }
            runs_info.append(run_info)

        para_info = {
            'text': paragraph.text,
            'style': style_name,
            'alignment': alignment_name,
            'runs': runs_info
        }
        formatted_content.append(para_info)

    return formatted_content


def get_document_statistics(doc_path: Union[str, Path]) -> Dict[str, Any]:
    """Get statistics about a Word document.

    Args:
        doc_path: Path to the document.

    Returns:
        A dictionary containing statistics about the document.
    """
    from docx import Document

    doc = Document(doc_path)

    # Count words
    word_count = 0
    for para in doc.paragraphs:
        word_count += len(para.text.split())

    # Count characters
    char_count = sum(len(para.text) for para in doc.paragraphs)
    char_count_no_spaces = sum(len(para.text.replace(" ", "")) for para in doc.paragraphs)

    # Count paragraphs
    para_count = sum(1 for para in doc.paragraphs if para.text.strip())

    # Count tables and cells
    table_count = len(doc.tables)
    cell_count = sum(len(table.rows) * len(table.columns) for table in doc.tables)

    # Find the most common words (excluding common stop words)
    stop_words = {"the", "and", "a", "an", "in", "on", "at", "to", "for", "of", "with", "is", "are"}
    words = []
    for para in doc.paragraphs:
        words.extend(word.lower() for word in re.findall(r'\b\w+\b', para.text) if word.lower() not in stop_words)

    word_freq = {}
    for word in words:
        word_freq[word] = word_freq.get(word, 0) + 1

    most_common = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]

    return {
        'word_count': word_count,
        'character_count': char_count,
        'character_count_no_spaces': char_count_no_spaces,
        'paragraph_count': para_count,
        'table_count': table_count,
        'cell_count': cell_count,
        'most_common_words': most_common
    } 