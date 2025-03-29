#!/usr/bin/env python3
"""
Command-line interface (CLI) for LlamaDocx.

This module provides the main entry point for the LlamaDocx command-line interface,
allowing users to interact with LlamaDocx functionality from the terminal.
"""

import os
import sys
import argparse
import logging
from pathlib import Path
from typing import Dict, List, Optional, Union, Any, Tuple

# Import LlamaDocx modules
from llamadocx import __version__
from llamadocx import Document
from llamadocx.utils import ensure_directory

# Set up module-level logger
logger = logging.getLogger(__name__)

__all__ = ['main', 'cli_app']


def set_up_logging(verbose: bool = False) -> None:
    """Configure logging based on verbosity level.

    Args:
        verbose: If True, set logging level to DEBUG, otherwise INFO.
    """
    log_level = logging.DEBUG if verbose else logging.INFO
    logging.basicConfig(
        level=log_level,
        format='%(asctime)s [%(levelname)s] %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )


def create_parser() -> argparse.ArgumentParser:
    """Create the command-line argument parser.

    Returns:
        An ArgumentParser object with all arguments defined.
    """
    parser = argparse.ArgumentParser(
        description='LlamaDocx - A powerful toolkit for working with Microsoft Word documents.',
        formatter_class=argparse.ArgumentDefaultsHelpFormatter
    )

    # Global arguments
    parser.add_argument(
        '-v', '--verbose',
        action='store_true',
        help='Enable verbose output'
    )
    parser.add_argument(
        '--version',
        action='store_true',
        help='Show version information and exit'
    )

    # Create subparsers for commands
    subparsers = parser.add_subparsers(
        title='commands',
        dest='command',
        help='Command to execute'
    )

    # Create subparser
    create_parser = subparsers.add_parser(
        'create',
        help='Create a new Word document'
    )
    create_parser.add_argument(
        'output',
        type=str,
        help='Path to save the new document'
    )
    create_parser.add_argument(
        '--title',
        type=str,
        help='Document title'
    )
    create_parser.add_argument(
        '--author',
        type=str,
        help='Document author'
    )

    # Extract subparser
    extract_parser = subparsers.add_parser(
        'extract',
        help='Extract content from a Word document'
    )
    extract_parser.add_argument(
        'input',
        type=str,
        help='Input Word document'
    )
    extract_parser.add_argument(
        '--output',
        type=str,
        help='Output file path (defaults to stdout)'
    )
    extract_parser.add_argument(
        '--format',
        type=str,
        choices=['text', 'markdown', 'html', 'json'],
        default='text',
        help='Output format'
    )
    extract_parser.add_argument(
        '--images-dir',
        type=str,
        help='Directory to save extracted images'
    )

    # Convert subparser
    convert_parser = subparsers.add_parser(
        'convert',
        help='Convert between document formats'
    )
    convert_parser.add_argument(
        'input',
        type=str,
        help='Input document'
    )
    convert_parser.add_argument(
        'output',
        type=str,
        help='Output document'
    )
    convert_parser.add_argument(
        '--from-format',
        type=str,
        help='Input format (auto-detected by default)'
    )
    convert_parser.add_argument(
        '--to-format',
        type=str,
        help='Output format (auto-detected by default)'
    )

    # Search subparser
    search_parser = subparsers.add_parser(
        'search',
        help='Search content in a Word document'
    )
    search_parser.add_argument(
        'input',
        type=str,
        help='Input Word document'
    )
    search_parser.add_argument(
        'query',
        type=str,
        help='Search query'
    )
    search_parser.add_argument(
        '--regex',
        action='store_true',
        help='Use regular expressions'
    )
    search_parser.add_argument(
        '--case-sensitive',
        action='store_true',
        help='Make search case-sensitive'
    )
    search_parser.add_argument(
        '--context',
        type=int,
        default=2,
        help='Number of context lines to show'
    )
    search_parser.add_argument(
        '--format',
        type=str,
        choices=['text', 'json'],
        default='text',
        help='Output format'
    )

    # Replace subparser
    replace_parser = subparsers.add_parser(
        'replace',
        help='Replace content in a Word document'
    )
    replace_parser.add_argument(
        'input',
        type=str,
        help='Input Word document'
    )
    replace_parser.add_argument(
        'pattern',
        type=str,
        help='Pattern to search for'
    )
    replace_parser.add_argument(
        'replacement',
        type=str,
        help='Replacement text'
    )
    replace_parser.add_argument(
        '--output',
        type=str,
        help='Output file (defaults to in-place modification)'
    )
    replace_parser.add_argument(
        '--regex',
        action='store_true',
        help='Use regular expressions'
    )
    replace_parser.add_argument(
        '--case-sensitive',
        action='store_true',
        help='Make replacement case-sensitive'
    )
    replace_parser.add_argument(
        '--count',
        action='store_true',
        help='Print the number of replacements made'
    )

    # Template subparser
    template_parser = subparsers.add_parser(
        'template',
        help='Process a Word document template'
    )
    template_parser.add_argument(
        'template',
        type=str,
        help='Template document'
    )
    template_parser.add_argument(
        'data',
        type=str,
        help='JSON data file with template values'
    )
    template_parser.add_argument(
        'output',
        type=str,
        help='Output document'
    )
    template_parser.add_argument(
        '--remove-empty',
        action='store_true',
        help='Remove empty fields'
    )
    template_parser.add_argument(
        '--list-fields',
        action='store_true',
        help='List fields in the template and exit'
    )

    # Meta subparser
    meta_parser = subparsers.add_parser(
        'meta',
        help='View or modify document metadata'
    )
    meta_group = meta_parser.add_mutually_exclusive_group(required=True)
    meta_group.add_argument(
        '--get',
        action='store_true',
        help='Get document metadata'
    )
    meta_group.add_argument(
        '--set',
        type=str,
        help='JSON file with metadata to set'
    )
    meta_group.add_argument(
        '--extract',
        type=str,
        help='Extract metadata to a JSON file'
    )
    meta_parser.add_argument(
        'document',
        type=str,
        help='Word document'
    )
    meta_parser.add_argument(
        '--output',
        type=str,
        help='Output document path (when setting metadata)'
    )
    meta_parser.add_argument(
        '--format',
        type=str,
        choices=['text', 'json'],
        default='text',
        help='Output format for --get'
    )

    return parser


def show_version() -> None:
    """Display version information for LlamaDocx and its dependencies."""
    import platform
    import docx
    
    print(f"LlamaDocx version: {__version__}")
    print(f"Python version: {platform.python_version()}")
    print(f"Platform: {platform.platform()}")
    print(f"python-docx version: {docx.__version__}")
    
    # Try to import optional dependencies
    try:
        import markdown
        print(f"markdown version: {markdown.__version__}")
    except ImportError:
        pass
    
    try:
        import bs4
        print(f"beautifulsoup4 version: {bs4.__version__}")
    except ImportError:
        pass
    
    try:
        import pandoc
        print(f"pypandoc version: {pandoc.__version__}")
    except ImportError:
        pass


def handle_create_command(args: argparse.Namespace) -> int:
    """Handle the 'create' command.
    
    Args:
        args: The parsed command-line arguments.
        
    Returns:
        0 on success, non-zero on failure.
    """
    try:
        doc = Document()
        
        # Set metadata if provided
        if args.title:
            from llamadocx.metadata import set_title
            set_title(doc, args.title)
        
        if args.author:
            from llamadocx.metadata import set_author
            set_author(doc, args.author)
        
        # Ensure output directory exists
        output_path = Path(args.output)
        ensure_directory(output_path.parent)
        
        # Save document
        doc.save(output_path)
        logger.info(f"Created new document: {output_path}")
        return 0
    except Exception as e:
        logger.error(f"Error creating document: {e}")
        return 1


def handle_extract_command(args: argparse.Namespace) -> int:
    """Handle the 'extract' command.
    
    Args:
        args: The parsed command-line arguments.
        
    Returns:
        0 on success, non-zero on failure.
    """
    try:
        # Open document
        doc = Document(args.input)
        
        # Handle extraction based on format
        if args.format == 'text':
            content = '\n\n'.join(p.text for p in doc.paragraphs)
        elif args.format == 'markdown':
            from llamadocx.convert import docx_to_markdown
            content = docx_to_markdown(args.input, None, return_content=True)
        elif args.format == 'html':
            from llamadocx.convert import docx_to_html
            content = docx_to_html(args.input, None, return_content=True)
        elif args.format == 'json':
            import json
            content_dict = {
                'paragraphs': [p.text for p in doc.paragraphs],
                'tables': []
            }
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                content_dict['tables'].append(table_data)
            content = json.dumps(content_dict, indent=2)
        
        # Extract images if requested
        if args.images_dir:
            from llamadocx.image import Image
            images_dir = Path(args.images_dir)
            ensure_directory(images_dir)
            
            count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_path = images_dir / f"image_{count:03d}.png"
                    with open(image_path, 'wb') as f:
                        f.write(rel.target_part.blob)
                    count += 1
            
            logger.info(f"Extracted {count} images to {images_dir}")
        
        # Output content
        if args.output:
            output_path = Path(args.output)
            ensure_directory(output_path.parent)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            logger.info(f"Content extracted to {output_path}")
        else:
            print(content)
        
        return 0
    except Exception as e:
        logger.error(f"Error extracting content: {e}")
        return 1


def handle_convert_command(args: argparse.Namespace) -> int:
    """Handle the 'convert' command.
    
    Args:
        args: The parsed command-line arguments.
        
    Returns:
        0 on success, non-zero on failure.
    """
    try:
        input_path = Path(args.input)
        output_path = Path(args.output)
        
        # Ensure output directory exists
        ensure_directory(output_path.parent)
        
        # Determine input and output formats
        in_format = args.from_format or input_path.suffix.lstrip('.')
        out_format = args.to_format or output_path.suffix.lstrip('.')
        
        # Handle conversion based on formats
        if in_format == 'docx' and out_format == 'pdf':
            from llamadocx.convert import docx_to_pdf
            docx_to_pdf(input_path, output_path)
        elif in_format == 'docx' and out_format == 'md':
            from llamadocx.convert import docx_to_markdown
            docx_to_markdown(input_path, output_path)
        elif in_format == 'docx' and out_format == 'html':
            from llamadocx.convert import docx_to_html
            docx_to_html(input_path, output_path)
        elif in_format == 'md' and out_format == 'docx':
            from llamadocx.convert import md_to_docx
            md_to_docx(input_path, output_path)
        elif in_format == 'html' and out_format == 'docx':
            from llamadocx.convert import html_to_docx
            html_to_docx(input_path, output_path)
        else:
            logger.error(f"Unsupported conversion: {in_format} to {out_format}")
            return 1
        
        logger.info(f"Converted {input_path} to {output_path}")
        return 0
    except Exception as e:
        logger.error(f"Error converting document: {e}")
        return 1


def handle_search_command(args: argparse.Namespace) -> int:
    """Handle the 'search' command.
    
    Args:
        args: The parsed command-line arguments.
        
    Returns:
        0 on success, non-zero on failure.
    """
    try:
        # Import search function
        from llamadocx.search import search_text
        
        # Open document
        doc = Document(args.input)
        
        # Perform search
        results = search_text(
            doc, 
            args.query, 
            regex=args.regex, 
            case_sensitive=args.case_sensitive,
            context_lines=args.context
        )
        
        # Format and output results
        if args.format == 'json':
            import json
            print(json.dumps(results, indent=2))
        else:
            if not results:
                print(f"No matches found for '{args.query}'")
            else:
                print(f"Found {len(results)} matches:")
                for i, result in enumerate(results):
                    print(f"\nMatch {i+1}:")
                    print(f"  Location: Paragraph {result['paragraph_index'] + 1}")
                    print(f"  Context: {result['context']}")
        
        return 0
    except Exception as e:
        logger.error(f"Error searching document: {e}")
        return 1


def handle_replace_command(args: argparse.Namespace) -> int:
    """Handle the 'replace' command.
    
    Args:
        args: The parsed command-line arguments.
        
    Returns:
        0 on success, non-zero on failure.
    """
    try:
        # Import replace function
        from llamadocx.search import replace_text
        
        # Open document
        doc = Document(args.input)
        
        # Perform replacement
        count = replace_text(
            doc, 
            args.pattern, 
            args.replacement, 
            regex=args.regex, 
            case_sensitive=args.case_sensitive
        )
        
        # Save document
        output_path = args.output or args.input
        doc.save(output_path)
        
        if args.count:
            print(f"Made {count} replacements")
        
        logger.info(f"Saved document with replacements to {output_path}")
        return 0
    except Exception as e:
        logger.error(f"Error replacing content: {e}")
        return 1


def handle_template_command(args: argparse.Namespace) -> int:
    """Handle the 'template' command.
    
    Args:
        args: The parsed command-line arguments.
        
    Returns:
        0 on success, non-zero on failure.
    """
    try:
        # Import Template class
        from llamadocx.template import Template
        
        # Create Template object
        template = Template(args.template)
        
        # List fields and exit if requested
        if args.list_fields:
            fields = template.get_fields()
            print("Fields in template:")
            for field in sorted(fields):
                print(f"  {field}")
            return 0
        
        # Load data
        import json
        with open(args.data, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Merge fields
        template.merge_fields(data, remove_empty=args.remove_empty)
        
        # Save document
        output_path = Path(args.output)
        ensure_directory(output_path.parent)
        template.save(output_path)
        
        logger.info(f"Template processed and saved to {output_path}")
        return 0
    except Exception as e:
        logger.error(f"Error processing template: {e}")
        return 1


def handle_meta_command(args: argparse.Namespace) -> int:
    """Handle the 'meta' command.
    
    Args:
        args: The parsed command-line arguments.
        
    Returns:
        0 on success, non-zero on failure.
    """
    try:
        # Import metadata functions
        from llamadocx.metadata import get_metadata, set_metadata, extract_metadata_to_file, update_metadata_from_file
        
        # Open document
        doc = Document(args.document)
        
        # Get metadata
        if args.get:
            metadata = get_metadata(doc)
            if args.format == 'json':
                import json
                print(json.dumps(metadata, indent=2, default=str))
            else:
                print("Document Metadata:")
                for key, value in metadata.items():
                    print(f"  {key}: {value}")
        
        # Set metadata
        elif args.set:
            import json
            with open(args.set, 'r', encoding='utf-8') as f:
                metadata = json.load(f)
            
            update_metadata_from_file(doc, args.set)
            
            # Save document
            output_path = args.output or args.document
            doc.save(output_path)
            logger.info(f"Metadata updated and saved to {output_path}")
        
        # Extract metadata
        elif args.extract:
            extract_path = Path(args.extract)
            ensure_directory(extract_path.parent)
            extract_metadata_to_file(doc, extract_path)
            logger.info(f"Metadata extracted to {extract_path}")
        
        return 0
    except Exception as e:
        logger.error(f"Error handling metadata: {e}")
        return 1


def cli_app(args: Optional[List[str]] = None) -> int:
    """Execute the CLI application.

    Args:
        args: Command line arguments (if None, sys.argv is used)

    Returns:
        Exit code (0 for success, non-zero for failure)
    """
    parser = create_parser()
    parsed_args = parser.parse_args(args)

    # Set up logging
    set_up_logging(parsed_args.verbose)

    # Show version and exit if requested
    if parsed_args.version:
        show_version()
        return 0

    # Handle commands
    if parsed_args.command == 'create':
        return handle_create_command(parsed_args)
    elif parsed_args.command == 'extract':
        return handle_extract_command(parsed_args)
    elif parsed_args.command == 'convert':
        return handle_convert_command(parsed_args)
    elif parsed_args.command == 'search':
        return handle_search_command(parsed_args)
    elif parsed_args.command == 'replace':
        return handle_replace_command(parsed_args)
    elif parsed_args.command == 'template':
        return handle_template_command(parsed_args)
    elif parsed_args.command == 'meta':
        return handle_meta_command(parsed_args)
    else:
        parser.print_help()
        return 1


def main() -> None:
    """Main entry point for the CLI application."""
    try:
        exit_code = cli_app()
        sys.exit(exit_code)
    except Exception as e:
        logger.error(f"Unhandled exception: {e}", exc_info=True)
        sys.exit(1)


if __name__ == "__main__":
    main() 