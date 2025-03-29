"""
Classes for working with paragraphs and text runs in Word documents.

This module provides the Paragraph and Run classes that wrap python-docx's
paragraph and run classes with enhanced functionality.
"""

from typing import Optional, Union, List, Iterator, Any

from docx.text.paragraph import Paragraph as _Paragraph
from docx.text.run import Run as _Run
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

from .style import Style, Font, ParagraphFormat


class Run:
    """
    A run of text with consistent formatting.

    This class wraps python-docx's Run class and provides additional functionality
    for text formatting and manipulation.

    Args:
        run (_Run): The underlying python-docx Run object

    Attributes:
        run (_Run): The underlying python-docx Run object
    """

    def __init__(self, run: _Run) -> None:
        self.run = run

    @property
    def text(self) -> str:
        """Get the run's text content."""
        return self.run.text

    @text.setter
    def text(self, value: str) -> None:
        """Set the run's text content."""
        self.run.text = value

    @property
    def bold(self) -> bool:
        """Get/set whether the run is bold."""
        return self.run.bold

    @bold.setter
    def bold(self, value: bool) -> None:
        self.run.bold = value

    @property
    def italic(self) -> bool:
        """Get/set whether the run is italic."""
        return self.run.italic

    @italic.setter
    def italic(self, value: bool) -> None:
        self.run.italic = value

    @property
    def underline(self) -> bool:
        """Get/set whether the run is underlined."""
        return bool(self.run.underline)

    @underline.setter
    def underline(self, value: bool) -> None:
        self.run.underline = value

    @property
    def font(self) -> Font:
        """Get the run's font settings."""
        return Font(self.run.font)

    def set_font(
        self,
        name: Optional[str] = None,
        size: Optional[Union[int, float]] = None,
        color: Optional[Union[str, tuple]] = None,
        bold: Optional[bool] = None,
        italic: Optional[bool] = None,
        underline: Optional[bool] = None,
    ) -> None:
        """
        Set multiple font properties at once.

        Args:
            name (str, optional): Font name
            size (Union[int, float], optional): Font size in points
            color (Union[str, tuple], optional): Font color as RGB tuple or hex string
            bold (bool, optional): Whether text is bold
            italic (bool, optional): Whether text is italic
            underline (bool, optional): Whether text is underlined
        """
        if name is not None:
            self.run.font.name = name
        if size is not None:
            self.run.font.size = Pt(size)
        if color is not None:
            if isinstance(color, str):
                # Convert hex color to RGB
                color = color.lstrip("#")
                r, g, b = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
            else:
                r, g, b = color
            self.run.font.color.rgb = RGBColor(r, g, b)
        if bold is not None:
            self.run.bold = bold
        if italic is not None:
            self.run.italic = italic
        if underline is not None:
            self.run.underline = underline

    def add_break(self, break_type: Optional[int] = None) -> None:
        """
        Add a break to the run.

        Args:
            break_type (int, optional): Type of break to add
                (WD_BREAK.LINE, WD_BREAK.PAGE, etc.)
        """
        self.run.add_break(break_type)

    def clear(self) -> None:
        """Clear the run's text content."""
        self.run.text = ""

    def copy_format_from(self, other: Union["Run", _Run]) -> None:
        """
        Copy formatting from another run.

        Args:
            other (Union[Run, _Run]): Run to copy formatting from
        """
        other_run = other.run if isinstance(other, Run) else other
        self.run.bold = other_run.bold
        self.run.italic = other_run.italic
        self.run.underline = other_run.underline
        self.run.font.name = other_run.font.name
        self.run.font.size = other_run.font.size
        if other_run.font.color.rgb is not None:
            self.run.font.color.rgb = other_run.font.color.rgb


class Paragraph:
    """
    A paragraph in a Word document.

    This class wraps python-docx's Paragraph class and provides additional
    functionality for paragraph formatting and manipulation.

    Args:
        paragraph (_Paragraph): The underlying python-docx Paragraph object

    Attributes:
        paragraph (_Paragraph): The underlying python-docx Paragraph object
    """

    def __init__(self, paragraph: _Paragraph) -> None:
        self.paragraph = paragraph

    @property
    def text(self) -> str:
        """Get the paragraph's text content."""
        return self.paragraph.text

    @text.setter
    def text(self, value: str) -> None:
        """Set the paragraph's text content."""
        self.clear()
        self.add_run(value)

    @property
    def style(self) -> Optional[Style]:
        """Get/set the paragraph's style."""
        if self.paragraph.style is None:
            return None
        return Style(self.paragraph.style)

    @style.setter
    def style(self, value: Optional[Union[str, Style]]) -> None:
        if isinstance(value, Style):
            self.paragraph.style = value.style
        else:
            self.paragraph.style = value

    @property
    def alignment(self) -> Optional[int]:
        """Get/set the paragraph's alignment."""
        return self.paragraph.alignment

    @alignment.setter
    def alignment(self, value: Optional[int]) -> None:
        self.paragraph.alignment = value

    @property
    def format(self) -> ParagraphFormat:
        """Get the paragraph's formatting."""
        return ParagraphFormat(self.paragraph.paragraph_format)

    def set_alignment(self, alignment: str) -> None:
        """
        Set the paragraph's alignment.

        Args:
            alignment (str): Alignment type ("left", "center", "right", "justify")
        """
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        if alignment.lower() not in alignment_map:
            raise ValueError(f"Invalid alignment: {alignment}")
        self.paragraph.alignment = alignment_map[alignment.lower()]

    def set_line_spacing(
        self,
        spacing: Union[float, str],
        rule: Optional[str] = None
    ) -> None:
        """
        Set the paragraph's line spacing.

        Args:
            spacing (Union[float, str]): Line spacing value or preset
            rule (str, optional): Line spacing rule
                ("single", "1.5 lines", "double", "exactly", "multiple")
        """
        if isinstance(spacing, str):
            rule = spacing
            if rule == "single":
                spacing = 1.0
            elif rule == "1.5 lines":
                spacing = 1.5
            elif rule == "double":
                spacing = 2.0
            else:
                raise ValueError(f"Invalid line spacing preset: {spacing}")

        rule_map = {
            "single": WD_LINE_SPACING.SINGLE,
            "1.5 lines": WD_LINE_SPACING.ONE_POINT_FIVE,
            "double": WD_LINE_SPACING.DOUBLE,
            "exactly": WD_LINE_SPACING.EXACTLY,
            "multiple": WD_LINE_SPACING.MULTIPLE,
        }

        if rule is not None:
            if rule.lower() not in rule_map:
                raise ValueError(f"Invalid line spacing rule: {rule}")
            self.paragraph.paragraph_format.line_spacing_rule = rule_map[rule.lower()]

        self.paragraph.paragraph_format.line_spacing = spacing

    def set_spacing(
        self,
        before: Optional[Union[int, float]] = None,
        after: Optional[Union[int, float]] = None
    ) -> None:
        """
        Set spacing before and after the paragraph.

        Args:
            before (Union[int, float], optional): Spacing before in points
            after (Union[int, float], optional): Spacing after in points
        """
        if before is not None:
            self.paragraph.paragraph_format.space_before = Pt(before)
        if after is not None:
            self.paragraph.paragraph_format.space_after = Pt(after)

    def set_indentation(
        self,
        left: Optional[Union[int, float]] = None,
        right: Optional[Union[int, float]] = None,
        first_line: Optional[Union[int, float]] = None,
        hanging: Optional[Union[int, float]] = None
    ) -> None:
        """
        Set paragraph indentation.

        Args:
            left (Union[int, float], optional): Left indentation in inches
            right (Union[int, float], optional): Right indentation in inches
            first_line (Union[int, float], optional): First line indentation in inches
            hanging (Union[int, float], optional): Hanging indentation in inches
        """
        if left is not None:
            self.paragraph.paragraph_format.left_indent = Inches(left)
        if right is not None:
            self.paragraph.paragraph_format.right_indent = Inches(right)
        if first_line is not None:
            self.paragraph.paragraph_format.first_line_indent = Inches(first_line)
        if hanging is not None:
            self.paragraph.paragraph_format.first_line_indent = Inches(-hanging)

    def add_run(
        self,
        text: str = "",
        style: Optional[Union[str, Style]] = None
    ) -> Run:
        """
        Add a run of text to the paragraph.

        Args:
            text (str): Text content
            style (Union[str, Style], optional): Run style

        Returns:
            Run: The created run
        """
        run = self.paragraph.add_run(text)
        if style is not None:
            if isinstance(style, Style):
                run.style = style.style
            else:
                run.style = style
        return Run(run)

    def insert_run_before(
        self,
        text: str = "",
        style: Optional[Union[str, Style]] = None
    ) -> Run:
        """
        Insert a run of text at the start of the paragraph.

        Args:
            text (str): Text content
            style (Union[str, Style], optional): Run style

        Returns:
            Run: The created run
        """
        run = Run(self.paragraph.add_run(text))
        if style is not None:
            run.style = style
        # Move run to start of paragraph
        self.paragraph._p.insert(0, run.run._r)
        return run

    def clear(self) -> None:
        """Clear all runs from the paragraph."""
        for run in self.runs:
            self.paragraph._p.remove(run.run._r)

    @property
    def runs(self) -> List[Run]:
        """Get all runs in the paragraph."""
        return [Run(run) for run in self.paragraph.runs]

    def __iter__(self) -> Iterator[Run]:
        """Iterate over runs in the paragraph."""
        return iter(self.runs)

    def __len__(self) -> int:
        """Get the number of runs in the paragraph."""
        return len(self.runs) 