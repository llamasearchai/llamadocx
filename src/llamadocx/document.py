"""
Core Document class for working with Word documents.

This module provides the main Document class that serves as the primary interface
for working with .docx files. It wraps python-docx's Document class and provides
additional functionality for document processing.
"""

import os
from pathlib import Path
from typing import Optional, Union, List, Dict, Any, Iterator

from docx import Document as DocxDocument
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph as _Paragraph
from docx.table import _Cell, _Row, _Column, Table as _Table

from .paragraph import Paragraph
from .table import Table
from .style import Style
from .image import Image, ImageProcessor
from .utils import ensure_path, validate_docx


class Document:
    """
    Main class for working with Word documents.

    This class provides methods for reading, writing, and modifying Word documents.
    It wraps python-docx's Document class and adds enhanced functionality.

    Args:
        path (Union[str, Path, None]): Path to an existing .docx file, or None to create a new document
        template (Union[str, Path, None]): Path to a template .docx file to use as a base

    Attributes:
        doc (_Document): The underlying python-docx Document object
        path (Path): Path to the document file
    """

    def __init__(
        self,
        path: Optional[Union[str, Path]] = None,
        template: Optional[Union[str, Path]] = None,
    ) -> None:
        if path is not None:
            path = ensure_path(path)
            validate_docx(path)
            self.doc = DocxDocument(path)
            self.path = path
        elif template is not None:
            template = ensure_path(template)
            validate_docx(template)
            self.doc = DocxDocument(template)
            self.path = None
        else:
            self.doc = DocxDocument()
            self.path = None

    def save(self, path: Optional[Union[str, Path]] = None) -> None:
        """
        Save the document to a file.

        Args:
            path (Union[str, Path, None]): Path to save the document to.
                If None, uses the original path.

        Raises:
            ValueError: If no path is provided and the document wasn't loaded from a file
        """
        if path is None:
            if self.path is None:
                raise ValueError("No path provided to save the document")
            path = self.path
        else:
            path = ensure_path(path)

        self.doc.save(str(path))
        self.path = path

    def get_text(self, include_tables: bool = True) -> str:
        """
        Extract all text from the document.

        Args:
            include_tables (bool): Whether to include text from tables

        Returns:
            str: The extracted text
        """
        text = []
        for paragraph in self.paragraphs:
            text.append(paragraph.text)
        
        if include_tables:
            for table in self.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
        
        return "\n".join(text)

    def add_heading(self, text: str, level: int = 1) -> Paragraph:
        """
        Add a heading to the document.

        Args:
            text (str): The heading text
            level (int): The heading level (1-9)

        Returns:
            Paragraph: The created heading paragraph
        """
        if not 1 <= level <= 9:
            raise ValueError("Heading level must be between 1 and 9")
        
        heading = self.doc.add_heading(text, level)
        return Paragraph(heading)

    def add_paragraph(
        self,
        text: str = "",
        style: Optional[Union[str, Style]] = None
    ) -> Paragraph:
        """
        Add a paragraph to the document.

        Args:
            text (str): The paragraph text
            style (Union[str, Style, None]): The paragraph style

        Returns:
            Paragraph: The created paragraph
        """
        style_name = style.name if isinstance(style, Style) else style
        para = self.doc.add_paragraph(text, style_name)
        return Paragraph(para)

    def add_table(
        self,
        rows: int,
        cols: int,
        style: Optional[Union[str, Style]] = None
    ) -> Table:
        """
        Add a table to the document.

        Args:
            rows (int): Number of rows
            cols (int): Number of columns
            style (Union[str, Style, None]): The table style

        Returns:
            Table: The created table
        """
        style_name = style.name if isinstance(style, Style) else style
        table = self.doc.add_table(rows, cols, style_name)
        return Table(table)

    def add_image(
        self,
        path: Union[str, Path],
        width: Optional[float] = None,
        height: Optional[float] = None
    ) -> Image:
        """
        Add an image to the document.

        Args:
            path (Union[str, Path]): Path to the image file
            width (float, optional): Desired width in inches
            height (float, optional): Desired height in inches

        Returns:
            Image: The added image
        """
        path = ensure_path(path)
        if not path.exists():
            raise FileNotFoundError(f"Image file not found: {path}")

        image = self.doc.add_picture(str(path))
        if width is not None:
            image.width = int(width * 914400)  # convert to EMU
        if height is not None:
            image.height = int(height * 914400)  # convert to EMU
        
        return Image(image)

    def search_text(
        self,
        pattern: str,
        regex: bool = False,
        case_sensitive: bool = False
    ) -> List[Dict[str, Any]]:
        """
        Search for text in the document.

        Args:
            pattern (str): Text pattern to search for
            regex (bool): Whether to treat pattern as regex
            case_sensitive (bool): Whether to match case

        Returns:
            List[Dict[str, Any]]: List of matches with location info
        """
        from .search import search_text
        return search_text(self, pattern, regex, case_sensitive)

    def replace_text(
        self,
        pattern: str,
        replacement: str,
        regex: bool = False,
        case_sensitive: bool = False
    ) -> int:
        """
        Replace text in the document.

        Args:
            pattern (str): Text pattern to search for
            replacement (str): Replacement text
            regex (bool): Whether to treat pattern as regex
            case_sensitive (bool): Whether to match case

        Returns:
            int: Number of replacements made
        """
        from .search import replace_text
        return replace_text(self, pattern, replacement, regex, case_sensitive)

    @property
    def paragraphs(self) -> List[Paragraph]:
        """Get all paragraphs in the document."""
        return [Paragraph(p) for p in self.doc.paragraphs]

    @property
    def tables(self) -> List[Table]:
        """Get all tables in the document."""
        return [Table(t) for t in self.doc.tables]

    @property
    def styles(self) -> Dict[str, Style]:
        """Get all styles defined in the document."""
        return {name: Style(style) for name, style in self.doc.styles.items()}

    @property
    def core_properties(self) -> Dict[str, Any]:
        """Get the document's core properties."""
        props = self.doc.core_properties
        return {
            "author": props.author,
            "category": props.category,
            "comments": props.comments,
            "content_status": props.content_status,
            "created": props.created,
            "identifier": props.identifier,
            "keywords": props.keywords,
            "language": props.language,
            "last_modified_by": props.last_modified_by,
            "last_printed": props.last_printed,
            "modified": props.modified,
            "revision": props.revision,
            "subject": props.subject,
            "title": props.title,
            "version": props.version,
        }

    def __iter__(self) -> Iterator[Union[Paragraph, Table]]:
        """Iterate over paragraphs and tables in document order."""
        for element in self.doc.element.body:
            if isinstance(element, CT_P):
                yield Paragraph(_Paragraph(element, self.doc))
            elif isinstance(element, _Table._tbl_elm):
                yield Table(_Table(element, self.doc))

    def __enter__(self) -> "Document":
        return self

    def __exit__(self, exc_type, exc_val, exc_tb) -> None:
        if self.path is not None:
            self.save() 