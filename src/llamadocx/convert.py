"""
Format conversion functions for LlamaDocx.

This module provides functions for converting between different document formats,
including DOCX, Markdown, HTML, and PDF.
"""

import os
import re
import logging
import subprocess
import tempfile
from pathlib import Path
from typing import Optional, Union, List, Dict, Any, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Inches

# Set up module-level logger
logger = logging.getLogger(__name__)


def convert_to_docx(
    source_path: Union[str, Path],
    output_path: Optional[Union[str, Path]] = None,
    source_format: Optional[str] = None,
    extra_args: Optional[List[str]] = None
) -> Path:
    """Convert a document to DOCX format.

    This function attempts to automatically detect the source format
    based on the file extension, or you can specify it explicitly.

    Args:
        source_path: Path to the source document
        output_path: Path for the output DOCX file. If None, uses the 
            source filename with .docx extension
        source_format: Format of the source document. If None, 
            detected from file extension
        extra_args: Additional arguments to pass to the conversion tool

    Returns:
        Path to the output DOCX file

    Raises:
        FileNotFoundError: If the source file doesn't exist
        ValueError: For unsupported formats or conversion errors
    """
    source_path = Path(source_path)
    if not source_path.exists():
        raise FileNotFoundError(f"Source file not found: {source_path}")

    # Determine source format if not specified
    if not source_format:
        source_format = _get_format_from_extension(source_path)
        if not source_format:
            raise ValueError(f"Unable to determine format for {source_path}. "
                             "Please specify source_format.")

    # Determine output path if not specified
    if not output_path:
        output_path = source_path.with_suffix('.docx')
    output_path = Path(output_path)

    # Create parent directory for output if it doesn't exist
    if not output_path.parent.exists():
        output_path.parent.mkdir(parents=True, exist_ok=True)

    # Convert based on source format
    if source_format.lower() == 'md' or source_format.lower() == 'markdown':
        return md_to_docx(source_path, output_path)
    elif source_format.lower() == 'html':
        return html_to_docx(source_path, output_path)
    else:
        # Try to use pandoc for other formats
        if has_pandoc():
            extra_args = extra_args or []
            try:
                import pypandoc
                pypandoc.convert_file(
                    str(source_path),
                    'docx',
                    outputfile=str(output_path),
                    extra_args=extra_args
                )
                return output_path
            except Exception as e:
                raise ValueError(f"Pandoc conversion failed: {e}")
        else:
            raise ValueError(f"Conversion from {source_format} to DOCX not supported. "
                           "Please install pandoc for additional format support.")


def convert_from_docx(
    source_path: Union[str, Path],
    output_path: Optional[Union[str, Path]] = None,
    output_format: str = "markdown",
    extra_args: Optional[List[str]] = None
) -> Path:
    """Convert a DOCX document to another format.

    Args:
        source_path: Path to the source DOCX document
        output_path: Path for the output file. If None, uses the source
            filename with the appropriate extension
        output_format: Format for the output document (default: markdown)
        extra_args: Additional arguments to pass to the conversion tool

    Returns:
        Path to the output file

    Raises:
        FileNotFoundError: If the source file doesn't exist
        ValueError: For unsupported formats or conversion errors
    """
    source_path = Path(source_path)
    if not source_path.exists():
        raise FileNotFoundError(f"Source file not found: {source_path}")

    # Check if file is a DOCX file
    if source_path.suffix.lower() != '.docx':
        raise ValueError(f"Source file must be a DOCX file: {source_path}")

    # Normalize output format
    output_format = output_format.lower()
    
    # Determine output path if not specified
    if not output_path:
        extension = _get_extension_for_format(output_format)
        if not extension:
            raise ValueError(f"Unsupported output format: {output_format}")
        output_path = source_path.with_suffix(extension)
    output_path = Path(output_path)

    # Create parent directory for output if it doesn't exist
    if not output_path.parent.exists():
        output_path.parent.mkdir(parents=True, exist_ok=True)

    # Convert based on target format
    if output_format in ('md', 'markdown'):
        return docx_to_markdown(source_path, output_path)
    elif output_format == 'html':
        return docx_to_html(source_path, output_path)
    elif output_format == 'pdf':
        docx_to_pdf(source_path, output_path)
        return output_path
    else:
        # Try to use pandoc for other formats
        if has_pandoc():
            extra_args = extra_args or []
            try:
                import pypandoc
                pypandoc.convert_file(
                    str(source_path),
                    output_format,
                    outputfile=str(output_path),
                    extra_args=extra_args
                )
                return output_path
            except Exception as e:
                raise ValueError(f"Pandoc conversion failed: {e}")
        else:
            raise ValueError(f"Conversion from DOCX to {output_format} not supported. "
                           "Please install pandoc for additional format support.")


def _get_format_from_extension(path: Path) -> Optional[str]:
    """Get document format from file extension.

    Args:
        path: Path to the document

    Returns:
        Format string or None if unknown extension
    """
    extension = path.suffix.lower().lstrip('.')
    
    # Map extensions to formats
    format_map = {
        'docx': 'docx',
        'md': 'markdown',
        'markdown': 'markdown',
        'html': 'html',
        'htm': 'html',
        'txt': 'text',
        'pdf': 'pdf',
        'rtf': 'rtf',
        'odt': 'odt',
        'tex': 'latex',
    }
    
    return format_map.get(extension)


def _get_extension_for_format(format: str) -> Optional[str]:
    """Get file extension for document format.

    Args:
        format: Document format

    Returns:
        File extension (with dot) or None if unknown format
    """
    # Map formats to extensions
    extension_map = {
        'docx': '.docx',
        'markdown': '.md',
        'md': '.md',
        'html': '.html',
        'text': '.txt',
        'pdf': '.pdf',
        'rtf': '.rtf',
        'odt': '.odt',
        'latex': '.tex',
    }
    
    return extension_map.get(format.lower())


def has_pandoc() -> bool:
    """Check if pandoc is installed and available.

    Returns:
        True if pandoc is available, False otherwise
    """
    try:
        import pypandoc
        return True
    except ImportError:
        return False
    
    # Also check for pandoc executable
    try:
        result = subprocess.run(
            ["pandoc", "--version"], 
            stdout=subprocess.PIPE, 
            stderr=subprocess.PIPE
        )
        return result.returncode == 0
    except (FileNotFoundError, subprocess.SubprocessError):
        return False


def get_pandoc_version() -> Optional[str]:
    """Get the installed pandoc version.

    Returns:
        Pandoc version string or None if not installed
    """
    if not has_pandoc():
        return None
    
    try:
        result = subprocess.run(
            ["pandoc", "--version"], 
            stdout=subprocess.PIPE, 
            stderr=subprocess.PIPE,
            text=True
        )
        if result.returncode == 0:
            version_match = re.search(r"pandoc\s+(\d+\.\d+(?:\.\d+)?)", result.stdout)
            if version_match:
                return version_match.group(1)
        return "Unknown"
    except (FileNotFoundError, subprocess.SubprocessError):
        return None


def get_supported_formats() -> tuple:
    """Get lists of supported input and output formats.

    Returns:
        Tuple of (input_formats, output_formats)
    """
    # Basic formats always supported
    input_formats = ["docx", "markdown", "html"]
    output_formats = ["docx", "markdown", "html", "pdf"]
    
    # Check for pandoc to add more formats
    if has_pandoc():
        try:
            import pypandoc
            return (pypandoc.get_pandoc_formats()[0], pypandoc.get_pandoc_formats()[1])
        except Exception:
            pass
    
    return (input_formats, output_formats)


def md_to_docx(
    md_path: Union[str, Path],
    output_path: Union[str, Path],
    return_content: bool = False
) -> Optional[Document]:
    """Convert a Markdown file to DOCX format.

    Args:
        md_path: Path to the Markdown file
        output_path: Path for the output DOCX file
        return_content: If True, return the Document object

    Returns:
        If return_content is True, returns the Document object,
        otherwise returns None

    Raises:
        FileNotFoundError: If the Markdown file doesn't exist
        ImportError: If the required libraries aren't installed
    """
    try:
        import markdown
    except ImportError:
        raise ImportError("The 'markdown' package is required for Markdown conversion. "
                          "Install it with 'pip install markdown'.")

    md_path = Path(md_path)
    if not md_path.exists():
        raise FileNotFoundError(f"Markdown file not found: {md_path}")
    
    output_path = Path(output_path)
    
    # Read markdown content
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Convert to HTML
    html_content = markdown.markdown(
        md_content,
        extensions=[
            'markdown.extensions.extra',
            'markdown.extensions.nl2br',
            'markdown.extensions.sane_lists',
            'markdown.extensions.smarty',
            'markdown.extensions.tables',
        ]
    )
    
    # Create a temporary HTML file
    with tempfile.NamedTemporaryFile(suffix='.html', delete=False, mode='w', encoding='utf-8') as f:
        f.write(html_content)
        temp_html = f.name
    
    try:
        # Convert HTML to DOCX
        doc = html_to_docx(temp_html, output_path, return_content=True)
    finally:
        # Remove the temporary HTML file
        try:
            os.unlink(temp_html)
        except Exception:
            pass
    
    if return_content:
        return doc
    return None


def html_to_docx(
    html_path: Union[str, Path],
    output_path: Union[str, Path],
    return_content: bool = False
) -> Optional[Document]:
    """Convert an HTML file to DOCX format.

    Args:
        html_path: Path to the HTML file
        output_path: Path for the output DOCX file
        return_content: If True, return the Document object

    Returns:
        If return_content is True, returns the Document object,
        otherwise returns None

    Raises:
        FileNotFoundError: If the HTML file doesn't exist
        ImportError: If the required libraries aren't installed
    """
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError("The 'beautifulsoup4' package is required for HTML conversion. "
                          "Install it with 'pip install beautifulsoup4'.")

    html_path = Path(html_path)
    if not html_path.exists():
        raise FileNotFoundError(f"HTML file not found: {html_path}")
    
    output_path = Path(output_path)
    
    # Read HTML content
    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Create a new Word document
    doc = Document()
    
    # Process HTML content
    for element in soup.body.children:
        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'h4':
            doc.add_heading(element.get_text(), level=4)
        elif element.name == 'h5':
            doc.add_heading(element.get_text(), level=5)
        elif element.name == 'h6':
            doc.add_heading(element.get_text(), level=6)
        elif element.name == 'p':
            p = doc.add_paragraph()
            _process_html_element(element, p)
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet')
                _process_html_element(li, p)
        elif element.name == 'ol':
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Number')
                _process_html_element(li, p)
        elif element.name == 'table':
            # Count rows and columns
            rows = element.find_all('tr')
            if rows:
                # Determine the maximum number of columns
                max_cols = 0
                for row in rows:
                    cols = row.find_all(['td', 'th'])
                    max_cols = max(max_cols, len(cols))
                
                # Create table
                if max_cols > 0:
                    table = doc.add_table(rows=len(rows), cols=max_cols)
                    
                    # Fill table
                    for i, row in enumerate(rows):
                        cells = row.find_all(['td', 'th'])
                        for j, cell in enumerate(cells):
                            if j < max_cols:  # Ensure we don't exceed the table dimensions
                                table.cell(i, j).text = cell.get_text().strip()
        elif element.name == 'pre':
            # Handle code blocks
            code = element.get_text()
            p = doc.add_paragraph()
            
            # Add code with Courier New font
            code_run = p.add_run(code)
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(9)
    
    # Save document
    doc.save(output_path)
    
    if return_content:
        return doc
    return None


def _process_html_element(element, paragraph):
    """Process an HTML element and add its content to a paragraph.

    Args:
        element: BeautifulSoup element
        paragraph: docx Paragraph object
    """
    # If the element is just text, add it directly
    if isinstance(element, str):
        paragraph.add_run(element)
        return
    
    # Process the content of the element
    for content in element.contents:
        # If it's a string, add it as text
        if isinstance(content, str):
            paragraph.add_run(content)
        # If it's a tag, process it based on its type
        elif hasattr(content, 'name'):
            if content.name == 'strong' or content.name == 'b':
                run = paragraph.add_run(content.get_text())
                run.bold = True
            elif content.name == 'em' or content.name == 'i':
                run = paragraph.add_run(content.get_text())
                run.italic = True
            elif content.name == 'u':
                run = paragraph.add_run(content.get_text())
                run.underline = True
            elif content.name == 'code':
                run = paragraph.add_run(content.get_text())
                run.font.name = 'Courier New'
            elif content.name == 'a':
                # Add hyperlink
                run = paragraph.add_run(content.get_text())
                run.underline = True
                run.font.color.rgb = (0, 0, 255)  # Blue color
            elif content.name in ['span', 'p', 'div']:
                # Recursively process the content
                _process_html_element(content, paragraph)
            else:
                # For unhandled tags, just get the text
                paragraph.add_run(content.get_text())


def docx_to_pdf(
    docx_path: Union[str, Path],
    output_path: Union[str, Path]
) -> None:
    """Convert a DOCX file to PDF format.

    This function attempts to use various methods for conversion:
    1. docx2pdf (if installed)
    2. LibreOffice/OpenOffice command-line conversion (if available)
    3. Microsoft Word COM automation (Windows only)

    Args:
        docx_path: Path to the DOCX file
        output_path: Path for the output PDF file

    Raises:
        FileNotFoundError: If the DOCX file doesn't exist
        ImportError: If no suitable conversion method is available
        RuntimeError: If conversion fails
    """
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")
    
    output_path = Path(output_path)
    
    # Method 1: Try docx2pdf
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(output_path))
        logger.debug(f"Converted {docx_path} to {output_path} using docx2pdf")
        return
    except ImportError:
        logger.debug("docx2pdf not installed, trying alternative methods")
    except Exception as e:
        logger.debug(f"docx2pdf conversion failed: {e}")
    
    # Method 2: Try LibreOffice/OpenOffice
    try:
        # Check for LibreOffice/OpenOffice
        libreoffice_paths = [
            "libreoffice",
            "soffice",
            "/usr/bin/libreoffice",
            "/usr/bin/soffice",
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
            "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
        ]
        
        office_path = None
        for path in libreoffice_paths:
            try:
                subprocess.run([path, "--version"], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                office_path = path
                break
            except (FileNotFoundError, subprocess.SubprocessError):
                continue
        
        if office_path:
            cmd = [
                office_path,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(output_path.parent),
                str(docx_path)
            ]
            
            process = subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True
            )
            
            if process.returncode == 0:
                # LibreOffice saves the PDF with the same name in the specified output directory
                generated_pdf = output_path.parent / f"{docx_path.stem}.pdf"
                if generated_pdf.exists() and generated_pdf != output_path:
                    os.rename(generated_pdf, output_path)
                logger.debug(f"Converted {docx_path} to {output_path} using LibreOffice")
                return
            else:
                logger.debug(f"LibreOffice conversion failed: {process.stderr}")
    except Exception as e:
        logger.debug(f"LibreOffice conversion attempt failed: {e}")
    
    # Method 3: Try Microsoft Word COM automation (Windows only)
    if os.name == 'nt':  # Windows only
        try:
            import win32com.client
            
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            wb = word.Documents.Open(str(docx_path.absolute()))
            wb.SaveAs(str(output_path.absolute()), FileFormat=17)  # 17 = PDF
            wb.Close()
            word.Quit()
            
            logger.debug(f"Converted {docx_path} to {output_path} using Microsoft Word")
            return
        except Exception as e:
            logger.debug(f"Microsoft Word conversion attempt failed: {e}")
    
    # If all methods failed
    raise RuntimeError(f"PDF conversion failed. Please install docx2pdf with 'pip install docx2pdf' "
                      "or ensure LibreOffice or Microsoft Word is available.")


def docx_to_markdown(
    docx_path: Union[str, Path],
    output_path: Optional[Union[str, Path]] = None,
    return_content: bool = False
) -> Optional[str]:
    """Convert a DOCX file to Markdown format.

    Args:
        docx_path: Path to the DOCX file
        output_path: Path for the output Markdown file.
            If None, content is not saved to a file.
        return_content: If True, return the Markdown content as a string

    Returns:
        If return_content is True, returns the Markdown content as a string,
        otherwise returns the output Path object

    Raises:
        FileNotFoundError: If the DOCX file doesn't exist
        ImportError: If no suitable conversion method is available
    """
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")
    
    if output_path:
        output_path = Path(output_path)
    
    # Try to use pandoc (recommended)
    if has_pandoc():
        try:
            import pypandoc
            md_content = pypandoc.convert_file(
                str(docx_path),
                'markdown',
                extra_args=['--wrap=none', '--atx-headers']
            )
            
            if return_content:
                return md_content
            
            if output_path:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(md_content)
                return output_path
        except Exception as e:
            logger.warning(f"Pandoc conversion failed: {e}")
    
    # Fallback to basic conversion using python-docx
    doc = Document(docx_path)
    
    # Convert document content to Markdown
    md_lines = []
    
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            md_lines.append("")
            continue
        
        # Check if it's a heading
        if paragraph.style.name.startswith('Heading'):
            level = int(paragraph.style.name[-1]) if paragraph.style.name[-1].isdigit() else 1
            md_lines.append(f"{'#' * level} {paragraph.text}")
        else:
            # Handle text formatting
            formatted_text = ""
            for run in paragraph.runs:
                text = run.text
                if run.bold:
                    text = f"**{text}**"
                if run.italic:
                    text = f"*{text}*"
                if run.underline:
                    text = f"_{text}_"
                formatted_text += text
            
            md_lines.append(formatted_text)
    
    # Add tables
    for table in doc.tables:
        # Add separator line
        md_lines.append("")
        
        # Create table header
        header_row = table.rows[0]
        md_lines.append("| " + " | ".join(cell.text for cell in header_row.cells) + " |")
        
        # Add separator line
        md_lines.append("| " + " | ".join(["---"] * len(header_row.cells)) + " |")
        
        # Add data rows
        for row in table.rows[1:]:
            md_lines.append("| " + " | ".join(cell.text for cell in row.cells) + " |")
    
    md_content = "\n".join(md_lines)
    
    if return_content:
        return md_content
    
    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        return output_path


def docx_to_html(
    docx_path: Union[str, Path],
    output_path: Optional[Union[str, Path]] = None,
    return_content: bool = False
) -> Optional[str]:
    """Convert a DOCX file to HTML format.

    Args:
        docx_path: Path to the DOCX file
        output_path: Path for the output HTML file.
            If None, content is not saved to a file.
        return_content: If True, return the HTML content as a string

    Returns:
        If return_content is True, returns the HTML content as a string,
        otherwise returns the output Path object

    Raises:
        FileNotFoundError: If the DOCX file doesn't exist
        ImportError: If no suitable conversion method is available
    """
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")
    
    if output_path:
        output_path = Path(output_path)
    
    # Try to use pandoc (recommended)
    if has_pandoc():
        try:
            import pypandoc
            html_content = pypandoc.convert_file(
                str(docx_path),
                'html'
            )
            
            if return_content:
                return html_content
            
            if output_path:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                return output_path
        except Exception as e:
            logger.warning(f"Pandoc conversion failed: {e}")
    
    # Fallback to basic conversion using python-docx
    doc = Document(docx_path)
    
    # Start HTML document
    html_lines = [
        "<!DOCTYPE html>",
        "<html>",
        "<head>",
        f"<title>{docx_path.stem}</title>",
        "<style>",
        "body { font-family: Arial, sans-serif; line-height: 1.6; margin: 20px; }",
        "h1, h2, h3, h4, h5, h6 { color: #333; }",
        "table { border-collapse: collapse; width: 100%; }",
        "table, th, td { border: 1px solid #ddd; padding: 8px; }",
        "th { background-color: #f2f2f2; }",
        "</style>",
        "</head>",
        "<body>"
    ]
    
    # Convert paragraphs to HTML
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            html_lines.append("<p>&nbsp;</p>")
            continue
        
        # Check if it's a heading
        if paragraph.style.name.startswith('Heading'):
            level = int(paragraph.style.name[-1]) if paragraph.style.name[-1].isdigit() else 1
            html_lines.append(f"<h{level}>{paragraph.text}</h{level}>")
        else:
            # Handle text formatting
            html_parts = []
            for run in paragraph.runs:
                text = run.text
                if run.bold:
                    text = f"<strong>{text}</strong>"
                if run.italic:
                    text = f"<em>{text}</em>"
                if run.underline:
                    text = f"<u>{text}</u>"
                html_parts.append(text)
            
            html_lines.append(f"<p>{''.join(html_parts)}</p>")
    
    # Add tables
    for table in doc.tables:
        html_lines.append("<table>")
        
        for i, row in enumerate(table.rows):
            html_lines.append("<tr>")
            
            for cell in row.cells:
                # Use th for header row
                tag = "th" if i == 0 else "td"
                html_lines.append(f"<{tag}>{cell.text}</{tag}>")
            
            html_lines.append("</tr>")
        
        html_lines.append("</table>")
    
    # Close HTML document
    html_lines.append("</body>")
    html_lines.append("</html>")
    
    html_content = "\n".join(html_lines)
    
    if return_content:
        return html_content
    
    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        return output_path 