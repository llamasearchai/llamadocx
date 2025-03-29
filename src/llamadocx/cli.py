#!/usr/bin/env python
"""
Command Line Interface for LlamaDocx.

This module provides a command-line interface for the LlamaDocx toolkit,
allowing users to perform operations like document creation, conversion,
template processing, and more without writing Python code.
"""

import os
import sys
import json
import logging
import argparse
from pathlib import Path
from typing import Dict, Any, List, Optional, Union

from docx import Document

from llamadocx import __version__
from llamadocx.convert import (
    convert_to_docx, convert_from_docx, docx_to_pdf, docx_to_markdown,
    docx_to_html, md_to_docx, html_to_docx, get_supported_formats,
    has_pandoc, get_pandoc_version
)
from llamadocx.template import Template
from llamadocx.metadata import (
    get_metadata, set_metadata, update_metadata_from_file,
    extract_metadata_to_file
)
from llamadocx.search import search_text, replace_text

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


def cli():
    """Execute the command line interface for LlamaDocx."""
    parser = argparse.ArgumentParser(
        description="LlamaDocx - A toolkit for working with Microsoft Word documents",
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    
    parser.add_argument(
        "--version", action="version", 
        version=f"LlamaDocx {__version__}"
    )
    
    parser.add_argument(
        "--verbose", "-v", action="store_true",
        help="Enable verbose output"
    )
    
    subparsers = parser.add_subparsers(dest="command", help="Commands")
    
    # Create command
    create_parser = subparsers.add_parser(
        "create", help="Create a new Word document"
    )
    create_parser.add_argument(
        "output", help="Path to save the new document"
    )
    create_parser.add_argument(
        "--title", help="Document title"
    )
    create_parser.add_argument(
        "--author", help="Document author"
    )
    create_parser.add_argument(
        "--subject", help="Document subject"
    )
    create_parser.add_argument(
        "--template", help="Template file to use as a base"
    )
    
    # Extract command
    extract_parser = subparsers.add_parser(
        "extract", help="Extract content from a Word document"
    )
    extract_parser.add_argument(
        "input", help="Path to the input document"
    )
    extract_parser.add_argument(
        "output", help="Path to save the extracted content"
    )
    extract_parser.add_argument(
        "--format", choices=["txt", "json", "md", "html"],
        default="txt", help="Output format (default: txt)"
    )
    extract_parser.add_argument(
        "--images", action="store_true",
        help="Extract images from the document"
    )
    extract_parser.add_argument(
        "--tables", action="store_true",
        help="Extract tables from the document to CSV files"
    )
    
    # Convert command
    convert_parser = subparsers.add_parser(
        "convert", help="Convert between document formats"
    )
    convert_parser.add_argument(
        "input", help="Path to the input document"
    )
    convert_parser.add_argument(
        "output", help="Path to save the converted document"
    )
    convert_parser.add_argument(
        "--from-format", help="Input format (if not detected from extension)"
    )
    convert_parser.add_argument(
        "--to-format", help="Output format (if not detected from extension)"
    )
    convert_parser.add_argument(
        "--list-formats", action="store_true",
        help="List supported formats and exit"
    )
    
    # Search command
    search_parser = subparsers.add_parser(
        "search", help="Search content in a Word document"
    )
    search_parser.add_argument(
        "input", help="Path to the input document"
    )
    search_parser.add_argument(
        "query", help="Text to search for"
    )
    search_parser.add_argument(
        "--regex", action="store_true",
        help="Use regular expressions for search"
    )
    search_parser.add_argument(
        "--case-sensitive", action="store_true",
        help="Make search case-sensitive"
    )
    search_parser.add_argument(
        "--output", help="Save search results to a file"
    )
    search_parser.add_argument(
        "--format", choices=["txt", "json"],
        default="txt", help="Output format (default: txt)"
    )
    
    # Replace command
    replace_parser = subparsers.add_parser(
        "replace", help="Replace content in a Word document"
    )
    replace_parser.add_argument(
        "input", help="Path to the input document"
    )
    replace_parser.add_argument(
        "output", help="Path to save the modified document"
    )
    replace_parser.add_argument(
        "pattern", help="Text pattern to replace"
    )
    replace_parser.add_argument(
        "replacement", help="Replacement text"
    )
    replace_parser.add_argument(
        "--regex", action="store_true",
        help="Use regular expressions for pattern matching"
    )
    replace_parser.add_argument(
        "--case-sensitive", action="store_true",
        help="Make replacements case-sensitive"
    )
    replace_parser.add_argument(
        "--count", action="store_true",
        help="Just count replacements without making changes"
    )
    
    # Template command
    template_parser = subparsers.add_parser(
        "template", help="Process a Word document template"
    )
    template_parser.add_argument(
        "template", help="Path to the template document"
    )
    template_parser.add_argument(
        "output", help="Path to save the processed document"
    )
    template_group = template_parser.add_mutually_exclusive_group(required=True)
    template_group.add_argument(
        "--data", help="JSON string with template data"
    )
    template_group.add_argument(
        "--data-file", help="Path to JSON file with template data"
    )
    template_parser.add_argument(
        "--list-fields", action="store_true",
        help="List all fields in the template and exit"
    )
    template_parser.add_argument(
        "--delimiters", nargs=2, metavar=("START", "END"),
        default=[r"{{", r"}}"],
        help="Field delimiters (default: {{ }})"
    )
    
    # Metadata command
    meta_parser = subparsers.add_parser(
        "meta", help="View or modify document metadata"
    )
    meta_parser.add_argument(
        "document", help="Path to the document"
    )
    meta_group = meta_parser.add_mutually_exclusive_group(required=True)
    meta_group.add_argument(
        "--get", action="store_true",
        help="Get document metadata"
    )
    meta_group.add_argument(
        "--set", metavar="KEY=VALUE",
        nargs="+", help="Set metadata properties (e.g., --set title='My Doc' author='Me')"
    )
    meta_group.add_argument(
        "--from-file", metavar="JSON_FILE",
        help="Update metadata from a JSON file"
    )
    meta_group.add_argument(
        "--to-file", metavar="JSON_FILE",
        help="Extract metadata to a JSON file"
    )
    
    # Parse arguments
    args = parser.parse_args()
    
    # Configure logging based on verbosity
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    # If no command is specified, show help
    if not args.command:
        parser.print_help()
        return
    
    # Execute the requested command
    try:
        if args.command == "create":
            cmd_create(args)
        elif args.command == "extract":
            cmd_extract(args)
        elif args.command == "convert":
            cmd_convert(args)
        elif args.command == "search":
            cmd_search(args)
        elif args.command == "replace":
            cmd_replace(args)
        elif args.command == "template":
            cmd_template(args)
        elif args.command == "meta":
            cmd_metadata(args)
    except Exception as e:
        logger.error(f"Error: {e}")
        if args.verbose:
            logger.exception("Detailed error information:")
        sys.exit(1)


def cmd_create(args):
    """Handle the 'create' command."""
    output_path = Path(args.output)
    
    # Create document from template or blank
    if args.template:
        template_path = Path(args.template)
        if not template_path.exists():
            raise FileNotFoundError(f"Template file not found: {template_path}")
        
        # Copy from template
        doc = Document(template_path)
        logger.info(f"Created document from template: {template_path}")
    else:
        # Create a blank document
        doc = Document()
        
        # Add a title if specified
        if args.title:
            doc.add_heading(args.title, level=1)
        
        logger.info("Created a blank document")
    
    # Set metadata if specified
    metadata = {}
    if args.title:
        metadata["title"] = args.title
    if args.author:
        metadata["author"] = args.author
    if args.subject:
        metadata["subject"] = args.subject
    
    if metadata:
        set_metadata(doc, metadata)
        properties = ", ".join(f"{k}='{v}'" for k, v in metadata.items())
        logger.info(f"Set document properties: {properties}")
    
    # Save the document
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    logger.info(f"Saved document to: {output_path}")


def cmd_extract(args):
    """Handle the 'extract' command."""
    input_path = Path(args.input)
    output_path = Path(args.output)
    
    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")
    
    # Create output directory if needed
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Load the document
    doc = Document(input_path)
    logger.info(f"Loaded document: {input_path}")
    
    # Extract content based on format
    if args.format == "txt":
        # Extract as plain text
        text_content = []
        for para in doc.paragraphs:
            text_content.append(para.text)
        
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n\n".join(text_content))
        
        logger.info(f"Extracted text content to: {output_path}")
    
    elif args.format == "json":
        # Extract as JSON with structure
        content = {
            "metadata": get_metadata(doc),
            "paragraphs": [para.text for para in doc.paragraphs if para.text.strip()],
            "tables": []
        }
        
        # Extract tables if requested
        if args.tables:
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                content["tables"].append(table_data)
        
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(content, f, indent=2)
        
        logger.info(f"Extracted content as JSON to: {output_path}")
    
    elif args.format == "md":
        # Extract as Markdown
        md_content = docx_to_markdown(input_path, return_content=True)
        
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(md_content)
        
        logger.info(f"Extracted content as Markdown to: {output_path}")
    
    elif args.format == "html":
        # Extract as HTML
        html_content = docx_to_html(input_path, return_content=True)
        
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        
        logger.info(f"Extracted content as HTML to: {output_path}")
    
    # Extract images if requested
    if args.images:
        image_dir = output_path.parent / f"{output_path.stem}_images"
        image_dir.mkdir(exist_ok=True)
        
        image_count = 0
        # Extract images from the document
        for rel in doc.part.rels.values():
            if rel.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image':
                image_count += 1
                image_data = rel.target_part.blob
                image_ext = rel.target_ref.split('.')[-1]
                image_path = image_dir / f"image_{image_count}.{image_ext}"
                
                with open(image_path, "wb") as f:
                    f.write(image_data)
        
        logger.info(f"Extracted {image_count} images to: {image_dir}")
    
    # Extract tables as CSV if requested
    if args.tables and args.format != "json":
        import csv
        
        table_dir = output_path.parent / f"{output_path.stem}_tables"
        table_dir.mkdir(exist_ok=True)
        
        for i, table in enumerate(doc.tables):
            table_path = table_dir / f"table_{i+1}.csv"
            
            with open(table_path, "w", newline="", encoding="utf-8") as f:
                writer = csv.writer(f)
                for row in table.rows:
                    writer.writerow([cell.text for cell in row.cells])
        
        logger.info(f"Extracted {len(doc.tables)} tables to: {table_dir}")


def cmd_convert(args):
    """Handle the 'convert' command."""
    # List supported formats if requested
    if args.list_formats:
        input_formats, output_formats = get_supported_formats()
        
        print("Supported input formats:")
        for fmt in sorted(input_formats):
            print(f"  - {fmt}")
        
        print("\nSupported output formats:")
        for fmt in sorted(output_formats):
            print(f"  - {fmt}")
        
        print(f"\nPandoc {'installed' if has_pandoc() else 'not installed'}")
        if has_pandoc():
            print(f"Pandoc version: {get_pandoc_version()}")
        
        return
    
    input_path = Path(args.input)
    output_path = Path(args.output)
    
    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")
    
    # Determine input and output formats
    input_format = args.from_format
    output_format = args.to_format
    
    # Guess formats from file extensions if not specified
    if not input_format:
        input_format = input_path.suffix.lstrip('.').lower()
    
    if not output_format:
        output_format = output_path.suffix.lstrip('.').lower()
    
    # Perform the conversion based on formats
    if output_format == "docx":
        # Convert to DOCX
        result = convert_to_docx(input_path, output_path, source_format=input_format)
        logger.info(f"Converted {input_path} to DOCX: {result}")
    elif input_format == "docx":
        # Convert from DOCX
        result = convert_from_docx(input_path, output_path, output_format=output_format)
        logger.info(f"Converted DOCX to {output_format}: {result}")
    else:
        # Two-step conversion via docx
        logger.info(f"Performing two-step conversion from {input_format} to {output_format}")
        temp_path = output_path.parent / f"temp_{output_path.stem}.docx"
        try:
            # First convert to docx
            convert_to_docx(input_path, temp_path, source_format=input_format)
            # Then convert to target format
            result = convert_from_docx(temp_path, output_path, output_format=output_format)
            logger.info(f"Converted {input_path} to {output_format}: {result}")
        finally:
            # Clean up temporary file
            if temp_path.exists():
                temp_path.unlink()


def cmd_search(args):
    """Handle the 'search' command."""
    input_path = Path(args.input)
    
    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")
    
    # Load the document
    doc = Document(input_path)
    
    # Perform the search
    results = search_text(
        doc, 
        args.query, 
        regex=args.regex, 
        case_sensitive=args.case_sensitive
    )
    
    # Display or save results
    if args.format == "txt":
        # Format as plain text
        output_lines = [
            f"Search results for '{args.query}' in {input_path}:",
            f"Found {len(results)} matches.\n"
        ]
        
        for i, result in enumerate(results, 1):
            output_lines.append(f"Match {i}:")
            output_lines.append(f"  Location: {result['type']} {result['index']}")
            output_lines.append(f"  Text: {result['text']}")
            if 'context' in result:
                output_lines.append(f"  Context: ...{result['context']}...")
            output_lines.append("")
        
        output_text = "\n".join(output_lines)
        
        if args.output:
            with open(args.output, "w", encoding="utf-8") as f:
                f.write(output_text)
            logger.info(f"Search results saved to: {args.output}")
        else:
            print(output_text)
    
    elif args.format == "json":
        # Format as JSON
        output_data = {
            "query": args.query,
            "document": str(input_path),
            "count": len(results),
            "matches": results
        }
        
        if args.output:
            with open(args.output, "w", encoding="utf-8") as f:
                json.dump(output_data, f, indent=2)
            logger.info(f"Search results saved to: {args.output}")
        else:
            print(json.dumps(output_data, indent=2))


def cmd_replace(args):
    """Handle the 'replace' command."""
    input_path = Path(args.input)
    output_path = Path(args.output)
    
    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")
    
    # Load the document
    doc = Document(input_path)
    
    # Perform replacements or count matches
    if args.count:
        # Just count matches without replacing
        results = search_text(
            doc, 
            args.pattern, 
            regex=args.regex, 
            case_sensitive=args.case_sensitive
        )
        count = len(results)
        print(f"Found {count} matches for '{args.pattern}' in {input_path}")
    else:
        # Replace text and save
        count = replace_text(
            doc, 
            args.pattern, 
            args.replacement, 
            regex=args.regex, 
            case_sensitive=args.case_sensitive
        )
        
        # Save the modified document
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        
        logger.info(f"Replaced {count} occurrences of '{args.pattern}' with '{args.replacement}'")
        logger.info(f"Saved modified document to: {output_path}")


def cmd_template(args):
    """Handle the 'template' command."""
    template_path = Path(args.template)
    output_path = Path(args.output)
    
    if not template_path.exists():
        raise FileNotFoundError(f"Template file not found: {template_path}")
    
    # Create Template object
    template = Template(template_path, field_delimiters=args.delimiters)
    
    # List fields if requested
    if args.list_fields:
        fields = template.get_fields()
        print(f"Fields in template {template_path}:")
        for field in sorted(fields):
            print(f"  - {field}")
        return
    
    # Load template data
    if args.data_file:
        data_path = Path(args.data_file)
        if not data_path.exists():
            raise FileNotFoundError(f"Data file not found: {data_path}")
        
        with open(data_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        
        logger.info(f"Loaded template data from: {data_path}")
    else:
        data = json.loads(args.data)
        logger.info("Loaded template data from command line")
    
    # Process the template
    template.merge_fields(data)
    
    # Save the processed document
    output_path.parent.mkdir(parents=True, exist_ok=True)
    template.save(output_path)
    
    logger.info(f"Saved processed template to: {output_path}")


def cmd_metadata(args):
    """Handle the 'meta' command."""
    doc_path = Path(args.document)
    
    if not doc_path.exists():
        raise FileNotFoundError(f"Document not found: {doc_path}")
    
    # Load the document
    doc = Document(doc_path)
    
    # Process based on operation
    if args.get:
        # Get metadata
        metadata = get_metadata(doc)
        print(json.dumps(metadata, indent=2))
    
    elif args.set:
        # Set metadata properties
        properties = {}
        for prop in args.set:
            if "=" not in prop:
                logger.warning(f"Ignoring invalid property format: {prop}")
                continue
            
            key, value = prop.split("=", 1)
            # Remove quotes if present
            value = value.strip("'\"")
            properties[key] = value
        
        set_metadata(doc, properties)
        doc.save(doc_path)
        
        logger.info(f"Updated metadata for: {doc_path}")
        logger.info(f"Properties set: {', '.join(properties.keys())}")
    
    elif args.from_file:
        # Update metadata from file
        file_path = Path(args.from_file)
        if not file_path.exists():
            raise FileNotFoundError(f"Metadata file not found: {file_path}")
        
        update_metadata_from_file(doc, file_path)
        doc.save(doc_path)
        
        logger.info(f"Updated metadata for {doc_path} from {file_path}")
    
    elif args.to_file:
        # Extract metadata to file
        file_path = Path(args.to_file)
        file_path.parent.mkdir(parents=True, exist_ok=True)
        
        extract_metadata_to_file(doc, file_path)
        
        logger.info(f"Extracted metadata from {doc_path} to {file_path}")


if __name__ == "__main__":
    cli() 