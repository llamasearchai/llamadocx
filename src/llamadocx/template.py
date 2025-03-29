"""
Classes for working with Word document templates.

This module provides classes for using Word documents as templates,
with support for field merging and template processing.
"""

import re
from pathlib import Path
from typing import Optional, Union, Dict, Any, List, Tuple

from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .document import Document
from .paragraph import Paragraph
from .table import Table, Cell


class Template:
    """
    A Word document template.

    This class provides functionality for using Word documents as templates,
    with support for field merging and template processing.

    Args:
        path (Union[str, Path]): Path to the template document
        field_delimiters (Tuple[str, str]): Field delimiter characters (default: {{ }})

    Attributes:
        doc (Document): The document object
        field_delimiters (Tuple[str, str]): Field delimiter characters
    """

    def __init__(
        self,
        path: Union[str, Path],
        field_delimiters: Tuple[str, str] = ("{{", "}}")
    ) -> None:
        self.doc = Document(path)
        self.field_delimiters = field_delimiters
        self._field_pattern = re.compile(
            f"{re.escape(field_delimiters[0])}\\s*([^}}]+?)\\s*{re.escape(field_delimiters[1])}"
        )

    def get_fields(self) -> List[str]:
        """
        Get all field names in the template.

        Returns:
            List[str]: List of field names
        """
        fields = set()
        for paragraph in self.doc.paragraphs:
            matches = self._field_pattern.finditer(paragraph.text)
            fields.update(match.group(1) for match in matches)

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = self._field_pattern.finditer(paragraph.text)
                        fields.update(match.group(1) for match in matches)

        return sorted(list(fields))

    def merge_fields(
        self,
        field_values: Dict[str, Any],
        remove_empty: bool = True
    ) -> None:
        """
        Merge field values into the template.

        Args:
            field_values (Dict[str, Any]): Dictionary of field names and values
            remove_empty (bool): Whether to remove empty fields
        """
        # Process paragraphs
        for paragraph in self.doc.paragraphs:
            self._merge_paragraph_fields(paragraph, field_values, remove_empty)

        # Process tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._merge_paragraph_fields(paragraph, field_values, remove_empty)

    def _merge_paragraph_fields(
        self,
        paragraph: Paragraph,
        field_values: Dict[str, Any],
        remove_empty: bool
    ) -> None:
        """
        Merge fields in a paragraph.

        Args:
            paragraph (Paragraph): The paragraph to process
            field_values (Dict[str, Any]): Dictionary of field values
            remove_empty (bool): Whether to remove empty fields
        """
        text = paragraph.text
        matches = list(self._field_pattern.finditer(text))

        if not matches:
            return

        # Clear existing runs
        paragraph.clear()

        last_end = 0
        for match in matches:
            # Add text before field
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])

            # Get field value
            field_name = match.group(1)
            field_value = field_values.get(field_name, "")

            # Add field value or keep field if empty and not removing
            if field_value or not remove_empty:
                if not field_value and not remove_empty:
                    field_value = match.group(0)
                paragraph.add_run(str(field_value))

            last_end = match.end()

        # Add remaining text
        if last_end < len(text):
            paragraph.add_run(text[last_end:])

    def add_field(
        self,
        field_name: str,
        paragraph: Optional[Paragraph] = None,
        cell: Optional[Cell] = None
    ) -> None:
        """
        Add a merge field to the document.

        Args:
            field_name (str): Name of the field
            paragraph (Paragraph, optional): Paragraph to add field to
            cell (Cell, optional): Cell to add field to

        Raises:
            ValueError: If neither paragraph nor cell is provided
        """
        if paragraph is None and cell is None:
            raise ValueError("Must provide either paragraph or cell")

        field_text = f"{self.field_delimiters[0]} {field_name} {self.field_delimiters[1]}"

        if paragraph is not None:
            paragraph.add_run(field_text)
        else:
            cell.add_paragraph().add_run(field_text)

    def add_repeating_section(
        self,
        section_name: str,
        content: Union[Paragraph, Table]
    ) -> None:
        """
        Add a repeating section to the template.

        Args:
            section_name (str): Name of the repeating section
            content (Union[Paragraph, Table]): Content to repeat
        """
        # Add section start marker
        start_marker = self.doc.add_paragraph()
        start_marker.add_run(f"{self.field_delimiters[0]} start_{section_name} {self.field_delimiters[1]}")

        # Add content
        if isinstance(content, Paragraph):
            self.doc.add_paragraph().text = content.text
        else:  # Table
            new_table = self.doc.add_table(len(content.rows), len(content.columns))
            for i, row in enumerate(content.rows):
                for j, cell in enumerate(row.cells):
                    new_table[i, j].text = cell.text

        # Add section end marker
        end_marker = self.doc.add_paragraph()
        end_marker.add_run(f"{self.field_delimiters[0]} end_{section_name} {self.field_delimiters[1]}")

    def merge_repeating_section(
        self,
        section_name: str,
        items: List[Dict[str, Any]]
    ) -> None:
        """
        Merge data into a repeating section.

        Args:
            section_name (str): Name of the repeating section
            items (List[Dict[str, Any]]): List of field values for each repetition
        """
        # Find section markers
        start_pattern = f"{self.field_delimiters[0]}\\s*start_{section_name}\\s*{self.field_delimiters[1]}"
        end_pattern = f"{self.field_delimiters[0]}\\s*end_{section_name}\\s*{self.field_delimiters[1]}"

        start_para = None
        end_para = None
        content_elements = []

        # Find section boundaries and content
        for element in self.doc:
            if isinstance(element, Paragraph):
                if re.search(start_pattern, element.text):
                    start_para = element
                    continue
                elif re.search(end_pattern, element.text):
                    end_para = element
                    break

            if start_para is not None and end_para is None:
                content_elements.append(element)

        if start_para is None or end_para is None:
            raise ValueError(f"Section '{section_name}' not found in template")

        # Remove original content
        for element in content_elements:
            if isinstance(element, Paragraph):
                self.doc.doc._body._body.remove(element.paragraph._p)
            else:  # Table
                self.doc.doc._body._body.remove(element.table._tbl)

        # Insert repeated content
        insert_idx = self.doc.doc._body._body.index(start_para.paragraph._p) + 1
        for item in items:
            for element in content_elements:
                if isinstance(element, Paragraph):
                    new_para = self.doc.add_paragraph()
                    new_para.text = element.text
                    self._merge_paragraph_fields(new_para, item, True)
                    self.doc.doc._body._body.insert(insert_idx, new_para.paragraph._p)
                else:  # Table
                    new_table = self.doc.add_table(len(element.rows), len(element.columns))
                    for i, row in enumerate(element.rows):
                        for j, cell in enumerate(row.cells):
                            new_cell = new_table[i, j]
                            new_cell.text = cell.text
                            for para in new_cell.paragraphs:
                                self._merge_paragraph_fields(para, item, True)
                    self.doc.doc._body._body.insert(insert_idx, new_table.table._tbl)
                insert_idx += 1

    def save(self, path: Optional[Union[str, Path]] = None) -> None:
        """
        Save the processed template.

        Args:
            path (Union[str, Path, None]): Path to save the document to
        """
        self.doc.save(path) 