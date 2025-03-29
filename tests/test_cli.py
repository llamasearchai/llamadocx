"""
CLI tests for LlamaDocx.

This module contains tests for the command-line interface of the LlamaDocx package.
"""

import os
import json
import tempfile
import subprocess
from pathlib import Path
import pytest
from docx import Document


@pytest.fixture
def temp_docx():
    """Create a temporary DOCX file for testing CLI commands."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        doc.add_heading('CLI Test Document', level=1)
        doc.add_paragraph('This is a test paragraph for LlamaDocx CLI testing.')
        doc.add_paragraph('This paragraph contains the word "replace_me" that will be replaced.')
        
        # Create a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = 'Header 1'
        table.cell(0, 1).text = 'Header 2'
        table.cell(1, 0).text = 'Data 1'
        table.cell(1, 1).text = 'Data 2'
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


@pytest.fixture
def temp_template():
    """Create a temporary template document for testing CLI commands."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        doc.add_heading('{{title}}', level=1)
        doc.add_paragraph('Name: {{name}}')
        doc.add_paragraph('Date: {{date}}')
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


@pytest.fixture
def temp_data_file():
    """Create a temporary JSON data file for template testing."""
    with tempfile.NamedTemporaryFile(suffix='.json', mode='w', delete=False) as tmp:
        data = {
            'title': 'CLI Template Test',
            'name': 'CLI Tester',
            'date': '2023-06-15'
        }
        json.dump(data, tmp)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


@pytest.fixture
def temp_markdown():
    """Create a temporary Markdown file for testing CLI conversion."""
    with tempfile.NamedTemporaryFile(suffix='.md', mode='w', delete=False) as tmp:
        content = """# CLI Markdown Test

This is a test file for CLI markdown conversion.

## Test Section

- Item 1
- Item 2
- Item 3
"""
        tmp.write(content)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


def run_llamadocx_command(cmd_args):
    """Run a llamadocx CLI command and return the result."""
    full_cmd = ['llamadocx'] + cmd_args
    result = subprocess.run(
        full_cmd,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        check=False
    )
    return result


@pytest.mark.cli
def test_version_command():
    """Test the --version command."""
    result = run_llamadocx_command(['--version'])
    assert result.returncode == 0
    assert 'LlamaDocx' in result.stdout


@pytest.mark.cli
def test_create_command():
    """Test the create command."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
    
    try:
        result = run_llamadocx_command([
            'create',
            output_path,
            '--title', 'CLI Created Document',
            '--author', 'CLI Test'
        ])
        
        assert result.returncode == 0
        assert os.path.exists(output_path)
        
        # Verify document content
        doc = Document(output_path)
        assert doc.paragraphs[0].text == 'CLI Created Document'  # Title should be first heading
        
        # Verify metadata
        core_properties = doc.core_properties
        assert core_properties.title == 'CLI Created Document'
        assert core_properties.author == 'CLI Test'
    finally:
        if os.path.exists(output_path):
            os.unlink(output_path)


@pytest.mark.cli
def test_extract_command(temp_docx):
    """Test the extract command."""
    # Test text extraction
    with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp:
        output_path = tmp.name
    
    try:
        result = run_llamadocx_command([
            'extract',
            temp_docx,
            output_path,
            '--format', 'txt'
        ])
        
        assert result.returncode == 0
        assert os.path.exists(output_path)
        
        # Verify extracted content
        with open(output_path, 'r') as f:
            content = f.read()
            assert 'CLI Test Document' in content
            assert 'test paragraph' in content
    finally:
        if os.path.exists(output_path):
            os.unlink(output_path)


@pytest.mark.cli
def test_search_command(temp_docx):
    """Test the search command."""
    # Test basic search
    result = run_llamadocx_command([
        'search',
        temp_docx,
        'test',
    ])
    
    assert result.returncode == 0
    assert 'test paragraph' in result.stdout


@pytest.mark.cli
def test_replace_command(temp_docx):
    """Test the replace command."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
    
    try:
        result = run_llamadocx_command([
            'replace',
            temp_docx,
            output_path,
            'replace_me',
            'REPLACED'
        ])
        
        assert result.returncode == 0
        assert os.path.exists(output_path)
        
        # Verify replacement
        doc = Document(output_path)
        text = '\n'.join(p.text for p in doc.paragraphs)
        assert 'REPLACED' in text
        assert 'replace_me' not in text
    finally:
        if os.path.exists(output_path):
            os.unlink(output_path)


@pytest.mark.cli
def test_template_command(temp_template, temp_data_file):
    """Test the template command."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
    
    try:
        result = run_llamadocx_command([
            'template',
            temp_template,
            output_path,
            '--data-file', temp_data_file
        ])
        
        assert result.returncode == 0
        assert os.path.exists(output_path)
        
        # Verify template processing
        doc = Document(output_path)
        text = '\n'.join(p.text for p in doc.paragraphs)
        assert 'CLI Template Test' in text
        assert 'Name: CLI Tester' in text
        assert 'Date: 2023-06-15' in text
    finally:
        if os.path.exists(output_path):
            os.unlink(output_path)


@pytest.mark.cli
def test_metadata_command(temp_docx):
    """Test the meta command."""
    # Test getting metadata
    result = run_llamadocx_command([
        'meta',
        temp_docx,
        '--get'
    ])
    
    assert result.returncode == 0
    
    # Test setting metadata
    result = run_llamadocx_command([
        'meta',
        temp_docx,
        '--set',
        'title=New Title',
        'author=New Author'
    ])
    
    assert result.returncode == 0
    
    # Verify metadata was set
    doc = Document(temp_docx)
    assert doc.core_properties.title == 'New Title'
    assert doc.core_properties.author == 'New Author'


@pytest.mark.cli
def test_convert_command(temp_docx, temp_markdown):
    """Test the convert command."""
    # Test DOCX to MD conversion
    with tempfile.NamedTemporaryFile(suffix='.md', delete=False) as md_tmp:
        md_output_path = md_tmp.name
    
    # Test MD to DOCX conversion
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as docx_tmp:
        docx_output_path = docx_tmp.name
    
    try:
        # DOCX to MD
        result = run_llamadocx_command([
            'convert',
            temp_docx,
            md_output_path
        ])
        
        assert result.returncode == 0
        assert os.path.exists(md_output_path)
        
        # Verify conversion
        with open(md_output_path, 'r') as f:
            content = f.read()
            assert '# CLI Test Document' in content
        
        # MD to DOCX
        result = run_llamadocx_command([
            'convert',
            temp_markdown,
            docx_output_path
        ])
        
        assert result.returncode == 0
        assert os.path.exists(docx_output_path)
        
        # Verify conversion
        doc = Document(docx_output_path)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        assert 'CLI Markdown Test' in headings
    finally:
        for path in [md_output_path, docx_output_path]:
            if os.path.exists(path):
                os.unlink(path)


if __name__ == '__main__':
    pytest.main(['-v', __file__]) 