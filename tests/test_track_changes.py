"""
Track Changes tests for LlamaDocx.

This module contains tests for the track changes functionality of the LlamaDocx package.
"""

import os
import tempfile
import pytest
from docx import Document

from llamadocx.track_changes import (
    enable_track_changes,
    disable_track_changes,
    is_track_changes_enabled,
    get_tracked_changes,
    accept_all_changes,
    reject_all_changes,
    accept_change,
    reject_change,
    get_change_author,
    get_change_date,
    get_change_type,
    get_original_text,
    get_changed_text
)


@pytest.fixture
def sample_document():
    """Create a sample document for testing track changes."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        
        # Add heading
        doc.add_heading('Track Changes Test Document', level=1)
        
        # Add paragraphs
        doc.add_paragraph('This is the first paragraph of the test document.')
        doc.add_paragraph('This is the second paragraph that will be edited.')
        doc.add_paragraph('This is the third paragraph that will be deleted.')
        doc.add_paragraph('This is the fourth paragraph that will remain unchanged.')
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


def test_enable_track_changes(sample_document):
    """Test enabling track changes in a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Enable track changes
    result = enable_track_changes(doc, author="Test User")
    
    # Verify track changes was enabled
    assert result is True
    assert is_track_changes_enabled(doc) is True
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify track changes is still enabled
        assert is_track_changes_enabled(doc2) is True
    finally:
        # Clean up
        os.unlink(output_path)


def test_disable_track_changes(sample_document):
    """Test disabling track changes in a document."""
    # Load the document
    doc = Document(sample_document)
    
    # First enable track changes
    enable_result = enable_track_changes(doc, author="Test User")
    assert enable_result is True
    assert is_track_changes_enabled(doc) is True
    
    # Now disable track changes
    disable_result = disable_track_changes(doc)
    
    # Verify track changes was disabled
    assert disable_result is True
    assert is_track_changes_enabled(doc) is False
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify track changes is still disabled
        assert is_track_changes_enabled(doc2) is False
    finally:
        # Clean up
        os.unlink(output_path)


def test_track_text_changes(sample_document):
    """Test tracking text changes in a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Enable track changes
    enable_result = enable_track_changes(doc, author="Test User")
    assert enable_result is True
    
    # Get the second paragraph and modify it
    paragraph = doc.paragraphs[2]
    original_text = paragraph.text
    
    # Replace the paragraph text (this should track the change)
    new_text = "This is the modified second paragraph."
    paragraph.text = new_text
    
    # Get tracked changes
    changes = get_tracked_changes(doc)
    
    # Verify changes were tracked
    assert len(changes) > 0
    
    # Find our specific change
    found_change = False
    for change in changes:
        if get_change_type(change) == "modification" and original_text in get_original_text(change):
            found_change = True
            assert get_changed_text(change) == new_text
            assert get_change_author(change) == "Test User"
            break
    
    assert found_change is True
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get tracked changes
        changes2 = get_tracked_changes(doc2)
        
        # Verify changes persisted
        assert len(changes2) > 0
        
        # Find our specific change again
        found_change2 = False
        for change in changes2:
            if get_change_type(change) == "modification" and original_text in get_original_text(change):
                found_change2 = True
                assert get_changed_text(change) == new_text
                assert get_change_author(change) == "Test User"
                break
        
        assert found_change2 is True
    finally:
        # Clean up
        os.unlink(output_path)


def test_track_deletion(sample_document):
    """Test tracking deletions in a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Enable track changes
    enable_result = enable_track_changes(doc, author="Test User")
    assert enable_result is True
    
    # Get the third paragraph which will be deleted
    paragraph_to_delete = doc.paragraphs[3]
    deleted_text = paragraph_to_delete.text
    
    # Delete the paragraph (this should track the deletion)
    # In python-docx, we have to use a workaround since direct paragraph deletion is not supported
    # We'll just remove all runs from the paragraph, making it empty
    for run in paragraph_to_delete.runs:
        run.text = ""
    
    # Get tracked changes
    changes = get_tracked_changes(doc)
    
    # Verify changes were tracked
    assert len(changes) > 0
    
    # Find our specific deletion
    found_deletion = False
    for change in changes:
        if get_change_type(change) == "deletion" and deleted_text in get_original_text(change):
            found_deletion = True
            assert get_change_author(change) == "Test User"
            break
    
    assert found_deletion is True
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get tracked changes
        changes2 = get_tracked_changes(doc2)
        
        # Verify changes persisted
        assert len(changes2) > 0
        
        # Find our specific deletion again
        found_deletion2 = False
        for change in changes2:
            if get_change_type(change) == "deletion" and deleted_text in get_original_text(change):
                found_deletion2 = True
                assert get_change_author(change) == "Test User"
                break
        
        assert found_deletion2 is True
    finally:
        # Clean up
        os.unlink(output_path)


def test_track_insertion(sample_document):
    """Test tracking insertions in a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Enable track changes
    enable_result = enable_track_changes(doc, author="Test User")
    assert enable_result is True
    
    # Add a new paragraph (this should track the insertion)
    inserted_text = "This is a newly inserted paragraph."
    doc.add_paragraph(inserted_text)
    
    # Get tracked changes
    changes = get_tracked_changes(doc)
    
    # Verify changes were tracked
    assert len(changes) > 0
    
    # Find our specific insertion
    found_insertion = False
    for change in changes:
        if get_change_type(change) == "insertion" and inserted_text in get_changed_text(change):
            found_insertion = True
            assert get_change_author(change) == "Test User"
            break
    
    assert found_insertion is True
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get tracked changes
        changes2 = get_tracked_changes(doc2)
        
        # Verify changes persisted
        assert len(changes2) > 0
        
        # Find our specific insertion again
        found_insertion2 = False
        for change in changes2:
            if get_change_type(change) == "insertion" and inserted_text in get_changed_text(change):
                found_insertion2 = True
                assert get_change_author(change) == "Test User"
                break
        
        assert found_insertion2 is True
    finally:
        # Clean up
        os.unlink(output_path)


def test_accept_all_changes(sample_document):
    """Test accepting all tracked changes in a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Enable track changes
    enable_result = enable_track_changes(doc, author="Test User")
    assert enable_result is True
    
    # Make multiple changes
    # 1. Modify paragraph 2
    doc.paragraphs[2].text = "This is the modified second paragraph."
    
    # 2. Clear paragraph 3 (simulating deletion)
    for run in doc.paragraphs[3].runs:
        run.text = ""
    
    # 3. Add a new paragraph
    doc.add_paragraph("This is a newly inserted paragraph.")
    
    # Verify changes were tracked
    changes = get_tracked_changes(doc)
    assert len(changes) > 0
    
    # Accept all changes
    accept_result = accept_all_changes(doc)
    
    # Verify all changes were accepted
    assert accept_result is True
    remaining_changes = get_tracked_changes(doc)
    assert len(remaining_changes) == 0
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify no tracked changes remain
        remaining_changes2 = get_tracked_changes(doc2)
        assert len(remaining_changes2) == 0
    finally:
        # Clean up
        os.unlink(output_path)


def test_reject_all_changes(sample_document):
    """Test rejecting all tracked changes in a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Get the original text of paragraphs for comparison later
    original_texts = [para.text for para in doc.paragraphs]
    original_paragraph_count = len(doc.paragraphs)
    
    # Enable track changes
    enable_result = enable_track_changes(doc, author="Test User")
    assert enable_result is True
    
    # Make multiple changes
    # 1. Modify paragraph 2
    doc.paragraphs[2].text = "This is the modified second paragraph."
    
    # 2. Clear paragraph 3 (simulating deletion)
    for run in doc.paragraphs[3].runs:
        run.text = ""
    
    # 3. Add a new paragraph
    doc.add_paragraph("This is a newly inserted paragraph.")
    
    # Verify changes were tracked
    changes = get_tracked_changes(doc)
    assert len(changes) > 0
    
    # Reject all changes
    reject_result = reject_all_changes(doc)
    
    # Verify all changes were rejected
    assert reject_result is True
    remaining_changes = get_tracked_changes(doc)
    assert len(remaining_changes) == 0
    
    # Verify document has been restored to original state
    # Note: This might be approximate since the actual implementation
    # of rejection can be complex and not all changes might revert exactly
    # as expected in python-docx
    assert len(doc.paragraphs) == original_paragraph_count
    for i, para in enumerate(doc.paragraphs):
        if i < len(original_texts):
            # The text might not be exactly the same after rejection due to
            # how Word handles track changes internally
            assert para.text == original_texts[i] or para.text.strip() == original_texts[i].strip()
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify no tracked changes remain
        remaining_changes2 = get_tracked_changes(doc2)
        assert len(remaining_changes2) == 0
    finally:
        # Clean up
        os.unlink(output_path)


def test_accept_specific_change(sample_document):
    """Test accepting a specific tracked change in a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Enable track changes
    enable_result = enable_track_changes(doc, author="Test User")
    assert enable_result is True
    
    # Make multiple changes
    # 1. Modify paragraph 2
    doc.paragraphs[2].text = "This is the modified second paragraph."
    
    # 2. Add a new paragraph
    doc.add_paragraph("This is a newly inserted paragraph.")
    
    # Get tracked changes
    changes = get_tracked_changes(doc)
    assert len(changes) >= 2
    
    # Find the insertion change
    insertion_change = None
    for change in changes:
        if get_change_type(change) == "insertion" and "newly inserted" in get_changed_text(change):
            insertion_change = change
            break
    
    assert insertion_change is not None
    
    # Accept just the insertion change
    accept_result = accept_change(doc, insertion_change)
    
    # Verify the change was accepted
    assert accept_result is True
    
    # Get remaining changes
    remaining_changes = get_tracked_changes(doc)
    
    # Verify only one change remains and it's not the insertion
    assert len(remaining_changes) == len(changes) - 1
    for change in remaining_changes:
        assert get_change_type(change) != "insertion" or "newly inserted" not in get_changed_text(change)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify the correct changes remain
        remaining_changes2 = get_tracked_changes(doc2)
        assert len(remaining_changes2) == len(changes) - 1
        for change in remaining_changes2:
            assert get_change_type(change) != "insertion" or "newly inserted" not in get_changed_text(change)
    finally:
        # Clean up
        os.unlink(output_path)


def test_reject_specific_change(sample_document):
    """Test rejecting a specific tracked change in a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Get the original text of paragraph 2
    original_para2_text = doc.paragraphs[2].text
    
    # Enable track changes
    enable_result = enable_track_changes(doc, author="Test User")
    assert enable_result is True
    
    # Make multiple changes
    # 1. Modify paragraph 2
    doc.paragraphs[2].text = "This is the modified second paragraph."
    
    # 2. Add a new paragraph
    doc.add_paragraph("This is a newly inserted paragraph.")
    
    # Get tracked changes
    changes = get_tracked_changes(doc)
    assert len(changes) >= 2
    
    # Find the modification change for paragraph 2
    modification_change = None
    for change in changes:
        if get_change_type(change) == "modification" and "second paragraph" in get_original_text(change):
            modification_change = change
            break
    
    assert modification_change is not None
    
    # Reject just the modification change
    reject_result = reject_change(doc, modification_change)
    
    # Verify the change was rejected
    assert reject_result is True
    
    # Verify paragraph 2 has been restored to original text
    assert doc.paragraphs[2].text == original_para2_text or doc.paragraphs[2].text.strip() == original_para2_text.strip()
    
    # Get remaining changes
    remaining_changes = get_tracked_changes(doc)
    
    # Verify only one change remains and it's not the modification
    assert len(remaining_changes) == len(changes) - 1
    for change in remaining_changes:
        assert get_change_type(change) != "modification" or "second paragraph" not in get_original_text(change)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify the correct changes remain
        remaining_changes2 = get_tracked_changes(doc2)
        assert len(remaining_changes2) == len(changes) - 1
        for change in remaining_changes2:
            assert get_change_type(change) != "modification" or "second paragraph" not in get_original_text(change)
    finally:
        # Clean up
        os.unlink(output_path)


def test_change_metadata(sample_document):
    """Test retrieving metadata for tracked changes."""
    # Load the document
    doc = Document(sample_document)
    
    # Enable track changes with a specific author
    author = "Test Author"
    enable_result = enable_track_changes(doc, author=author)
    assert enable_result is True
    
    # Make a change
    doc.paragraphs[2].text = "This is the modified second paragraph."
    
    # Get tracked changes
    changes = get_tracked_changes(doc)
    assert len(changes) > 0
    
    # Check metadata for the change
    change = changes[0]
    
    # Verify author
    assert get_change_author(change) == author
    
    # Verify change date exists
    change_date = get_change_date(change)
    assert change_date is not None
    
    # Verify change type
    change_type = get_change_type(change)
    assert change_type in ["insertion", "deletion", "modification"]
    
    # Verify original and changed text
    if change_type == "modification":
        assert get_original_text(change) is not None
        assert get_changed_text(change) is not None
    elif change_type == "insertion":
        assert get_changed_text(change) is not None
    elif change_type == "deletion":
        assert get_original_text(change) is not None


def test_multiple_authors(sample_document):
    """Test track changes with multiple authors."""
    # Load the document
    doc = Document(sample_document)
    
    # Enable track changes with first author
    author1 = "Author One"
    enable_result = enable_track_changes(doc, author=author1)
    assert enable_result is True
    
    # Make a change with first author
    doc.paragraphs[2].text = "This is edited by Author One."
    
    # Change to second author
    author2 = "Author Two"
    enable_result = enable_track_changes(doc, author=author2)
    assert enable_result is True
    
    # Make a change with second author
    doc.add_paragraph("This paragraph was added by Author Two.")
    
    # Get tracked changes
    changes = get_tracked_changes(doc)
    assert len(changes) >= 2
    
    # Verify both authors are represented in the changes
    authors = [get_change_author(change) for change in changes]
    assert author1 in authors
    assert author2 in authors
    
    # Verify specific changes by each author
    for change in changes:
        author = get_change_author(change)
        if author == author1:
            if get_change_type(change) == "modification":
                assert "Author One" in get_changed_text(change)
        elif author == author2:
            if get_change_type(change) == "insertion":
                assert "Author Two" in get_changed_text(change)


def test_complex_document_with_track_changes(sample_document):
    """Test track changes in a more complex document with various change types."""
    # Load the document
    doc = Document(sample_document)
    
    # Enable track changes
    author = "Complex Test"
    enable_result = enable_track_changes(doc, author=author)
    assert enable_result is True
    
    # Make various changes
    
    # 1. Text modification
    doc.paragraphs[2].text = "This paragraph has been completely replaced."
    
    # 2. Partial text change within a run (if supported by the API)
    # This will depend on how your API handles runs and character-level changes
    # For this test, we'll simulate by replacing the paragraph
    doc.paragraphs[1].text = "This is the first paragraph with some modified words."
    
    # 3. Insertion
    doc.add_paragraph("This is a brand new paragraph inserted into the document.")
    
    # 4. Deletion (simulated)
    for run in doc.paragraphs[3].runs:
        run.text = ""
    
    # 5. Formatting change (if tracked by the API)
    # This might not be tracked in all implementations
    if len(doc.paragraphs[4].runs) > 0:
        doc.paragraphs[4].runs[0].bold = True
    
    # Get all tracked changes
    changes = get_tracked_changes(doc)
    
    # Verify multiple changes were tracked
    # The exact number depends on how the API handles changes
    assert len(changes) > 0
    
    # Verify different change types were captured
    change_types = set(get_change_type(change) for change in changes)
    # We should have at least insertions and modifications
    assert "insertion" in change_types or "modification" in change_types
    
    # Verify all changes have the correct author
    for change in changes:
        assert get_change_author(change) == author
    
    # Accept some changes and reject others
    if len(changes) >= 2:
        # Accept the first change
        accept_result = accept_change(doc, changes[0])
        assert accept_result is True
        
        # Reject the second change
        reject_result = reject_change(doc, changes[1])
        assert reject_result is True
        
        # Verify remaining changes
        remaining_changes = get_tracked_changes(doc)
        assert len(remaining_changes) == len(changes) - 2
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify tracked changes state persisted
        remaining_changes2 = get_tracked_changes(doc2)
        if len(changes) >= 2:
            assert len(remaining_changes2) == len(changes) - 2
    finally:
        # Clean up
        os.unlink(output_path)


if __name__ == '__main__':
    pytest.main(['-v', __file__]) 