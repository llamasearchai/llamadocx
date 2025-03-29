"""
Watermarks tests for LlamaDocx.

This module contains tests for the watermarks functionality of the LlamaDocx package.
"""

import os
import tempfile
from pathlib import Path
import pytest
from docx import Document

from llamadocx.watermarks import (
    add_text_watermark,
    add_image_watermark,
    remove_watermark,
    get_watermark_text,
    has_watermark,
    update_text_watermark
)


@pytest.fixture
def sample_document():
    """Create a sample document for testing watermarks."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        
        # Add heading
        doc.add_heading('Watermarks Test Document', level=1)
        
        # Add paragraphs
        doc.add_paragraph('This is the first paragraph of the test document.')
        doc.add_paragraph('This is the second paragraph of the test document.')
        doc.add_paragraph('This is the third paragraph of the test document.')
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


@pytest.fixture
def sample_image():
    """Create a sample image for testing image watermarks."""
    # This is a placeholder. In a real test, you would use a real image file
    # For this test, we'll use a path to a sample image that should exist in the test directory
    image_path = Path(__file__).parent / "resources" / "sample_image.png"
    if not image_path.exists():
        # Create a directory for resources if it doesn't exist
        (Path(__file__).parent / "resources").mkdir(exist_ok=True)
        
        # If the image doesn't exist, we'll create a dummy text file instead
        # In a real scenario, you would have an actual image file
        with open(image_path, 'w') as f:
            f.write("This is a dummy file to simulate an image for testing purposes.")
    
    yield str(image_path)
    
    # We won't delete the image as it might be used by other tests


def test_add_text_watermark(sample_document):
    """Test adding a text watermark to a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a text watermark
    watermark_text = "CONFIDENTIAL"
    result = add_text_watermark(doc, watermark_text)
    
    # Verify watermark was added successfully
    assert result is True
    assert has_watermark(doc) is True
    
    # Get the watermark text and verify it matches what we set
    assert get_watermark_text(doc) == watermark_text
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify watermark persisted
        assert has_watermark(doc2) is True
        assert get_watermark_text(doc2) == watermark_text
    finally:
        # Clean up
        os.unlink(output_path)


def test_add_text_watermark_with_options(sample_document):
    """Test adding a text watermark with various formatting options."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a text watermark with custom options
    watermark_text = "DRAFT"
    result = add_text_watermark(
        doc, 
        watermark_text, 
        font_name="Arial", 
        font_size=48, 
        color="#FF0000",  # Red
        opacity=0.3,
        rotation=45
    )
    
    # Verify watermark was added successfully
    assert result is True
    assert has_watermark(doc) is True
    
    # Get the watermark text and verify it matches what we set
    assert get_watermark_text(doc) == watermark_text
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify watermark persisted
        assert has_watermark(doc2) is True
        assert get_watermark_text(doc2) == watermark_text
    finally:
        # Clean up
        os.unlink(output_path)


def test_add_image_watermark(sample_document, sample_image):
    """Test adding an image watermark to a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add an image watermark
    result = add_image_watermark(doc, sample_image)
    
    # Verify watermark was added successfully
    assert result is True
    assert has_watermark(doc) is True
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify watermark persisted
        assert has_watermark(doc2) is True
    finally:
        # Clean up
        os.unlink(output_path)


def test_add_image_watermark_with_options(sample_document, sample_image):
    """Test adding an image watermark with custom options."""
    # Load the document
    doc = Document(sample_document)
    
    # Add an image watermark with custom options
    result = add_image_watermark(
        doc, 
        sample_image, 
        opacity=0.2,
        scale=0.5,
        rotation=30
    )
    
    # Verify watermark was added successfully
    assert result is True
    assert has_watermark(doc) is True
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify watermark persisted
        assert has_watermark(doc2) is True
    finally:
        # Clean up
        os.unlink(output_path)


def test_update_text_watermark(sample_document):
    """Test updating a text watermark in a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add initial watermark
    initial_text = "DRAFT"
    result = add_text_watermark(doc, initial_text)
    
    # Verify initial watermark was added
    assert result is True
    assert has_watermark(doc) is True
    assert get_watermark_text(doc) == initial_text
    
    # Update the watermark
    updated_text = "FINAL"
    update_result = update_text_watermark(doc, updated_text)
    
    # Verify watermark was updated
    assert update_result is True
    assert has_watermark(doc) is True
    assert get_watermark_text(doc) == updated_text
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify updated watermark persisted
        assert has_watermark(doc2) is True
        assert get_watermark_text(doc2) == updated_text
    finally:
        # Clean up
        os.unlink(output_path)


def test_update_text_watermark_with_options(sample_document):
    """Test updating a text watermark with different formatting options."""
    # Load the document
    doc = Document(sample_document)
    
    # Add initial watermark
    initial_text = "DRAFT"
    result = add_text_watermark(doc, initial_text)
    
    # Verify initial watermark was added
    assert result is True
    assert get_watermark_text(doc) == initial_text
    
    # Update the watermark with different formatting options
    updated_text = "CONFIDENTIAL"
    update_result = update_text_watermark(
        doc, 
        updated_text,
        font_name="Times New Roman",
        font_size=72,
        color="#0000FF",  # Blue
        opacity=0.5,
        rotation=0
    )
    
    # Verify watermark was updated
    assert update_result is True
    assert get_watermark_text(doc) == updated_text
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify updated watermark persisted
        assert has_watermark(doc2) is True
        assert get_watermark_text(doc2) == updated_text
    finally:
        # Clean up
        os.unlink(output_path)


def test_remove_watermark(sample_document):
    """Test removing a watermark from a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a watermark
    watermark_text = "DRAFT"
    add_result = add_text_watermark(doc, watermark_text)
    
    # Verify watermark was added
    assert add_result is True
    assert has_watermark(doc) is True
    
    # Remove the watermark
    remove_result = remove_watermark(doc)
    
    # Verify watermark was removed
    assert remove_result is True
    assert has_watermark(doc) is False
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify watermark removal persisted
        assert has_watermark(doc2) is False
    finally:
        # Clean up
        os.unlink(output_path)


def test_remove_image_watermark(sample_document, sample_image):
    """Test removing an image watermark from a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add an image watermark
    add_result = add_image_watermark(doc, sample_image)
    
    # Verify watermark was added
    assert add_result is True
    assert has_watermark(doc) is True
    
    # Remove the watermark
    remove_result = remove_watermark(doc)
    
    # Verify watermark was removed
    assert remove_result is True
    assert has_watermark(doc) is False
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify watermark removal persisted
        assert has_watermark(doc2) is False
    finally:
        # Clean up
        os.unlink(output_path)


def test_get_watermark_text_no_watermark(sample_document):
    """Test getting watermark text from a document with no watermark."""
    # Load the document
    doc = Document(sample_document)
    
    # Verify no watermark exists
    assert has_watermark(doc) is False
    
    # Get watermark text should return None or empty string
    assert get_watermark_text(doc) is None or get_watermark_text(doc) == ""


def test_multiple_watermark_operations(sample_document, sample_image):
    """Test a series of watermark operations on a document."""
    # Load the document
    doc = Document(sample_document)
    
    # 1. Add a text watermark
    text_watermark = "DRAFT"
    add_text_result = add_text_watermark(doc, text_watermark)
    assert add_text_result is True
    assert has_watermark(doc) is True
    assert get_watermark_text(doc) == text_watermark
    
    # 2. Remove the watermark
    remove_result = remove_watermark(doc)
    assert remove_result is True
    assert has_watermark(doc) is False
    
    # 3. Add an image watermark
    add_image_result = add_image_watermark(doc, sample_image)
    assert add_image_result is True
    assert has_watermark(doc) is True
    
    # 4. Remove again
    remove_result = remove_watermark(doc)
    assert remove_result is True
    assert has_watermark(doc) is False
    
    # 5. Add another text watermark
    final_text = "FINAL"
    add_final_result = add_text_watermark(doc, final_text)
    assert add_final_result is True
    assert has_watermark(doc) is True
    assert get_watermark_text(doc) == final_text
    
    # Save and reload to verify last operation persisted
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify final watermark persisted
        assert has_watermark(doc2) is True
        assert get_watermark_text(doc2) == final_text
    finally:
        # Clean up
        os.unlink(output_path)


if __name__ == '__main__':
    pytest.main(['-v', __file__]) 