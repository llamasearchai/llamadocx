"""
Conversion tests for LlamaDocx.

This module contains tests for the format conversion functionality of the LlamaDocx package.
"""

import os
import tempfile
from pathlib import Path
import pytest
import shutil

from llamadocx.convert import (
    convert_docx_to_markdown,
    convert_markdown_to_docx,
    convert_docx_to_pdf,
    convert_docx_to_html,
    is_pandoc_available
)


# Skip all tests if pandoc is not available
pytestmark = pytest.mark.skipif(
    not is_pandoc_available(),
    reason="Pandoc is not available for conversion tests"
)


@pytest.fixture
def sample_docx():
    """Create a simple DOCX document for testing conversions."""
    from docx import Document
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        
        # Add content
        doc.add_heading('Test Document', level=1)
        doc.add_paragraph('This is a paragraph with *italic* and **bold** text.')
        
        # Add a bulleted list
        paragraph = doc.add_paragraph('Here is a list:')
        doc.add_paragraph('• Item 1', style='List Bullet')
        doc.add_paragraph('• Item 2', style='List Bullet')
        doc.add_paragraph('• Item 3', style='List Bullet')
        
        # Add a simple table
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        table.cell(0, 0).text = 'Header 1'
        table.cell(0, 1).text = 'Header 2'
        table.cell(1, 0).text = 'Cell 1'
        table.cell(1, 1).text = 'Cell 2'
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


@pytest.fixture
def sample_markdown():
    """Create a sample Markdown file for testing conversions."""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as tmp:
        # Write Markdown content
        content = """# Test Markdown Document

This is a paragraph with *italic* and **bold** text.

## Section One

Here is a list:

- Item 1
- Item 2
- Item 3

## Section Two

| Header 1 | Header 2 |
|----------|----------|
| Cell 1   | Cell 2   |

"""
        tmp.write(content)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


def test_docx_to_markdown_conversion(sample_docx):
    """Test converting a DOCX file to Markdown."""
    # Create output path
    with tempfile.NamedTemporaryFile(suffix='.md', delete=False) as tmp:
        output_path = tmp.name
    
    try:
        # Perform conversion
        result = convert_docx_to_markdown(sample_docx, output_path)
        
        # Check conversion result
        assert result is True
        assert os.path.exists(output_path)
        
        # Verify content
        with open(output_path, 'r') as f:
            content = f.read()
            
        # Check for expected Markdown elements
        assert '# Test Document' in content
        assert '*italic*' in content or '_italic_' in content  # Different Markdown styles possible
        assert '**bold**' in content or '__bold__' in content
        
        # Check for list elements
        assert '- Item 1' in content or '* Item 1' in content
        assert '- Item 2' in content or '* Item 2' in content
        assert '- Item 3' in content or '* Item 3' in content
        
        # Check for table elements (may vary by pandoc version)
        assert 'Header 1' in content
        assert 'Header 2' in content
        assert 'Cell 1' in content
        assert 'Cell 2' in content
        
    finally:
        # Clean up
        if os.path.exists(output_path):
            os.unlink(output_path)


def test_markdown_to_docx_conversion(sample_markdown):
    """Test converting a Markdown file to DOCX."""
    # Create output path
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
    
    try:
        # Perform conversion
        result = convert_markdown_to_docx(sample_markdown, output_path)
        
        # Check conversion result
        assert result is True
        assert os.path.exists(output_path)
        
        # Verify content (limited check since we can't easily read DOCX content in tests)
        from docx import Document
        doc = Document(output_path)
        
        # Extract all text
        text = ' '.join([p.text for p in doc.paragraphs])
        
        # Check for expected content
        assert 'Test Markdown Document' in text
        assert 'Section One' in text
        assert 'Section Two' in text
        assert 'Item 1' in text
        assert 'Item 2' in text
        assert 'Item 3' in text
        
    finally:
        # Clean up
        if os.path.exists(output_path):
            os.unlink(output_path)


@pytest.mark.skipif(
    not shutil.which('wkhtmltopdf') and not shutil.which('libreoffice'),
    reason="PDF conversion tools not available"
)
def test_docx_to_pdf_conversion(sample_docx):
    """Test converting a DOCX file to PDF."""
    # Create output path
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
        output_path = tmp.name
    
    try:
        # Perform conversion
        result = convert_docx_to_pdf(sample_docx, output_path)
        
        # Check conversion result
        assert result is True
        assert os.path.exists(output_path)
        
        # Check that file is not empty
        assert os.path.getsize(output_path) > 0
        
    finally:
        # Clean up
        if os.path.exists(output_path):
            os.unlink(output_path)


def test_docx_to_html_conversion(sample_docx):
    """Test converting a DOCX file to HTML."""
    # Create output path
    with tempfile.NamedTemporaryFile(suffix='.html', delete=False) as tmp:
        output_path = tmp.name
    
    try:
        # Perform conversion
        result = convert_docx_to_html(sample_docx, output_path)
        
        # Check conversion result
        assert result is True
        assert os.path.exists(output_path)
        
        # Verify content
        with open(output_path, 'r') as f:
            content = f.read()
            
        # Check for expected HTML elements
        assert '<html' in content.lower()
        assert '<body' in content.lower()
        assert 'Test Document' in content
        assert '<table' in content.lower()
        assert '<li' in content.lower()  # List items
        
    finally:
        # Clean up
        if os.path.exists(output_path):
            os.unlink(output_path)


def test_conversion_error_handling():
    """Test error handling during conversion."""
    # Test with non-existent input file
    with tempfile.NamedTemporaryFile(suffix='.md') as tmp:
        result = convert_docx_to_markdown('nonexistent.docx', tmp.name)
        assert result is False
    
    # Test with invalid output directory
    with tempfile.NamedTemporaryFile(suffix='.docx') as tmp:
        result = convert_markdown_to_docx(tmp.name, '/nonexistent/directory/output.docx')
        assert result is False


def test_bulk_conversion():
    """Test bulk conversion of multiple files."""
    # Create multiple test files
    temp_dir = tempfile.mkdtemp()
    try:
        # Create markdown files
        md_files = []
        for i in range(3):
            md_path = os.path.join(temp_dir, f'test{i}.md')
            with open(md_path, 'w') as f:
                f.write(f"# Test Document {i}\n\nThis is test document {i}.")
            md_files.append(md_path)
        
        # Create output directory
        output_dir = os.path.join(temp_dir, 'output')
        os.makedirs(output_dir)
        
        # Convert each file
        for md_file in md_files:
            base_name = os.path.basename(md_file)
            name, _ = os.path.splitext(base_name)
            output_path = os.path.join(output_dir, f"{name}.docx")
            
            result = convert_markdown_to_docx(md_file, output_path)
            assert result is True
            assert os.path.exists(output_path)
            
        # Check all files were created
        assert len(os.listdir(output_dir)) == 3
        
    finally:
        # Clean up
        shutil.rmtree(temp_dir)


def test_conversion_with_options(sample_docx, sample_markdown):
    """Test conversion with additional options."""
    # Create output paths
    md_output = tempfile.NamedTemporaryFile(suffix='.md', delete=False).name
    docx_output = tempfile.NamedTemporaryFile(suffix='.docx', delete=False).name
    
    try:
        # Convert DOCX to Markdown with options
        result1 = convert_docx_to_markdown(
            sample_docx, 
            md_output,
            extra_args=['--wrap=none', '--standalone']
        )
        
        assert result1 is True
        assert os.path.exists(md_output)
        
        # Convert Markdown to DOCX with options
        result2 = convert_markdown_to_docx(
            sample_markdown,
            docx_output,
            extra_args=['--toc', '--toc-depth=2']
        )
        
        assert result2 is True
        assert os.path.exists(docx_output)
        
    finally:
        # Clean up
        for path in [md_output, docx_output]:
            if os.path.exists(path):
                os.unlink(path)


if __name__ == '__main__':
    pytest.main(['-v', __file__]) 