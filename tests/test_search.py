"""
Search and Replace tests for LlamaDocx.

This module contains tests for the search and replacement functionality of the LlamaDocx package.
"""

import os
import tempfile
import re
import pytest
from docx import Document

from llamadocx.search import search_text, replace_text, find_all_text


@pytest.fixture
def sample_document():
    """Create a sample document with searchable text."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        
        # Add heading
        doc.add_heading('Sample Document for Search Testing', level=1)
        
        # Add paragraphs with searchable content
        doc.add_paragraph('This is a sample paragraph with some text to search for.')
        doc.add_paragraph('We need multiple instances of certain words to test the search functionality.')
        doc.add_paragraph('For example, the word "example" appears multiple times in this example document.')
        doc.add_paragraph('Case sensitivity testing: Example vs example vs EXAMPLE.')
        
        # Add paragraphs with special characters
        doc.add_paragraph('Special characters: $100.00, 50%, @username, #hashtag.')
        doc.add_paragraph('Punctuation: Hello, world! How are you? This is a test; it has various punctuation marks.')
        
        # Add a table with searchable content
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        table.cell(0, 0).text = 'Item'
        table.cell(0, 1).text = 'Description'
        
        table.cell(1, 0).text = 'Item 1'
        table.cell(1, 1).text = 'This is the first item in our example.'
        
        table.cell(2, 0).text = 'Item 2'
        table.cell(2, 1).text = 'This is the second item in our example table.'
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


def test_search_text_simple(sample_document):
    """Test basic text search functionality."""
    doc = Document(sample_document)
    
    # Search for a word that appears multiple times
    results = search_text(doc, 'example')
    
    # Check that we found the expected number of occurrences
    assert len(results) >= 4
    
    # Check that the results include paragraph and run indices
    for result in results:
        assert 'paragraph_index' in result
        assert 'run_index' in result
        assert 'text' in result
        assert 'example' in result['text'].lower()


def test_search_text_case_sensitive(sample_document):
    """Test case-sensitive text search."""
    doc = Document(sample_document)
    
    # Search with case sensitivity
    lower_results = search_text(doc, 'example', case_sensitive=True)
    upper_results = search_text(doc, 'Example', case_sensitive=True)
    
    # Check that we get different results for different cases
    assert len(lower_results) != len(upper_results)
    
    # Search without case sensitivity
    insensitive_results = search_text(doc, 'example', case_sensitive=False)
    
    # Check that insensitive search finds more instances
    assert len(insensitive_results) >= len(lower_results) + len(upper_results)


def test_search_text_with_regex(sample_document):
    """Test regex-based text search."""
    doc = Document(sample_document)
    
    # Search using regex pattern
    results = search_text(doc, r'\b\w+@\w+\b', use_regex=True)
    
    # Check that we find the email-like pattern
    assert len(results) == 1
    assert '@username' in results[0]['text']
    
    # Search for all words that start with 'ex'
    ex_results = search_text(doc, r'\bex\w*\b', use_regex=True, case_sensitive=False)
    
    # There should be multiple matches
    assert len(ex_results) >= 4


def test_search_text_in_tables(sample_document):
    """Test searching text within tables."""
    doc = Document(sample_document)
    
    # Search for text that appears in table cells
    results = search_text(doc, 'item', include_tables=True, case_sensitive=False)
    
    # Check that we found instances in tables
    assert len(results) >= 3
    
    # When tables are excluded, we should find fewer instances
    no_table_results = search_text(doc, 'item', include_tables=False, case_sensitive=False)
    assert len(no_table_results) < len(results)


def test_find_all_text(sample_document):
    """Test the find_all_text function for counting occurrences."""
    doc = Document(sample_document)
    
    # Count occurrences of a word
    count = find_all_text(doc, 'example', case_sensitive=False)
    
    # Check that we found the expected number
    assert count >= 4
    
    # Test with regex pattern
    regex_count = find_all_text(doc, r'\btest\w*\b', use_regex=True, case_sensitive=False)
    assert regex_count >= 2


def test_replace_text_simple(sample_document):
    """Test basic text replacement."""
    # Load the document
    doc = Document(sample_document)
    
    # Replace a specific term
    replacements = replace_text(doc, 'example', 'REPLACED')
    
    # Check that replacements were made
    assert replacements > 0
    
    # Save the modified document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the modified document and check replacements
        modified_doc = Document(output_path)
        
        # Original term should no longer be present
        count = find_all_text(modified_doc, 'example', case_sensitive=True)
        assert count == 0
        
        # Replacement term should be present
        replaced_count = find_all_text(modified_doc, 'REPLACED', case_sensitive=True)
        assert replaced_count == replacements
    finally:
        # Clean up
        os.unlink(output_path)


def test_replace_text_case_sensitive(sample_document):
    """Test case-sensitive text replacement."""
    # Load the document
    doc = Document(sample_document)
    
    # Replace only capitalized instances
    replacements = replace_text(doc, 'Example', 'REPLACED', case_sensitive=True)
    
    # Check that replacements were made
    assert replacements > 0
    
    # Save the modified document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the modified document and check replacements
        modified_doc = Document(output_path)
        
        # Capitalized term should no longer be present
        count = find_all_text(modified_doc, 'Example', case_sensitive=True)
        assert count == 0
        
        # Lowercase instances should still be present
        lower_count = find_all_text(modified_doc, 'example', case_sensitive=True)
        assert lower_count > 0
    finally:
        # Clean up
        os.unlink(output_path)


def test_replace_text_with_regex(sample_document):
    """Test regex-based text replacement."""
    # Load the document
    doc = Document(sample_document)
    
    # Replace all currency amounts using regex
    pattern = r'\$\d+\.\d+'
    replacements = replace_text(doc, pattern, '[AMOUNT]', use_regex=True)
    
    # Check that replacements were made
    assert replacements > 0
    
    # Save the modified document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the modified document and check replacements
        modified_doc = Document(output_path)
        
        # Currency pattern should no longer be present
        count = find_all_text(modified_doc, r'\$\d+\.\d+', use_regex=True)
        assert count == 0
        
        # Replacement term should be present
        replaced_count = find_all_text(modified_doc, '[AMOUNT]')
        assert replaced_count == replacements
    finally:
        # Clean up
        os.unlink(output_path)


def test_replace_text_with_function(sample_document):
    """Test replacing text using a function."""
    # Load the document
    doc = Document(sample_document)
    
    # Define a replacement function to uppercase matches
    def uppercase_match(match):
        return match.group(0).upper()
    
    # Replace "item" with its uppercase version
    replacements = replace_text(doc, r'item \d', uppercase_match, use_regex=True)
    
    # Check that replacements were made
    assert replacements > 0
    
    # Save the modified document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the modified document and check replacements
        modified_doc = Document(output_path)
        
        # Check that uppercase versions are present
        content = " ".join([p.text for p in modified_doc.paragraphs] + 
                         [cell.text for table in modified_doc.tables for row in table.rows for cell in row.cells])
        
        assert "ITEM 1" in content
        assert "ITEM 2" in content
    finally:
        # Clean up
        os.unlink(output_path)


def test_replace_in_tables(sample_document):
    """Test replacing text in table cells."""
    # Load the document
    doc = Document(sample_document)
    
    # Replace text only in tables
    replacements = replace_text(doc, "item", "product", include_tables=True, case_sensitive=False)
    
    # Check that replacements were made
    assert replacements > 0
    
    # Save the modified document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the modified document and check replacements in tables
        modified_doc = Document(output_path)
        
        # Check table content
        table_text = " ".join([cell.text for table in modified_doc.tables for row in table.rows for cell in row.cells])
        assert "Product 1" in table_text
        assert "Product 2" in table_text
    finally:
        # Clean up
        os.unlink(output_path)


def test_multiple_replacements(sample_document):
    """Test multiple sequential replacements on the same document."""
    # Load the document
    doc = Document(sample_document)
    
    # Perform multiple replacements
    replace1 = replace_text(doc, "example", "instance")
    replace2 = replace_text(doc, "instance", "case")
    replace3 = replace_text(doc, "sample", "test")
    
    # Check that replacements were made
    assert replace1 > 0
    assert replace2 > 0
    assert replace3 > 0
    
    # Save the modified document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the modified document
        modified_doc = Document(output_path)
        
        # "example" should be replaced with "case" after two replacements
        example_count = find_all_text(modified_doc, "example", case_sensitive=False)
        assert example_count == 0
        
        instance_count = find_all_text(modified_doc, "instance", case_sensitive=False)
        assert instance_count == 0
        
        case_count = find_all_text(modified_doc, "case", case_sensitive=False)
        assert case_count >= replace1
        
        # "sample" should be replaced with "test"
        sample_count = find_all_text(modified_doc, "sample", case_sensitive=False)
        assert sample_count == 0
        
        test_count = find_all_text(modified_doc, "test", case_sensitive=False)
        assert test_count >= replace3
    finally:
        # Clean up
        os.unlink(output_path)


if __name__ == '__main__':
    pytest.main(['-v', __file__]) 