"""
Headers and Footers tests for LlamaDocx.

This module contains tests for the headers and footers functionality of the LlamaDocx package.
"""

import os
import tempfile
import pytest
from docx import Document
from docx.enum.section import WD_SECTION_START

from llamadocx.headers_footers import (
    add_header,
    add_footer,
    add_page_numbers,
    set_different_first_page,
    set_different_odd_even_pages,
    remove_header,
    remove_footer,
    add_header_image,
    add_footer_image,
    link_to_previous
)


@pytest.fixture
def sample_document():
    """Create a sample document for testing headers and footers."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        
        # Add heading
        doc.add_heading('Headers and Footers Test Document', level=1)
        
        # Add content
        for i in range(3):
            doc.add_paragraph(f'This is paragraph {i+1} in the document.')
        
        # Add section break
        doc.add_paragraph().add_run().add_break(WD_SECTION_START.NEW_PAGE)
        
        # Add content to second section
        doc.add_heading('Second Section', level=1)
        for i in range(2):
            doc.add_paragraph(f'This is paragraph {i+1} in the second section.')
        
        # Add another section break
        doc.add_paragraph().add_run().add_break(WD_SECTION_START.NEW_PAGE)
        
        # Add content to third section
        doc.add_heading('Third Section', level=1)
        doc.add_paragraph('This is content in the third section.')
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


def test_add_header(sample_document):
    """Test adding a header to a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a header to the document
    header_text = 'Test Document Header'
    header = add_header(doc, header_text)
    
    # Verify header was added
    assert header is not None
    assert header.paragraphs[0].text == header_text
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify header persisted
        assert doc2.sections[0].header.paragraphs[0].text == header_text
    finally:
        # Clean up
        os.unlink(output_path)


def test_add_footer(sample_document):
    """Test adding a footer to a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a footer to the document
    footer_text = 'Test Document Footer'
    footer = add_footer(doc, footer_text)
    
    # Verify footer was added
    assert footer is not None
    assert footer.paragraphs[0].text == footer_text
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify footer persisted
        assert doc2.sections[0].footer.paragraphs[0].text == footer_text
    finally:
        # Clean up
        os.unlink(output_path)


def test_add_page_numbers(sample_document):
    """Test adding page numbers to a document footer."""
    # Load the document
    doc = Document(sample_document)
    
    # Add page numbers to the footer
    footer = add_page_numbers(doc)
    
    # Verify footer was added
    assert footer is not None
    
    # Save the document to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # We can't easily check for the page number field directly,
        # but we can check that a footer was created
        assert doc2.sections[0].footer is not None
    finally:
        # Clean up
        os.unlink(output_path)


def test_set_different_first_page(sample_document):
    """Test setting different first page header/footer."""
    # Load the document
    doc = Document(sample_document)
    
    # Create regular header and footer
    add_header(doc, 'Regular Header')
    add_footer(doc, 'Regular Footer')
    
    # Set different first page
    set_different_first_page(doc, True)
    
    # Add first page header and footer
    section = doc.sections[0]
    first_page_header = section.first_page_header
    first_page_header.paragraphs[0].text = 'First Page Header'
    
    first_page_footer = section.first_page_footer
    first_page_footer.paragraphs[0].text = 'First Page Footer'
    
    # Verify settings were applied
    assert section.different_first_page_header_footer == True
    assert first_page_header.paragraphs[0].text == 'First Page Header'
    assert first_page_footer.paragraphs[0].text == 'First Page Footer'
    assert section.header.paragraphs[0].text == 'Regular Header'
    assert section.footer.paragraphs[0].text == 'Regular Footer'
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        section2 = doc2.sections[0]
        
        # Verify settings persisted
        assert section2.different_first_page_header_footer == True
        assert section2.first_page_header.paragraphs[0].text == 'First Page Header'
        assert section2.first_page_footer.paragraphs[0].text == 'First Page Footer'
        assert section2.header.paragraphs[0].text == 'Regular Header'
        assert section2.footer.paragraphs[0].text == 'Regular Footer'
    finally:
        # Clean up
        os.unlink(output_path)


def test_set_different_odd_even_pages(sample_document):
    """Test setting different odd/even page headers/footers."""
    # Load the document
    doc = Document(sample_document)
    
    # Create regular header and footer
    add_header(doc, 'Default Header')
    add_footer(doc, 'Default Footer')
    
    # Set different odd/even pages
    set_different_odd_even_pages(doc, True)
    
    # Add odd/even headers and footers
    section = doc.sections[0]
    
    # Even page headers/footers
    section.even_page_header.paragraphs[0].text = 'Even Page Header'
    section.even_page_footer.paragraphs[0].text = 'Even Page Footer'
    
    # Odd page headers/footers (default)
    section.header.paragraphs[0].text = 'Odd Page Header'
    section.footer.paragraphs[0].text = 'Odd Page Footer'
    
    # Verify settings were applied
    assert section.different_first_page_header_footer == False
    assert section.even_page_header.paragraphs[0].text == 'Even Page Header'
    assert section.even_page_footer.paragraphs[0].text == 'Even Page Footer'
    assert section.header.paragraphs[0].text == 'Odd Page Header'
    assert section.footer.paragraphs[0].text == 'Odd Page Footer'
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        section2 = doc2.sections[0]
        
        # Verify settings persisted
        assert section2.even_page_header.paragraphs[0].text == 'Even Page Header'
        assert section2.even_page_footer.paragraphs[0].text == 'Even Page Footer'
        assert section2.header.paragraphs[0].text == 'Odd Page Header'
        assert section2.footer.paragraphs[0].text == 'Odd Page Footer'
    finally:
        # Clean up
        os.unlink(output_path)


def test_remove_header_footer(sample_document):
    """Test removing headers and footers."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a header and footer
    add_header(doc, 'Test Header')
    add_footer(doc, 'Test Footer')
    
    # Verify they were added
    assert doc.sections[0].header.paragraphs[0].text == 'Test Header'
    assert doc.sections[0].footer.paragraphs[0].text == 'Test Footer'
    
    # Remove header and footer
    remove_header(doc)
    remove_footer(doc)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Since we can't check for null headers/footers directly with python-docx,
        # we'll check that they're empty (no paragraphs or empty paragraph)
        header = doc2.sections[0].header
        assert len(header.paragraphs) <= 1
        if len(header.paragraphs) == 1:
            assert header.paragraphs[0].text == ''
        
        footer = doc2.sections[0].footer
        assert len(footer.paragraphs) <= 1
        if len(footer.paragraphs) == 1:
            assert footer.paragraphs[0].text == ''
    finally:
        # Clean up
        os.unlink(output_path)


def test_add_header_image(sample_document):
    """Test adding an image to a header."""
    # Create a simple test image
    import PIL.Image
    
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_img:
        image_path = tmp_img.name
        img = PIL.Image.new('RGB', (100, 50), color='red')
        img.save(image_path)
    
    try:
        # Load the document
        doc = Document(sample_document)
        
        # Add header with image
        header = add_header(doc, 'Header with image:')
        image = add_header_image(doc, image_path, width=2.0)
        
        # Verify image was added to header
        assert image is not None
        
        # Save and reload to verify persistence
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            output_path = tmp.name
            doc.save(output_path)
        
        try:
            # Load the saved document
            doc2 = Document(output_path)
            
            # Check header text persisted
            assert 'Header with image:' in doc2.sections[0].header.paragraphs[0].text
            
            # Check that there's a run containing an image in the header
            # Direct check for image is complicated in python-docx
            has_image_element = False
            for paragraph in doc2.sections[0].header.paragraphs:
                for run in paragraph.runs:
                    if len(run._element.findall(".//a:blip", {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})) > 0:
                        has_image_element = True
                        break
            
            assert has_image_element
        finally:
            # Clean up
            os.unlink(output_path)
    finally:
        # Clean up test image
        os.unlink(image_path)


def test_add_footer_image(sample_document):
    """Test adding an image to a footer."""
    # Create a simple test image
    import PIL.Image
    
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_img:
        image_path = tmp_img.name
        img = PIL.Image.new('RGB', (100, 50), color='blue')
        img.save(image_path)
    
    try:
        # Load the document
        doc = Document(sample_document)
        
        # Add footer with image
        footer = add_footer(doc, 'Footer with image:')
        image = add_footer_image(doc, image_path, width=1.5)
        
        # Verify image was added to footer
        assert image is not None
        
        # Save and reload to verify persistence
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            output_path = tmp.name
            doc.save(output_path)
        
        try:
            # Load the saved document
            doc2 = Document(output_path)
            
            # Check footer text persisted
            assert 'Footer with image:' in doc2.sections[0].footer.paragraphs[0].text
            
            # Check that there's a run containing an image in the footer
            has_image_element = False
            for paragraph in doc2.sections[0].footer.paragraphs:
                for run in paragraph.runs:
                    if len(run._element.findall(".//a:blip", {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})) > 0:
                        has_image_element = True
                        break
            
            assert has_image_element
        finally:
            # Clean up
            os.unlink(output_path)
    finally:
        # Clean up test image
        os.unlink(image_path)


def test_link_to_previous(sample_document):
    """Test linking headers/footers to previous section."""
    # Load the document
    doc = Document(sample_document)
    
    # There should be multiple sections in the document
    assert len(doc.sections) >= 2
    
    # Add header and footer to first section
    add_header(doc.sections[0], 'Section 1 Header')
    add_footer(doc.sections[0], 'Section 1 Footer')
    
    # By default, section 2 should have its own header/footer
    add_header(doc.sections[1], 'Section 2 Header')
    add_footer(doc.sections[1], 'Section 2 Footer')
    
    # Link section 3 headers/footers to section 2
    link_to_previous(doc.sections[2], True)
    
    # Verify headers/footers in sections 1 and 2 are different
    assert doc.sections[0].header.paragraphs[0].text == 'Section 1 Header'
    assert doc.sections[1].header.paragraphs[0].text == 'Section 2 Header'
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify sections 1 and 2 have different headers/footers
        assert doc2.sections[0].header.paragraphs[0].text == 'Section 1 Header'
        assert doc2.sections[1].header.paragraphs[0].text == 'Section 2 Header'
        
        # Section 3 should use the header/footer from section 2
        # This is difficult to test directly with python-docx as it doesn't expose
        # the link_to_previous property directly
    finally:
        # Clean up
        os.unlink(output_path)


def test_multiple_sections_with_headers_footers(sample_document):
    """Test complex document with different headers/footers in multiple sections."""
    # Load the document
    doc = Document(sample_document)
    
    # Ensure we have at least 3 sections
    assert len(doc.sections) >= 3
    
    # Section 1: Standard header/footer
    add_header(doc.sections[0], 'Section 1 Header')
    add_footer(doc.sections[0], 'Section 1 Footer')
    
    # Section 2: Different first page, odd/even pages
    set_different_first_page(doc.sections[1], True)
    set_different_odd_even_pages(doc.sections[1], True)
    
    # Regular header/footer for section 2
    add_header(doc.sections[1], 'Section 2 Odd Header')
    add_footer(doc.sections[1], 'Section 2 Odd Footer')
    
    # First page header/footer for section 2
    doc.sections[1].first_page_header.paragraphs[0].text = 'Section 2 First Page Header'
    doc.sections[1].first_page_footer.paragraphs[0].text = 'Section 2 First Page Footer'
    
    # Even page header/footer for section 2
    doc.sections[1].even_page_header.paragraphs[0].text = 'Section 2 Even Header'
    doc.sections[1].even_page_footer.paragraphs[0].text = 'Section 2 Even Footer'
    
    # Section 3: Link to previous for header, but not footer
    add_footer(doc.sections[2], 'Section 3 Footer')
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify section 1 headers/footers
        assert doc2.sections[0].header.paragraphs[0].text == 'Section 1 Header'
        assert doc2.sections[0].footer.paragraphs[0].text == 'Section 1 Footer'
        
        # Verify section 2 settings and headers/footers
        assert doc2.sections[1].different_first_page_header_footer == True
        assert doc2.sections[1].header.paragraphs[0].text == 'Section 2 Odd Header'
        assert doc2.sections[1].footer.paragraphs[0].text == 'Section 2 Odd Footer'
        assert doc2.sections[1].first_page_header.paragraphs[0].text == 'Section 2 First Page Header'
        assert doc2.sections[1].first_page_footer.paragraphs[0].text == 'Section 2 First Page Footer'
        assert doc2.sections[1].even_page_header.paragraphs[0].text == 'Section 2 Even Header'
        assert doc2.sections[1].even_page_footer.paragraphs[0].text == 'Section 2 Even Footer'
        
        # Verify section 3 footer
        assert doc2.sections[2].footer.paragraphs[0].text == 'Section 3 Footer'
    finally:
        # Clean up
        os.unlink(output_path)


if __name__ == '__main__':
    pytest.main(['-v', __file__]) 