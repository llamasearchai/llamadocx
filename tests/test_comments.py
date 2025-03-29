"""
Comments tests for LlamaDocx.

This module contains tests for the comment handling functionality of the LlamaDocx package.
"""

import os
import tempfile
import pytest
from docx import Document
from datetime import datetime

from llamadocx.comments import (
    add_comment,
    get_comments,
    get_comment_by_id,
    update_comment,
    delete_comment,
    reply_to_comment,
    get_comment_replies
)


@pytest.fixture
def sample_document():
    """Create a sample document for testing comments."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        
        # Add heading
        doc.add_heading('Comments Test Document', level=1)
        
        # Add paragraphs with content
        doc.add_paragraph('This is paragraph 1 for testing comments.')
        doc.add_paragraph('This is paragraph 2 for testing comments.')
        doc.add_paragraph('This is paragraph 3 for testing comments.')
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


def test_add_comment(sample_document):
    """Test adding a comment to text in a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a comment to paragraph 1
    paragraph = doc.paragraphs[1]
    comment_text = "This is a test comment"
    author = "Test Author"
    comment = add_comment(doc, paragraph, comment_text, author)
    
    # Verify comment was created
    assert comment is not None
    assert comment.text == comment_text
    assert comment.author == author
    assert isinstance(comment.date, datetime)  # Should be a datetime object
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get all comments and verify
        comments = get_comments(doc2)
        assert len(comments) == 1
        assert comments[0].text == comment_text
        assert comments[0].author == author
    finally:
        # Clean up
        os.unlink(output_path)


def test_add_multiple_comments(sample_document):
    """Test adding multiple comments to different parts of a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add comments to different paragraphs
    paragraph1 = doc.paragraphs[1]
    comment1 = add_comment(doc, paragraph1, "Comment on paragraph 1", "Author 1")
    
    paragraph2 = doc.paragraphs[2]
    comment2 = add_comment(doc, paragraph2, "Comment on paragraph 2", "Author 2")
    
    paragraph3 = doc.paragraphs[3]
    comment3 = add_comment(doc, paragraph3, "Comment on paragraph 3", "Author 3")
    
    # Verify comments were created
    assert comment1 is not None
    assert comment2 is not None
    assert comment3 is not None
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get all comments and verify
        comments = get_comments(doc2)
        assert len(comments) == 3
        
        # Verify content of comments
        comment_texts = [c.text for c in comments]
        assert "Comment on paragraph 1" in comment_texts
        assert "Comment on paragraph 2" in comment_texts
        assert "Comment on paragraph 3" in comment_texts
        
        # Verify authors
        comment_authors = [c.author for c in comments]
        assert "Author 1" in comment_authors
        assert "Author 2" in comment_authors
        assert "Author 3" in comment_authors
    finally:
        # Clean up
        os.unlink(output_path)


def test_get_comment_by_id(sample_document):
    """Test retrieving a comment by its ID."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a comment
    paragraph = doc.paragraphs[1]
    comment = add_comment(doc, paragraph, "Comment to retrieve by ID", "Test Author")
    
    # Get comment ID
    comment_id = comment.id
    
    # Retrieve the comment by ID
    retrieved_comment = get_comment_by_id(doc, comment_id)
    
    # Verify comment was retrieved correctly
    assert retrieved_comment is not None
    assert retrieved_comment.id == comment_id
    assert retrieved_comment.text == "Comment to retrieve by ID"
    assert retrieved_comment.author == "Test Author"
    
    # Test retrieving a non-existent comment ID
    non_existent_id = "non-existent-id"
    none_comment = get_comment_by_id(doc, non_existent_id)
    assert none_comment is None


def test_update_comment(sample_document):
    """Test updating an existing comment."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a comment
    paragraph = doc.paragraphs[1]
    original_text = "Original comment text"
    original_author = "Original Author"
    comment = add_comment(doc, paragraph, original_text, original_author)
    
    # Update the comment
    updated_text = "Updated comment text"
    updated_author = "Updated Author"
    update_result = update_comment(doc, comment.id, updated_text, updated_author)
    
    # Verify update succeeded
    assert update_result is True
    
    # Retrieve the comment and verify updates
    updated_comment = get_comment_by_id(doc, comment.id)
    assert updated_comment is not None
    assert updated_comment.text == updated_text
    assert updated_comment.author == updated_author
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Retrieve comment and verify updates persisted
        persisted_comment = get_comment_by_id(doc2, comment.id)
        assert persisted_comment is not None
        assert persisted_comment.text == updated_text
        assert persisted_comment.author == updated_author
    finally:
        # Clean up
        os.unlink(output_path)


def test_delete_comment(sample_document):
    """Test deleting a comment."""
    # Load the document
    doc = Document(sample_document)
    
    # Add comments
    paragraph1 = doc.paragraphs[1]
    comment1 = add_comment(doc, paragraph1, "Comment to keep", "Author 1")
    
    paragraph2 = doc.paragraphs[2]
    comment2 = add_comment(doc, paragraph2, "Comment to delete", "Author 2")
    
    # Verify both comments exist
    assert len(get_comments(doc)) == 2
    
    # Delete the second comment
    delete_result = delete_comment(doc, comment2.id)
    
    # Verify deletion succeeded
    assert delete_result is True
    
    # Verify only one comment remains
    remaining_comments = get_comments(doc)
    assert len(remaining_comments) == 1
    assert remaining_comments[0].id == comment1.id
    assert remaining_comments[0].text == "Comment to keep"
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify deletion persisted
        persisted_comments = get_comments(doc2)
        assert len(persisted_comments) == 1
        assert persisted_comments[0].text == "Comment to keep"
    finally:
        # Clean up
        os.unlink(output_path)


def test_reply_to_comment(sample_document):
    """Test replying to a comment."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a comment
    paragraph = doc.paragraphs[1]
    parent_comment = add_comment(doc, paragraph, "Parent comment", "Parent Author")
    
    # Add a reply to the comment
    reply = reply_to_comment(doc, parent_comment.id, "Reply to parent", "Reply Author")
    
    # Verify reply was created
    assert reply is not None
    assert reply.text == "Reply to parent"
    assert reply.author == "Reply Author"
    
    # Get replies to the parent comment
    replies = get_comment_replies(doc, parent_comment.id)
    
    # Verify replies
    assert len(replies) == 1
    assert replies[0].text == "Reply to parent"
    assert replies[0].author == "Reply Author"
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get parent comment
        parent = get_comment_by_id(doc2, parent_comment.id)
        assert parent is not None
        
        # Get replies and verify persistence
        persisted_replies = get_comment_replies(doc2, parent.id)
        assert len(persisted_replies) == 1
        assert persisted_replies[0].text == "Reply to parent"
        assert persisted_replies[0].author == "Reply Author"
    finally:
        # Clean up
        os.unlink(output_path)


def test_multiple_replies(sample_document):
    """Test adding multiple replies to a comment."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a comment
    paragraph = doc.paragraphs[1]
    parent_comment = add_comment(doc, paragraph, "Parent comment", "Parent Author")
    
    # Add multiple replies
    reply1 = reply_to_comment(doc, parent_comment.id, "First reply", "Reply Author 1")
    reply2 = reply_to_comment(doc, parent_comment.id, "Second reply", "Reply Author 2")
    reply3 = reply_to_comment(doc, parent_comment.id, "Third reply", "Reply Author 3")
    
    # Verify replies were created
    assert reply1 is not None
    assert reply2 is not None
    assert reply3 is not None
    
    # Get all replies
    replies = get_comment_replies(doc, parent_comment.id)
    
    # Verify replies
    assert len(replies) == 3
    
    # Verify content and order of replies
    assert replies[0].text == "First reply"
    assert replies[1].text == "Second reply"
    assert replies[2].text == "Third reply"
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get replies and verify persistence
        persisted_replies = get_comment_replies(doc2, parent_comment.id)
        assert len(persisted_replies) == 3
        
        # Verify content of persisted replies
        reply_texts = [r.text for r in persisted_replies]
        assert "First reply" in reply_texts
        assert "Second reply" in reply_texts
        assert "Third reply" in reply_texts
    finally:
        # Clean up
        os.unlink(output_path)


def test_nested_replies(sample_document):
    """Test nested replies (replying to a reply)."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a comment
    paragraph = doc.paragraphs[1]
    parent_comment = add_comment(doc, paragraph, "Parent comment", "Parent Author")
    
    # Add a reply to the parent
    first_reply = reply_to_comment(doc, parent_comment.id, "First level reply", "Reply Author 1")
    
    # Add a reply to the first reply (nested reply)
    nested_reply = reply_to_comment(doc, first_reply.id, "Nested reply", "Nested Author")
    
    # Verify nested reply was created
    assert nested_reply is not None
    assert nested_reply.text == "Nested reply"
    
    # Get replies to the first reply
    nested_replies = get_comment_replies(doc, first_reply.id)
    
    # Verify nested replies
    assert len(nested_replies) == 1
    assert nested_replies[0].text == "Nested reply"
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get the first reply
        persisted_first_reply = get_comment_by_id(doc2, first_reply.id)
        assert persisted_first_reply is not None
        
        # Get nested replies and verify persistence
        persisted_nested_replies = get_comment_replies(doc2, persisted_first_reply.id)
        assert len(persisted_nested_replies) == 1
        assert persisted_nested_replies[0].text == "Nested reply"
    finally:
        # Clean up
        os.unlink(output_path)


def test_delete_comment_with_replies(sample_document):
    """Test deleting a comment that has replies."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a comment
    paragraph = doc.paragraphs[1]
    parent_comment = add_comment(doc, paragraph, "Parent comment", "Parent Author")
    
    # Add replies
    reply1 = reply_to_comment(doc, parent_comment.id, "First reply", "Reply Author 1")
    reply2 = reply_to_comment(doc, parent_comment.id, "Second reply", "Reply Author 2")
    
    # Verify comment and replies exist
    assert get_comment_by_id(doc, parent_comment.id) is not None
    assert len(get_comment_replies(doc, parent_comment.id)) == 2
    
    # Delete the parent comment (should also delete replies)
    delete_result = delete_comment(doc, parent_comment.id)
    
    # Verify deletion succeeded
    assert delete_result is True
    
    # Verify parent comment and replies are gone
    assert get_comment_by_id(doc, parent_comment.id) is None
    assert get_comment_by_id(doc, reply1.id) is None
    assert get_comment_by_id(doc, reply2.id) is None
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify no comments exist
        assert len(get_comments(doc2)) == 0
    finally:
        # Clean up
        os.unlink(output_path)


if __name__ == '__main__':
    pytest.main(['-v', __file__]) 