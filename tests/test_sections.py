"""
Sections tests for LlamaDocx.

This module contains tests for the section handling functionality of the LlamaDocx package.
"""

import os
import tempfile
import pytest
from docx import Document
from docx.enum.section import WD_SECTION_START, WD_ORIENTATION
from docx.shared import Inches

from llamadocx.sections import (
    add_section_break,
    get_section_count,
    get_section,
    set_section_orientation,
    set_page_size,
    set_page_margins,
    set_columns,
    set_header_distance,
    set_footer_distance,
    set_gutter_margin,
    set_section_start_type
)


@pytest.fixture
def sample_document():
    """Create a sample document for testing section operations."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        
        # Add heading
        doc.add_heading('Section Handling Test Document', level=1)
        
        # Add some content to the first section
        doc.add_paragraph('This is content in section 1.')
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


def test_add_section_break(sample_document):
    """Test adding section breaks to a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Get initial section count
    initial_section_count = len(doc.sections)
    assert initial_section_count == 1
    
    # Add different types of section breaks
    break_types = [
        WD_SECTION_START.NEW_PAGE,
        WD_SECTION_START.EVEN_PAGE,
        WD_SECTION_START.ODD_PAGE,
        WD_SECTION_START.CONTINUOUS
    ]
    
    for break_type in break_types:
        # Add content before break
        doc.add_paragraph(f'Content before {break_type} section break')
        
        # Add section break
        add_section_break(doc, break_type)
        
        # Add content after break
        doc.add_paragraph(f'Content after {break_type} section break')
    
    # Verify section count increased
    assert len(doc.sections) == initial_section_count + len(break_types)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify section count persisted
        assert len(doc2.sections) == initial_section_count + len(break_types)
        
        # Verify section break types
        for i, break_type in enumerate(break_types, start=1):
            assert doc2.sections[i].start_type == break_type
    finally:
        # Clean up
        os.unlink(output_path)


def test_get_section_count(sample_document):
    """Test getting the section count of a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Get initial section count
    initial_count = get_section_count(doc)
    assert initial_count == 1
    
    # Add some section breaks
    add_section_break(doc, WD_SECTION_START.NEW_PAGE)
    add_section_break(doc, WD_SECTION_START.CONTINUOUS)
    
    # Verify count increased
    assert get_section_count(doc) == initial_count + 2


def test_get_section(sample_document):
    """Test getting a specific section from a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add some section breaks with different types
    add_section_break(doc, WD_SECTION_START.NEW_PAGE)
    add_section_break(doc, WD_SECTION_START.CONTINUOUS)
    
    # Get sections by index
    section0 = get_section(doc, 0)
    section1 = get_section(doc, 1)
    section2 = get_section(doc, 2)
    
    # Verify the sections were retrieved correctly
    assert section0 is doc.sections[0]
    assert section1 is doc.sections[1]
    assert section2 is doc.sections[2]
    
    assert section1.start_type == WD_SECTION_START.NEW_PAGE
    assert section2.start_type == WD_SECTION_START.CONTINUOUS
    
    # Test error handling for invalid index
    with pytest.raises(IndexError):
        get_section(doc, 3)  # Should be out of range


def test_set_section_orientation(sample_document):
    """Test setting section orientation."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a new section
    add_section_break(doc, WD_SECTION_START.NEW_PAGE)
    
    # Set different orientations for each section
    set_section_orientation(doc.sections[0], WD_ORIENTATION.PORTRAIT)
    set_section_orientation(doc.sections[1], WD_ORIENTATION.LANDSCAPE)
    
    # Verify orientations were set
    assert doc.sections[0].orientation == WD_ORIENTATION.PORTRAIT
    assert doc.sections[1].orientation == WD_ORIENTATION.LANDSCAPE
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify orientations persisted
        assert doc2.sections[0].orientation == WD_ORIENTATION.PORTRAIT
        assert doc2.sections[1].orientation == WD_ORIENTATION.LANDSCAPE
    finally:
        # Clean up
        os.unlink(output_path)


def test_set_page_size(sample_document):
    """Test setting page size for sections."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a new section
    add_section_break(doc, WD_SECTION_START.NEW_PAGE)
    
    # Set different page sizes for each section
    set_page_size(doc.sections[0], width=Inches(8.5), height=Inches(11))  # Letter
    set_page_size(doc.sections[1], width=Inches(11), height=Inches(17))   # Ledger
    
    # Verify page sizes were set
    assert doc.sections[0].page_width == Inches(8.5)
    assert doc.sections[0].page_height == Inches(11)
    assert doc.sections[1].page_width == Inches(11)
    assert doc.sections[1].page_height == Inches(17)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify page sizes persisted
        assert doc2.sections[0].page_width == Inches(8.5)
        assert doc2.sections[0].page_height == Inches(11)
        assert doc2.sections[1].page_width == Inches(11)
        assert doc2.sections[1].page_height == Inches(17)
    finally:
        # Clean up
        os.unlink(output_path)


def test_set_page_margins(sample_document):
    """Test setting page margins for sections."""
    # Load the document
    doc = Document(sample_document)
    
    # Set margins for the first section
    set_page_margins(doc.sections[0], 
                     top=Inches(1),
                     right=Inches(1.5),
                     bottom=Inches(1.25),
                     left=Inches(1.75))
    
    # Verify margins were set
    assert doc.sections[0].top_margin == Inches(1)
    assert doc.sections[0].right_margin == Inches(1.5)
    assert doc.sections[0].bottom_margin == Inches(1.25)
    assert doc.sections[0].left_margin == Inches(1.75)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify margins persisted
        assert doc2.sections[0].top_margin == Inches(1)
        assert doc2.sections[0].right_margin == Inches(1.5)
        assert doc2.sections[0].bottom_margin == Inches(1.25)
        assert doc2.sections[0].left_margin == Inches(1.75)
    finally:
        # Clean up
        os.unlink(output_path)


def test_set_columns(sample_document):
    """Test setting columns for sections."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a new section
    add_section_break(doc, WD_SECTION_START.NEW_PAGE)
    
    # Set different column layouts for each section
    set_columns(doc.sections[0], count=1)  # Single column
    set_columns(doc.sections[1], count=2, spacing=Inches(0.5))  # Two columns
    
    # Verify column settings were applied
    assert doc.sections[0].column_count == 1
    assert doc.sections[1].column_count == 2
    assert doc.sections[1].column_width == None  # Word calculates this automatically
    assert doc.sections[1].column_spacing == Inches(0.5)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify column settings persisted
        assert doc2.sections[0].column_count == 1
        assert doc2.sections[1].column_count == 2
        assert doc2.sections[1].column_spacing == Inches(0.5)
    finally:
        # Clean up
        os.unlink(output_path)


def test_set_header_footer_distance(sample_document):
    """Test setting header and footer distance from edge."""
    # Load the document
    doc = Document(sample_document)
    
    # Set header and footer distance
    set_header_distance(doc.sections[0], Inches(0.5))
    set_footer_distance(doc.sections[0], Inches(0.5))
    
    # Verify distances were set
    assert doc.sections[0].header_distance == Inches(0.5)
    assert doc.sections[0].footer_distance == Inches(0.5)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify distances persisted
        assert doc2.sections[0].header_distance == Inches(0.5)
        assert doc2.sections[0].footer_distance == Inches(0.5)
    finally:
        # Clean up
        os.unlink(output_path)


def test_set_gutter_margin(sample_document):
    """Test setting gutter margin for sections."""
    # Load the document
    doc = Document(sample_document)
    
    # Set gutter margin
    set_gutter_margin(doc.sections[0], Inches(0.75))
    
    # Verify gutter margin was set
    assert doc.sections[0].gutter == Inches(0.75)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify gutter margin persisted
        assert doc2.sections[0].gutter == Inches(0.75)
    finally:
        # Clean up
        os.unlink(output_path)


def test_set_section_start_type(sample_document):
    """Test setting section start type."""
    # Load the document
    doc = Document(sample_document)
    
    # Add some section breaks
    add_section_break(doc, WD_SECTION_START.NEW_PAGE)
    add_section_break(doc, WD_SECTION_START.NEW_PAGE)
    
    # Set different start types
    set_section_start_type(doc.sections[1], WD_SECTION_START.CONTINUOUS)
    set_section_start_type(doc.sections[2], WD_SECTION_START.ODD_PAGE)
    
    # Verify start types were set
    assert doc.sections[1].start_type == WD_SECTION_START.CONTINUOUS
    assert doc.sections[2].start_type == WD_SECTION_START.ODD_PAGE
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify start types persisted
        assert doc2.sections[1].start_type == WD_SECTION_START.CONTINUOUS
        assert doc2.sections[2].start_type == WD_SECTION_START.ODD_PAGE
    finally:
        # Clean up
        os.unlink(output_path)


def test_complex_section_settings(sample_document):
    """Test combining multiple section settings."""
    # Load the document
    doc = Document(sample_document)
    
    # Add two new sections
    add_section_break(doc, WD_SECTION_START.NEW_PAGE)
    add_section_break(doc, WD_SECTION_START.CONTINUOUS)
    
    # Configure section 1: Portrait, Letter size, 1-inch margins, single column
    section1 = doc.sections[0]
    set_section_orientation(section1, WD_ORIENTATION.PORTRAIT)
    set_page_size(section1, width=Inches(8.5), height=Inches(11))
    set_page_margins(section1, top=Inches(1), right=Inches(1), bottom=Inches(1), left=Inches(1))
    set_columns(section1, count=1)
    
    # Configure section 2: Landscape, Legal size, 0.75-inch margins, two columns
    section2 = doc.sections[1]
    set_section_orientation(section2, WD_ORIENTATION.LANDSCAPE)
    set_page_size(section2, width=Inches(14), height=Inches(8.5))  # Landscape measurements
    set_page_margins(section2, top=Inches(0.75), right=Inches(0.75), bottom=Inches(0.75), left=Inches(0.75))
    set_columns(section2, count=2, spacing=Inches(0.25))
    
    # Configure section 3: Portrait, custom size, narrow margins, three columns
    section3 = doc.sections[2]
    set_section_orientation(section3, WD_ORIENTATION.PORTRAIT)
    set_page_size(section3, width=Inches(7), height=Inches(9))
    set_page_margins(section3, top=Inches(0.5), right=Inches(0.5), bottom=Inches(0.5), left=Inches(0.5))
    set_columns(section3, count=3, spacing=Inches(0.2))
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify section 1 settings
        assert doc2.sections[0].orientation == WD_ORIENTATION.PORTRAIT
        assert doc2.sections[0].page_width == Inches(8.5)
        assert doc2.sections[0].page_height == Inches(11)
        assert doc2.sections[0].left_margin == Inches(1)
        assert doc2.sections[0].column_count == 1
        
        # Verify section 2 settings
        assert doc2.sections[1].orientation == WD_ORIENTATION.LANDSCAPE
        assert doc2.sections[1].page_width == Inches(14)
        assert doc2.sections[1].page_height == Inches(8.5)
        assert doc2.sections[1].left_margin == Inches(0.75)
        assert doc2.sections[1].column_count == 2
        assert doc2.sections[1].column_spacing == Inches(0.25)
        
        # Verify section 3 settings
        assert doc2.sections[2].orientation == WD_ORIENTATION.PORTRAIT
        assert doc2.sections[2].page_width == Inches(7)
        assert doc2.sections[2].page_height == Inches(9)
        assert doc2.sections[2].left_margin == Inches(0.5)
        assert doc2.sections[2].column_count == 3
        assert doc2.sections[2].column_spacing == Inches(0.2)
    finally:
        # Clean up
        os.unlink(output_path)


if __name__ == '__main__':
    pytest.main(['-v', __file__]) 