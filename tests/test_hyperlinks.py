"""
Hyperlinks tests for LlamaDocx.

This module contains tests for the hyperlink functionality of the LlamaDocx package.
"""

import os
import tempfile
import pytest
from docx import Document

from llamadocx.hyperlinks import (
    add_hyperlink,
    add_internal_hyperlink,
    add_email_hyperlink,
    get_hyperlinks,
    update_hyperlink,
    remove_hyperlink,
    get_hyperlink_url,
    add_bookmark,
    get_bookmarks,
    remove_bookmark
)


@pytest.fixture
def sample_document():
    """Create a sample document for testing hyperlinks."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        
        # Add heading
        doc.add_heading('Hyperlinks Test Document', level=1)
        
        # Add paragraphs
        doc.add_paragraph('This is a paragraph for testing external hyperlinks.')
        doc.add_paragraph('This is a paragraph for testing internal hyperlinks.')
        doc.add_paragraph('This is a paragraph for testing email hyperlinks.')
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


def test_add_hyperlink(sample_document):
    """Test adding an external hyperlink to a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a hyperlink to paragraph 1
    paragraph = doc.paragraphs[1]
    url = "https://www.example.com"
    text = "Example Website"
    hyperlink = add_hyperlink(doc, paragraph, url, text)
    
    # Verify hyperlink was created
    assert hyperlink is not None
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get all hyperlinks
        hyperlinks = get_hyperlinks(doc2)
        
        # Verify hyperlink exists
        assert len(hyperlinks) >= 1
        
        # Verify hyperlink properties
        hyperlink_found = False
        for link in hyperlinks:
            if get_hyperlink_url(link) == url:
                hyperlink_found = True
                assert text in link.text
                break
                
        assert hyperlink_found, "Added hyperlink not found in document"
    finally:
        # Clean up
        os.unlink(output_path)


def test_add_internal_hyperlink(sample_document):
    """Test adding an internal hyperlink (bookmark) to a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a bookmark
    paragraph = doc.paragraphs[2]
    bookmark_name = "test_bookmark"
    bookmark = add_bookmark(doc, paragraph, bookmark_name)
    
    # Add an internal hyperlink to the bookmark
    link_paragraph = doc.paragraphs[1]
    link_text = "Jump to bookmark"
    internal_link = add_internal_hyperlink(doc, link_paragraph, bookmark_name, link_text)
    
    # Verify link was created
    assert internal_link is not None
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get all hyperlinks
        hyperlinks = get_hyperlinks(doc2)
        
        # Verify hyperlink exists
        assert len(hyperlinks) >= 1
        
        # Verify bookmark exists
        bookmarks = get_bookmarks(doc2)
        assert bookmark_name in bookmarks
        
        # Verify link to bookmark
        link_found = False
        for link in hyperlinks:
            url = get_hyperlink_url(link)
            if url and bookmark_name in url:
                link_found = True
                assert link_text in link.text
                break
                
        assert link_found, "Internal hyperlink not found"
    finally:
        # Clean up
        os.unlink(output_path)


def test_add_email_hyperlink(sample_document):
    """Test adding an email hyperlink to a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add an email hyperlink
    paragraph = doc.paragraphs[3]
    email = "test@example.com"
    subject = "Test Subject"
    text = "Email Us"
    email_link = add_email_hyperlink(doc, paragraph, email, text, subject)
    
    # Verify link was created
    assert email_link is not None
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get all hyperlinks
        hyperlinks = get_hyperlinks(doc2)
        
        # Verify email hyperlink exists
        link_found = False
        for link in hyperlinks:
            url = get_hyperlink_url(link)
            if url and url.startswith("mailto:") and email in url:
                link_found = True
                assert text in link.text
                if subject:
                    assert "subject=" in url
                break
                
        assert link_found, "Email hyperlink not found"
    finally:
        # Clean up
        os.unlink(output_path)


def test_update_hyperlink(sample_document):
    """Test updating an existing hyperlink."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a hyperlink first
    paragraph = doc.paragraphs[1]
    original_url = "https://www.example.com"
    original_text = "Original Link"
    hyperlink = add_hyperlink(doc, paragraph, original_url, original_text)
    
    # Get hyperlinks
    hyperlinks = get_hyperlinks(doc)
    
    # Find the hyperlink we just added
    link_to_update = None
    for link in hyperlinks:
        if get_hyperlink_url(link) == original_url:
            link_to_update = link
            break
    
    assert link_to_update is not None
    
    # Update the hyperlink
    new_url = "https://www.updated-example.com"
    new_text = "Updated Link"
    update_result = update_hyperlink(doc, link_to_update, new_url, new_text)
    
    # Verify update was successful
    assert update_result is True
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get all hyperlinks
        hyperlinks = get_hyperlinks(doc2)
        
        # Verify updated hyperlink exists
        updated_link_found = False
        original_link_found = False
        
        for link in hyperlinks:
            url = get_hyperlink_url(link)
            if url == new_url:
                updated_link_found = True
                assert new_text in link.text
            elif url == original_url:
                original_link_found = True
        
        assert updated_link_found, "Updated hyperlink not found"
        assert not original_link_found, "Original hyperlink still exists"
    finally:
        # Clean up
        os.unlink(output_path)


def test_remove_hyperlink(sample_document):
    """Test removing a hyperlink from a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add multiple hyperlinks
    paragraph1 = doc.paragraphs[1]
    url1 = "https://www.example1.com"
    hyperlink1 = add_hyperlink(doc, paragraph1, url1, "Link 1")
    
    paragraph2 = doc.paragraphs[2]
    url2 = "https://www.example2.com"
    hyperlink2 = add_hyperlink(doc, paragraph2, url2, "Link 2")
    
    # Get all hyperlinks
    hyperlinks = get_hyperlinks(doc)
    
    # Find the first hyperlink
    link_to_remove = None
    for link in hyperlinks:
        if get_hyperlink_url(link) == url1:
            link_to_remove = link
            break
    
    assert link_to_remove is not None
    
    # Remove the first hyperlink
    remove_result = remove_hyperlink(doc, link_to_remove)
    
    # Verify removal was successful
    assert remove_result is True
    
    # Get remaining hyperlinks
    remaining_hyperlinks = get_hyperlinks(doc)
    
    # Verify first hyperlink was removed
    for link in remaining_hyperlinks:
        assert get_hyperlink_url(link) != url1
    
    # Verify second hyperlink still exists
    second_link_found = False
    for link in remaining_hyperlinks:
        if get_hyperlink_url(link) == url2:
            second_link_found = True
            break
    
    assert second_link_found, "Second hyperlink should still exist"
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get all hyperlinks
        hyperlinks = get_hyperlinks(doc2)
        
        # Verify first hyperlink was removed
        for link in hyperlinks:
            assert get_hyperlink_url(link) != url1
        
        # Verify second hyperlink still exists
        second_link_found = False
        for link in hyperlinks:
            if get_hyperlink_url(link) == url2:
                second_link_found = True
                break
        
        assert second_link_found, "Second hyperlink should still exist"
    finally:
        # Clean up
        os.unlink(output_path)


def test_bookmarks(sample_document):
    """Test adding, retrieving, and removing bookmarks."""
    # Load the document
    doc = Document(sample_document)
    
    # Add multiple bookmarks
    paragraph1 = doc.paragraphs[1]
    bookmark1_name = "bookmark1"
    bookmark1 = add_bookmark(doc, paragraph1, bookmark1_name)
    
    paragraph2 = doc.paragraphs[2]
    bookmark2_name = "bookmark2"
    bookmark2 = add_bookmark(doc, paragraph2, bookmark2_name)
    
    # Verify bookmarks were added
    assert bookmark1 is not None
    assert bookmark2 is not None
    
    # Get all bookmarks
    bookmarks = get_bookmarks(doc)
    
    # Verify both bookmarks exist
    assert bookmark1_name in bookmarks
    assert bookmark2_name in bookmarks
    
    # Remove the first bookmark
    remove_result = remove_bookmark(doc, bookmark1_name)
    
    # Verify removal was successful
    assert remove_result is True
    
    # Get remaining bookmarks
    remaining_bookmarks = get_bookmarks(doc)
    
    # Verify first bookmark was removed
    assert bookmark1_name not in remaining_bookmarks
    
    # Verify second bookmark still exists
    assert bookmark2_name in remaining_bookmarks
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get all bookmarks
        bookmarks = get_bookmarks(doc2)
        
        # Verify first bookmark was removed
        assert bookmark1_name not in bookmarks
        
        # Verify second bookmark still exists
        assert bookmark2_name in bookmarks
    finally:
        # Clean up
        os.unlink(output_path)


def test_complex_hyperlinks(sample_document):
    """Test complex hyperlink scenarios with parameters and fragments."""
    # Load the document
    doc = Document(sample_document)
    
    # Add complex URL with query parameters and fragment
    paragraph = doc.paragraphs[1]
    complex_url = "https://www.example.com/search?q=test&category=docs#section2"
    text = "Complex Link"
    complex_link = add_hyperlink(doc, paragraph, complex_url, text)
    
    # Verify link was created
    assert complex_link is not None
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get all hyperlinks
        hyperlinks = get_hyperlinks(doc2)
        
        # Verify complex hyperlink exists with all components
        complex_link_found = False
        for link in hyperlinks:
            url = get_hyperlink_url(link)
            if url and url == complex_url:
                complex_link_found = True
                assert text in link.text
                assert "q=test" in url
                assert "#section2" in url
                break
                
        assert complex_link_found, "Complex hyperlink not found or parameters missing"
    finally:
        # Clean up
        os.unlink(output_path)


def test_multiple_hyperlinks_in_paragraph(sample_document):
    """Test adding multiple hyperlinks to a single paragraph."""
    # Load the document
    doc = Document(sample_document)
    
    # Get a paragraph
    paragraph = doc.paragraphs[1]
    
    # Clear existing text
    paragraph.clear()
    
    # Add text and multiple hyperlinks to the same paragraph
    paragraph.add_run("This paragraph contains ")
    link1 = add_hyperlink(doc, paragraph, "https://www.example1.com", "first link")
    paragraph.add_run(" and ")
    link2 = add_hyperlink(doc, paragraph, "https://www.example2.com", "second link")
    paragraph.add_run(" in the same paragraph.")
    
    # Verify links were created
    assert link1 is not None
    assert link2 is not None
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get all hyperlinks
        hyperlinks = get_hyperlinks(doc2)
        
        # Verify both hyperlinks exist
        urls = [get_hyperlink_url(link) for link in hyperlinks]
        assert "https://www.example1.com" in urls
        assert "https://www.example2.com" in urls
        
        # Verify paragraph text contains both links
        para_text = doc2.paragraphs[1].text
        assert "first link" in para_text
        assert "second link" in para_text
        assert "This paragraph contains" in para_text
        assert "in the same paragraph" in para_text
    finally:
        # Clean up
        os.unlink(output_path)


if __name__ == '__main__':
    pytest.main(['-v', __file__]) 