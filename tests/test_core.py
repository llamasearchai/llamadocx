"""
Core tests for LlamaDocx.

This module contains tests for the core functionality of the LlamaDocx package.
"""

import os
import tempfile
from pathlib import Path
import pytest
from docx import Document

from llamadocx.template import Template
from llamadocx.metadata import get_metadata, set_metadata
from llamadocx.search import search_text, replace_text
from llamadocx.convert import docx_to_markdown, md_to_docx


# Fixture for creating a temporary document
@pytest.fixture
def temp_docx():
    """Create a temporary DOCX file for testing."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        doc.add_heading('Test Document', level=1)
        doc.add_paragraph('This is a test paragraph for LlamaDocx.')
        doc.add_paragraph('Another paragraph with some test content.')
        
        # Create a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = 'Cell 1'
        table.cell(0, 1).text = 'Cell 2'
        table.cell(1, 0).text = 'Cell 3'
        table.cell(1, 1).text = 'Cell 4'
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


# Fixture for creating a temporary template
@pytest.fixture
def temp_template():
    """Create a temporary template document for testing."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        doc.add_heading('{{title}}', level=1)
        doc.add_paragraph('Hello, {{name}}!')
        doc.add_paragraph('Date: {{date}}')
        
        # Create a table with template fields
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = 'Item'
        table.cell(0, 1).text = 'Price'
        table.cell(1, 0).text = '{{item}}'
        table.cell(1, 1).text = '{{price}}'
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


# Fixture for creating a temporary markdown file
@pytest.fixture
def temp_markdown():
    """Create a temporary Markdown file for testing."""
    with tempfile.NamedTemporaryFile(suffix='.md', mode='w', delete=False) as tmp:
        content = """# Test Markdown Document

This is a test paragraph for LlamaDocx.

Another paragraph with **bold** and *italic* text.

## Section 1

- List item 1
- List item 2
- List item 3

## Section 2

1. Numbered item 1
2. Numbered item 2
3. Numbered item 3
"""
        tmp.write(content)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


def test_document_metadata(temp_docx):
    """Test setting and retrieving document metadata."""
    # Set metadata
    doc = Document(temp_docx)
    metadata = {
        'title': 'Test Document',
        'author': 'LlamaDocx Tester',
        'subject': 'Unit Testing',
        'keywords': 'test,docx,python',
        'category': 'Testing'
    }
    set_metadata(doc, metadata)
    doc.save(temp_docx)
    
    # Verify metadata
    doc = Document(temp_docx)
    retrieved_metadata = get_metadata(doc)
    
    assert retrieved_metadata['title'] == metadata['title']
    assert retrieved_metadata['author'] == metadata['author']
    assert retrieved_metadata['subject'] == metadata['subject']
    assert retrieved_metadata['keywords'] == metadata['keywords']
    assert retrieved_metadata['category'] == metadata['category']


def test_search_text(temp_docx):
    """Test searching for text in a document."""
    doc = Document(temp_docx)
    
    # Basic search
    results = search_text(doc, "test")
    assert len(results) > 0
    assert all('test' in result['text'].lower() for result in results)
    
    # Case-sensitive search
    results = search_text(doc, "Test", case_sensitive=True)
    assert len(results) > 0
    assert all('Test' in result['text'] for result in results)
    
    # Regex search
    results = search_text(doc, r"test\s\w+", regex=True)
    assert len(results) > 0
    assert all(result['type'] == 'paragraph' for result in results)


def test_replace_text(temp_docx):
    """Test replacing text in a document."""
    doc = Document(temp_docx)
    
    # Basic replacement
    count = replace_text(doc, "test", "demo")
    assert count > 0
    
    # Verify replacement
    text_content = "\n".join(p.text for p in doc.paragraphs)
    assert "demo" in text_content
    assert "test" not in text_content.lower()  # Should be replaced


def test_template_processing(temp_template):
    """Test template processing with field merging."""
    # Create template and data
    template = Template(temp_template)
    data = {
        'title': 'Invoice',
        'name': 'John Doe',
        'date': '2023-06-15',
        'item': 'Product ABC',
        'price': '$99.99'
    }
    
    # Process template
    template.merge_fields(data)
    
    # Save to temporary file
    with tempfile.NamedTemporaryFile(suffix='.docx') as output_file:
        output_path = output_file.name
        template.save(output_path)
        
        # Verify merged content
        doc = Document(output_path)
        
        # Check heading
        assert doc.paragraphs[0].text == data['title']
        
        # Check paragraphs
        assert f"Hello, {data['name']}!" in doc.paragraphs[1].text
        assert f"Date: {data['date']}" in doc.paragraphs[2].text
        
        # Check table
        table = doc.tables[0]
        assert table.cell(1, 0).text == data['item']
        assert table.cell(1, 1).text == data['price']


def test_format_conversion(temp_docx, temp_markdown):
    """Test converting between document formats."""
    # Test DOCX to Markdown
    md_output = tempfile.NamedTemporaryFile(suffix='.md', delete=False).name
    try:
        docx_to_markdown(temp_docx, md_output)
        
        # Verify the Markdown file exists and has content
        assert os.path.exists(md_output)
        with open(md_output, 'r') as f:
            content = f.read()
            assert '# Test Document' in content
            assert 'This is a test paragraph' in content
    finally:
        if os.path.exists(md_output):
            os.unlink(md_output)
    
    # Test Markdown to DOCX
    docx_output = tempfile.NamedTemporaryFile(suffix='.docx', delete=False).name
    try:
        md_to_docx(temp_markdown, docx_output)
        
        # Verify the DOCX file exists and has content
        assert os.path.exists(docx_output)
        doc = Document(docx_output)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        assert 'Test Markdown Document' in headings
    finally:
        if os.path.exists(docx_output):
            os.unlink(docx_output)


if __name__ == '__main__':
    pytest.main(['-v', __file__]) 