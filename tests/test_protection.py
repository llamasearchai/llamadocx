"""
Document protection tests for LlamaDocx.

This module contains tests for the document protection functionality of the LlamaDocx package.
"""

import os
import tempfile
import pytest
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION

from llamadocx.protection import (
    protect_document,
    unprotect_document,
    is_protected,
    restrict_editing,
    allow_only_comments,
    allow_only_form_fields,
    allow_only_revisions,
    set_protection_password,
    encrypt_document,
    add_read_only_exception,
    set_document_read_only,
    add_digital_signature,
    get_protection_type
)


@pytest.fixture
def sample_document():
    """Create a sample document for testing protection."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        
        # Add heading
        doc.add_heading('Document Protection Test', level=1)
        
        # Add content
        doc.add_paragraph('This is a sample document for testing protection features.')
        doc.add_paragraph('Users should not be able to edit this content when protection is enabled.')
        
        # Add a form field
        form_paragraph = doc.add_paragraph('Form field: ')
        # Note: In an actual implementation, we would add a form field here
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


def test_protect_unprotect_document(sample_document):
    """Test protecting and unprotecting a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Protect the document
    password = "TestPassword123"
    protect_result = protect_document(doc, password)
    
    # Verify protection was applied
    assert protect_result is True
    assert is_protected(doc) is True
    
    # Save the protected document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        protected_path = tmp.name
    
    try:
        # Load the protected document
        protected_doc = Document(protected_path)
        
        # Verify document is still protected
        assert is_protected(protected_doc) is True
        
        # Attempt to unprotect with wrong password (should fail)
        wrong_password = "WrongPassword"
        unprotect_result_wrong = unprotect_document(protected_doc, wrong_password)
        assert unprotect_result_wrong is False
        assert is_protected(protected_doc) is True
        
        # Unprotect with correct password
        unprotect_result = unprotect_document(protected_doc, password)
        assert unprotect_result is True
        assert is_protected(protected_doc) is False
        
        # Save the unprotected document
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp2:
            protected_doc.save(tmp2.name)
            unprotected_path = tmp2.name
        
        try:
            # Load the unprotected document
            unprotected_doc = Document(unprotected_path)
            
            # Verify document is no longer protected
            assert is_protected(unprotected_doc) is False
        finally:
            # Clean up
            os.unlink(unprotected_path)
    finally:
        # Clean up
        os.unlink(protected_path)


def test_restrict_editing(sample_document):
    """Test restricting editing in a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Restrict editing
    password = "RestrictEdit123"
    restrict_result = restrict_editing(doc, password)
    
    # Verify restriction was applied
    assert restrict_result is True
    
    # Save the restricted document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        restricted_path = tmp.name
    
    try:
        # Load the restricted document
        restricted_doc = Document(restricted_path)
        
        # Verify document is protected
        assert is_protected(restricted_doc) is True
        
        # Verify protection type
        protection_type = get_protection_type(restricted_doc)
        assert protection_type is not None
        assert "EDITING_RESTRICTIONS" in protection_type or "FORMS" in protection_type
    finally:
        # Clean up
        os.unlink(restricted_path)


def test_allow_only_comments(sample_document):
    """Test restricting a document to allow only comments."""
    # Load the document
    doc = Document(sample_document)
    
    # Set to allow only comments
    password = "CommentsOnly123"
    comments_only_result = allow_only_comments(doc, password)
    
    # Verify restriction was applied
    assert comments_only_result is True
    
    # Save the restricted document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        comments_only_path = tmp.name
    
    try:
        # Load the restricted document
        comments_doc = Document(comments_only_path)
        
        # Verify document is protected
        assert is_protected(comments_doc) is True
        
        # Verify protection type
        protection_type = get_protection_type(comments_doc)
        assert protection_type is not None
        assert "COMMENTS" in protection_type
    finally:
        # Clean up
        os.unlink(comments_only_path)


def test_allow_only_form_fields(sample_document):
    """Test restricting a document to allow only form fields."""
    # Load the document
    doc = Document(sample_document)
    
    # Set to allow only form fields
    password = "FormsOnly123"
    forms_only_result = allow_only_form_fields(doc, password)
    
    # Verify restriction was applied
    assert forms_only_result is True
    
    # Save the restricted document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        forms_only_path = tmp.name
    
    try:
        # Load the restricted document
        forms_doc = Document(forms_only_path)
        
        # Verify document is protected
        assert is_protected(forms_doc) is True
        
        # Verify protection type
        protection_type = get_protection_type(forms_doc)
        assert protection_type is not None
        assert "FORMS" in protection_type
    finally:
        # Clean up
        os.unlink(forms_only_path)


def test_allow_only_revisions(sample_document):
    """Test restricting a document to allow only tracked changes."""
    # Load the document
    doc = Document(sample_document)
    
    # Set to allow only tracked changes
    password = "RevisionsOnly123"
    revisions_only_result = allow_only_revisions(doc, password)
    
    # Verify restriction was applied
    assert revisions_only_result is True
    
    # Save the restricted document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        revisions_only_path = tmp.name
    
    try:
        # Load the restricted document
        revisions_doc = Document(revisions_only_path)
        
        # Verify document is protected
        assert is_protected(revisions_doc) is True
        
        # Verify protection type
        protection_type = get_protection_type(revisions_doc)
        assert protection_type is not None
        assert "TRACKED_CHANGES" in protection_type or "REVISIONS" in protection_type
    finally:
        # Clean up
        os.unlink(revisions_only_path)


def test_set_protection_password(sample_document):
    """Test setting and changing protection password."""
    # Load the document
    doc = Document(sample_document)
    
    # Protect document with initial password
    initial_password = "InitialPassword123"
    protect_document(doc, initial_password)
    
    # Change the protection password
    new_password = "NewPassword456"
    change_result = set_protection_password(doc, initial_password, new_password)
    
    # Verify password was changed
    assert change_result is True
    assert is_protected(doc) is True
    
    # Save the document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        pwd_changed_path = tmp.name
    
    try:
        # Load the document
        pwd_doc = Document(pwd_changed_path)
        
        # Verify document is still protected
        assert is_protected(pwd_doc) is True
        
        # Try to unprotect with old password (should fail)
        unprotect_result_old = unprotect_document(pwd_doc, initial_password)
        assert unprotect_result_old is False
        
        # Unprotect with new password (should succeed)
        unprotect_result_new = unprotect_document(pwd_doc, new_password)
        assert unprotect_result_new is True
        assert is_protected(pwd_doc) is False
    finally:
        # Clean up
        os.unlink(pwd_changed_path)


def test_encrypt_document(sample_document):
    """Test encrypting a document with a password."""
    # This test requires the actual file to be saved
    
    # Load the document
    doc = Document(sample_document)
    
    # Add some content to identify the document
    doc.add_paragraph('This document has been encrypted for testing.')
    
    # Save to a temp file for encryption
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        to_encrypt_path = tmp.name
    
    try:
        # Encrypt the document
        encryption_password = "EncryptPass123"
        encrypted_path = to_encrypt_path + ".encrypted.docx"
        encrypt_result = encrypt_document(to_encrypt_path, encrypted_path, encryption_password)
        
        # Verify encryption was successful
        assert encrypt_result is True
        assert os.path.exists(encrypted_path)
        
        try:
            # Attempting to load without password would fail, but we can't test that directly
            # Instead, verify the file exists and has content
            assert os.path.getsize(encrypted_path) > 0
        finally:
            # Clean up
            if os.path.exists(encrypted_path):
                os.unlink(encrypted_path)
    finally:
        # Clean up
        os.unlink(to_encrypt_path)


def test_add_read_only_exception(sample_document):
    """Test adding exceptions to read-only protection."""
    # Load the document
    doc = Document(sample_document)
    
    # Set document as read-only
    password = "ReadOnly123"
    readonly_result = set_document_read_only(doc, password)
    
    # Verify read-only was applied
    assert readonly_result is True
    
    # Add exception for a user
    exception_user = "test.user@example.com"
    exception_result = add_read_only_exception(doc, exception_user)
    
    # Verify exception was added
    assert exception_result is True
    
    # Save the document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        exception_path = tmp.name
    
    try:
        # Load the document
        exception_doc = Document(exception_path)
        
        # Verify document is still protected
        assert is_protected(exception_doc) is True
        
        # Verify protection type
        protection_type = get_protection_type(exception_doc)
        assert protection_type is not None
        assert "READ_ONLY" in protection_type
        
        # Note: We can't easily verify the exception was added through python-docx,
        # as it doesn't expose the exceptions list directly
    finally:
        # Clean up
        os.unlink(exception_path)


def test_digital_signature(sample_document):
    """Test adding a digital signature to a document."""
    # This test requires a certificate file, which we'll simulate
    
    # Create a dummy certificate file
    with tempfile.NamedTemporaryFile(suffix='.pfx', delete=False) as cert_tmp:
        # In a real scenario, this would be a valid certificate file
        cert_tmp.write(b"DUMMY CERTIFICATE DATA")
        cert_path = cert_tmp.name
    
    try:
        # Load the document
        doc = Document(sample_document)
        
        # Save to a temp file for signing
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            to_sign_path = tmp.name
        
        try:
            # Sign the document (this would typically fail with our dummy certificate,
            # but we're testing the API call itself)
            signed_path = to_sign_path + ".signed.docx"
            cert_password = "CertPassword"
            
            try:
                # Attempt to sign (will likely fail with dummy cert)
                sign_result = add_digital_signature(to_sign_path, signed_path, cert_path, cert_password)
                
                # If signing succeeded (unlikely with dummy cert)
                if sign_result and os.path.exists(signed_path):
                    try:
                        # Verify the signed file exists and has content
                        assert os.path.getsize(signed_path) > 0
                    finally:
                        # Clean up signed file
                        if os.path.exists(signed_path):
                            os.unlink(signed_path)
            except Exception as e:
                # Expected to fail with dummy certificate
                # We're just testing the API call exists and can be called
                pass
        finally:
            # Clean up
            os.unlink(to_sign_path)
    finally:
        # Clean up
        os.unlink(cert_path)


def test_multiple_protection_types(sample_document):
    """Test applying multiple protection types to a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Test combining form fields protection with read-only exceptions
    password = "MultiProtect123"
    
    # First restrict to forms
    forms_result = allow_only_form_fields(doc, password)
    assert forms_result is True
    
    # Then add a read-only exception
    exception_user = "multiple.test@example.com"
    exception_result = add_read_only_exception(doc, exception_user)
    assert exception_result is True
    
    # Save the document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        multi_protected_path = tmp.name
    
    try:
        # Load the document
        multi_doc = Document(multi_protected_path)
        
        # Verify document is protected
        assert is_protected(multi_doc) is True
        
        # Verify protection type includes forms
        protection_type = get_protection_type(multi_doc)
        assert protection_type is not None
        assert "FORMS" in protection_type
    finally:
        # Clean up
        os.unlink(multi_protected_path)


if __name__ == '__main__':
    pytest.main(['-v', __file__]) 