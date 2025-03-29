"""
Lists tests for LlamaDocx.

This module contains tests for the list handling functionality of the LlamaDocx package.
"""

import os
import tempfile
import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

from llamadocx.lists import (
    add_bullet_list,
    add_numbered_list,
    add_multilevel_list,
    add_list_item,
    create_list_style,
    apply_list_style,
    set_list_level,
    get_list_styles,
    reset_list_numbering
)


@pytest.fixture
def sample_document():
    """Create a sample document for testing lists."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        
        # Add heading
        doc.add_heading('Lists Test Document', level=1)
        
        # Add content
        doc.add_paragraph('This document is used for testing list functionality.')
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


def test_add_bullet_list(sample_document):
    """Test creating a simple bullet list."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a bullet list
    items = ['First bullet item', 'Second bullet item', 'Third bullet item']
    list_paragraphs = add_bullet_list(doc, items)
    
    # Verify list was created with correct items
    assert len(list_paragraphs) == len(items)
    for i, paragraph in enumerate(list_paragraphs):
        assert items[i] in paragraph.text
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify list paragraphs exist
        # Note: we can't easily check for bullet formatting directly in python-docx
        for i, text in enumerate(items):
            assert text in doc2.paragraphs[i+2].text  # +2 for heading and intro paragraph
    finally:
        # Clean up
        os.unlink(output_path)


def test_add_numbered_list(sample_document):
    """Test creating a numbered list."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a numbered list
    items = ['First numbered item', 'Second numbered item', 'Third numbered item']
    list_paragraphs = add_numbered_list(doc, items)
    
    # Verify list was created with correct items
    assert len(list_paragraphs) == len(items)
    for i, paragraph in enumerate(list_paragraphs):
        assert items[i] in paragraph.text
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify list paragraphs exist
        # Note: we can't easily check for number formatting directly in python-docx
        for i, text in enumerate(items):
            assert text in doc2.paragraphs[i+2].text  # +2 for heading and intro paragraph
    finally:
        # Clean up
        os.unlink(output_path)


def test_add_multilevel_list(sample_document):
    """Test creating a multilevel list."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a multilevel list structure
    items = [
        ('First level item 1', 0),
        ('Second level item 1.1', 1),
        ('Second level item 1.2', 1),
        ('First level item 2', 0),
        ('Second level item 2.1', 1),
        ('Third level item 2.1.1', 2),
        ('First level item 3', 0)
    ]
    
    list_paragraphs = add_multilevel_list(doc, items)
    
    # Verify list was created with correct items and levels
    assert len(list_paragraphs) == len(items)
    for i, paragraph in enumerate(list_paragraphs):
        assert items[i][0] in paragraph.text
        # Note: level is harder to verify directly in python-docx
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify list paragraphs exist
        for i, (text, _) in enumerate(items):
            assert text in doc2.paragraphs[i+2].text  # +2 for heading and intro paragraph
    finally:
        # Clean up
        os.unlink(output_path)


def test_add_list_item(sample_document):
    """Test adding individual list items."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a list with individual items
    first_item = add_list_item(doc, 'First item', level=0, list_type='bullet')
    second_item = add_list_item(doc, 'Second item', level=0, list_type='bullet')
    
    # Add a sub-item
    sub_item = add_list_item(doc, 'Sub-item', level=1, list_type='bullet')
    
    # Continue with main level
    third_item = add_list_item(doc, 'Third item', level=0, list_type='bullet')
    
    # Verify items were created
    assert 'First item' in first_item.text
    assert 'Second item' in second_item.text
    assert 'Sub-item' in sub_item.text
    assert 'Third item' in third_item.text
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify list paragraphs exist
        list_texts = ['First item', 'Second item', 'Sub-item', 'Third item']
        for i, text in enumerate(list_texts):
            assert text in doc2.paragraphs[i+2].text  # +2 for heading and intro paragraph
    finally:
        # Clean up
        os.unlink(output_path)


def test_create_list_style(sample_document):
    """Test creating a custom list style."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a custom list style
    style_name = 'CustomListStyle'
    list_style = create_list_style(doc, style_name, 
                                  bullet_format='•',
                                  number_format='%1.',
                                  font_name='Arial',
                                  font_size=Pt(12))
    
    # Verify style was created
    assert list_style is not None
    
    # Create a list with the custom style
    items = ['First styled item', 'Second styled item', 'Third styled item']
    list_paragraphs = add_bullet_list(doc, items)
    
    # Apply the style to the list
    for paragraph in list_paragraphs:
        apply_list_style(paragraph, style_name)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify list style exists
        list_styles = get_list_styles(doc2)
        style_names = [style.name for style in list_styles if style is not None]
        assert style_name in style_names
        
        # Verify list paragraphs exist with content
        for i, text in enumerate(items):
            assert text in doc2.paragraphs[i+2].text  # +2 for heading and intro paragraph
    finally:
        # Clean up
        os.unlink(output_path)


def test_set_list_level(sample_document):
    """Test changing the level of list items."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a list
    items = ['First item', 'Second item', 'Third item', 'Fourth item']
    list_paragraphs = add_bullet_list(doc, items)
    
    # Change levels for some items
    set_list_level(list_paragraphs[1], 1)  # Second item to level 1
    set_list_level(list_paragraphs[2], 2)  # Third item to level 2
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify list paragraphs exist with content
        for i, text in enumerate(items):
            assert text in doc2.paragraphs[i+2].text  # +2 for heading and intro paragraph
            
        # Unfortunately, we can't easily verify the level directly with python-docx
    finally:
        # Clean up
        os.unlink(output_path)


def test_reset_list_numbering(sample_document):
    """Test resetting list numbering."""
    # Load the document
    doc = Document(sample_document)
    
    # Create first numbered list
    items1 = ['First list item 1', 'First list item 2', 'First list item 3']
    list_paragraphs1 = add_numbered_list(doc, items1)
    
    # Add a paragraph to separate lists
    doc.add_paragraph('This paragraph separates the two lists.')
    
    # Create second numbered list with reset numbering
    items2 = ['Second list item 1', 'Second list item 2', 'Second list item 3']
    list_paragraphs2 = add_numbered_list(doc, items2)
    
    # Reset numbering for the second list
    reset_list_numbering(list_paragraphs2[0])
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify all paragraphs exist
        paragraph_texts = items1 + ['This paragraph separates the two lists.'] + items2
        
        # Check content (numbering reset is difficult to verify directly)
        for i, text in enumerate(paragraph_texts):
            assert text in doc2.paragraphs[i+2].text  # +2 for heading and intro paragraph
    finally:
        # Clean up
        os.unlink(output_path)


def test_mixed_list_types(sample_document):
    """Test creating mixed bullet and numbered lists."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a bullet list
    bullet_items = ['Bullet item 1', 'Bullet item 2', 'Bullet item 3']
    bullet_paragraphs = add_bullet_list(doc, bullet_items)
    
    # Add a separator
    doc.add_paragraph('This separates bullet and numbered lists.')
    
    # Create a numbered list
    numbered_items = ['Numbered item 1', 'Numbered item 2', 'Numbered item 3']
    numbered_paragraphs = add_numbered_list(doc, numbered_items)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify all paragraphs exist
        paragraph_texts = bullet_items + ['This separates bullet and numbered lists.'] + numbered_items
        
        # Check content
        for i, text in enumerate(paragraph_texts):
            assert text in doc2.paragraphs[i+2].text  # +2 for heading and intro paragraph
    finally:
        # Clean up
        os.unlink(output_path)


def test_list_continuation(sample_document):
    """Test continuing a list after a non-list paragraph."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a numbered list
    items1 = ['First item', 'Second item', 'Third item']
    list_paragraphs1 = add_numbered_list(doc, items1)
    
    # Add a paragraph to separate list segments
    doc.add_paragraph('This paragraph interrupts the list.')
    
    # Continue the list with more items
    next_item = add_list_item(doc, 'Fourth item', level=0, list_type='numbered', continue_previous=True)
    another_item = add_list_item(doc, 'Fifth item', level=0, list_type='numbered', continue_previous=True)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify all paragraphs exist
        all_items = items1 + ['This paragraph interrupts the list.', 'Fourth item', 'Fifth item']
        
        # Check content
        for i, text in enumerate(all_items):
            assert text in doc2.paragraphs[i+2].text  # +2 for heading and intro paragraph
    finally:
        # Clean up
        os.unlink(output_path)


if __name__ == '__main__':
    pytest.main(['-v', __file__]) 