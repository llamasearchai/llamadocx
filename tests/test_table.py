"""
Table tests for LlamaDocx.

This module contains tests for the table handling functionality of the LlamaDocx package.
"""

import os
import tempfile
import pytest
from docx import Document

from llamadocx.table import (
    create_table, 
    add_row, 
    add_column, 
    delete_row, 
    delete_column,
    merge_cells,
    set_cell_text,
    apply_table_style,
    set_cell_background,
    set_column_width
)


@pytest.fixture
def sample_document():
    """Create a sample document for testing table operations."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        
        # Add heading
        doc.add_heading('Table Operations Test Document', level=1)
        
        # Add a paragraph
        doc.add_paragraph('This document is used for testing table operations.')
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


def test_create_table(sample_document):
    """Test creating a basic table."""
    # Load the document
    doc = Document(sample_document)
    
    # Initial table count
    initial_tables = len(doc.tables)
    
    # Create a table
    header_row = ['Header 1', 'Header 2', 'Header 3']
    data = [
        ['A1', 'A2', 'A3'],
        ['B1', 'B2', 'B3'],
        ['C1', 'C2', 'C3']
    ]
    
    table = create_table(doc, rows=4, cols=3, header_row=header_row, data=data)
    
    # Verify table was created
    assert len(doc.tables) == initial_tables + 1
    assert table is not None
    assert len(table.rows) == 4
    assert len(table.columns) == 3
    
    # Verify header row
    for i, header_text in enumerate(header_row):
        assert table.cell(0, i).text == header_text
    
    # Verify data rows
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            assert table.cell(row_idx + 1, col_idx).text == cell_data
    
    # Save the document to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify table persisted
        assert len(doc2.tables) == initial_tables + 1
        
        table2 = doc2.tables[-1]
        assert len(table2.rows) == 4
        assert len(table2.columns) == 3
        
        # Check header and data
        for i, header_text in enumerate(header_row):
            assert table2.cell(0, i).text == header_text
        
        for row_idx, row_data in enumerate(data):
            for col_idx, cell_data in enumerate(row_data):
                assert table2.cell(row_idx + 1, col_idx).text == cell_data
    finally:
        # Clean up
        os.unlink(output_path)


def test_add_row(sample_document):
    """Test adding a row to a table."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a table
    table = create_table(doc, rows=3, cols=3)
    initial_row_count = len(table.rows)
    
    # Add a row
    new_row_data = ['New1', 'New2', 'New3']
    add_row(table, new_row_data)
    
    # Verify row was added
    assert len(table.rows) == initial_row_count + 1
    
    # Verify row data
    for i, cell_text in enumerate(new_row_data):
        assert table.cell(initial_row_count, i).text == cell_text
    
    # Add another row at a specific position
    position_row_data = ['Pos1', 'Pos2', 'Pos3']
    add_row(table, position_row_data, position=1)
    
    # Verify row was added at the correct position
    assert len(table.rows) == initial_row_count + 2
    for i, cell_text in enumerate(position_row_data):
        assert table.cell(1, i).text == cell_text


def test_add_column(sample_document):
    """Test adding a column to a table."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a table with some data
    header_row = ['Header 1', 'Header 2']
    data = [
        ['A1', 'A2'],
        ['B1', 'B2'],
        ['C1', 'C2']
    ]
    
    table = create_table(doc, rows=4, cols=2, header_row=header_row, data=data)
    initial_column_count = len(table.columns)
    
    # Add a column
    new_column_data = ['Header 3', 'A3', 'B3', 'C3']
    add_column(table, new_column_data)
    
    # Verify column was added
    assert len(table.columns) == initial_column_count + 1
    
    # Verify column data
    for i, cell_text in enumerate(new_column_data):
        assert table.cell(i, initial_column_count).text == cell_text
    
    # Add another column at a specific position
    position_column_data = ['New Header', 'New A', 'New B', 'New C']
    add_column(table, position_column_data, position=1)
    
    # Verify column was added at the correct position
    assert len(table.columns) == initial_column_count + 2
    for i, cell_text in enumerate(position_column_data):
        assert table.cell(i, 1).text == cell_text


def test_delete_row(sample_document):
    """Test deleting a row from a table."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a table with data
    header_row = ['Header 1', 'Header 2', 'Header 3']
    data = [
        ['A1', 'A2', 'A3'],
        ['B1', 'B2', 'B3'],
        ['C1', 'C2', 'C3']
    ]
    
    table = create_table(doc, rows=4, cols=3, header_row=header_row, data=data)
    initial_row_count = len(table.rows)
    
    # Remember second row data for verification
    second_row_data = [table.cell(2, i).text for i in range(len(table.columns))]
    
    # Delete the second row (index 1, which is row "A")
    delete_row(table, 1)
    
    # Verify row was deleted
    assert len(table.rows) == initial_row_count - 1
    
    # Verify the second row now contains what was previously the third row
    for i, cell_text in enumerate(second_row_data):
        assert table.cell(1, i).text != cell_text
        assert table.cell(1, i).text == data[1][i]  # "B" row data is now at row 1


def test_delete_column(sample_document):
    """Test deleting a column from a table."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a table with data
    header_row = ['Header 1', 'Header 2', 'Header 3']
    data = [
        ['A1', 'A2', 'A3'],
        ['B1', 'B2', 'B3'],
        ['C1', 'C2', 'C3']
    ]
    
    table = create_table(doc, rows=4, cols=3, header_row=header_row, data=data)
    initial_column_count = len(table.columns)
    
    # Remember second column data for verification
    second_column_data = [table.cell(i, 1).text for i in range(len(table.rows))]
    
    # Delete the second column (index 1)
    delete_column(table, 1)
    
    # Verify column was deleted
    assert len(table.columns) == initial_column_count - 1
    
    # Verify the second column now contains what was previously the third column
    for i in range(len(table.rows)):
        if i == 0:
            assert table.cell(i, 1).text != second_column_data[i]
            assert table.cell(i, 1).text == header_row[2]
        else:
            assert table.cell(i, 1).text != second_column_data[i]
            assert table.cell(i, 1).text == data[i-1][2]


def test_merge_cells(sample_document):
    """Test merging cells in a table."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a table
    table = create_table(doc, rows=4, cols=4)
    
    # Set some text to verify later
    set_cell_text(table, 1, 1, "Top-left")
    
    # Merge a 2x2 grid of cells
    merged_cell = merge_cells(table, start_row=1, start_col=1, end_row=2, end_col=2)
    
    # Verify cells were merged
    assert merged_cell is not None
    assert merged_cell.text == "Top-left"
    
    # Verify the dimensions of the merged cell (checking it spans both rows and columns)
    # This is tricky to test directly, so we'll verify indirectly by checking
    # that certain cells are no longer accessible
    
    # Save the document to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        table2 = doc2.tables[0]
        
        # Verify merged cell text persisted
        assert table2.cell(1, 1).text == "Top-left"
        
        # Trying to access merged cells would normally cause issues but python-docx
        # handles this gracefully by returning the same cell object
        assert table2.cell(1, 1) is table2.cell(1, 2)
    finally:
        # Clean up
        os.unlink(output_path)


def test_set_cell_text(sample_document):
    """Test setting text in table cells."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a table
    table = create_table(doc, rows=3, cols=3)
    
    # Set text in various cells
    test_data = [
        (0, 0, "Header 1"),
        (0, 1, "Header 2"),
        (1, 0, "Data 1"),
        (1, 1, "Data 2"),
        (2, 2, "Footer")
    ]
    
    for row, col, text in test_data:
        set_cell_text(table, row, col, text)
    
    # Verify text was set correctly
    for row, col, text in test_data:
        assert table.cell(row, col).text == text


def test_apply_table_style(sample_document):
    """Test applying styles to tables."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a table
    table = create_table(doc, rows=3, cols=3)
    
    # Apply a style
    apply_table_style(table, "Table Grid")
    
    # Verify style was applied
    assert table.style.name == "Table Grid"
    
    # Apply another style
    apply_table_style(table, "Light Shading")
    
    # Verify new style was applied
    assert table.style.name == "Light Shading"


def test_set_cell_background(sample_document):
    """Test setting cell background colors."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a table
    table = create_table(doc, rows=3, cols=3)
    
    # Set background color for specific cells
    set_cell_background(table, 0, 0, "FF0000")  # Red
    set_cell_background(table, 0, 1, "00FF00")  # Green
    set_cell_background(table, 0, 2, "0000FF")  # Blue
    
    # The result is difficult to test programmatically as python-docx doesn't
    # provide a simple way to extract cell shading information.
    # We'll save the document and verify it exists as a basic test
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Just verify file exists and can be loaded
        doc2 = Document(output_path)
        assert len(doc2.tables) == 1
    finally:
        # Clean up
        os.unlink(output_path)


def test_set_column_width(sample_document):
    """Test setting column widths in a table."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a table
    table = create_table(doc, rows=3, cols=3)
    
    # Set column widths
    set_column_width(table, 0, 1000000)  # ~1 inch
    set_column_width(table, 1, 2000000)  # ~2 inches
    set_column_width(table, 2, 1500000)  # ~1.5 inches
    
    # Save the document to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Python-docx doesn't provide a direct way to check column widths
        # after saving, so we just verify the document can be loaded
        doc2 = Document(output_path)
        assert len(doc2.tables) == 1
    finally:
        # Clean up
        os.unlink(output_path)


def test_complex_table_operations(sample_document):
    """Test a combination of table operations in sequence."""
    # Load the document
    doc = Document(sample_document)
    
    # 1. Create a table
    table = create_table(doc, rows=4, cols=3)
    
    # 2. Set header row
    header_row = ['Product', 'Quantity', 'Price']
    for i, text in enumerate(header_row):
        set_cell_text(table, 0, i, text)
    
    # 3. Add data rows
    data = [
        ['Apples', '10', '$2.50'],
        ['Oranges', '5', '$3.00'],
        ['Bananas', '8', '$1.75']
    ]
    
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            set_cell_text(table, row_idx + 1, col_idx, cell_data)
    
    # 4. Apply a style
    apply_table_style(table, "Light Shading Accent 1")
    
    # 5. Add a new column for total
    add_column(table, ['Total', '$25.00', '$15.00', '$14.00'])
    
    # 6. Set column widths
    for i in range(4):
        set_column_width(table, i, 1500000)
    
    # 7. Set header background
    for i in range(4):
        set_cell_background(table, 0, i, "DDDDDD")
    
    # 8. Add a footer row for grand total
    add_row(table, ['Grand Total', '', '', '$54.00'])
    
    # 9. Merge cells for grand total description
    merge_cells(table, start_row=4, start_col=0, end_row=4, end_col=2)
    
    # Save the document to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document and verify key aspects
        doc2 = Document(output_path)
        table2 = doc2.tables[0]
        
        # Verify table dimensions
        assert len(table2.rows) == 5
        assert len(table2.columns) == 4
        
        # Verify headers
        for i, text in enumerate(header_row):
            assert table2.cell(0, i).text == text
        
        # Verify added column
        assert table2.cell(0, 3).text == 'Total'
        
        # Verify data
        for row_idx, row_data in enumerate(data):
            for col_idx, cell_data in enumerate(row_data):
                assert table2.cell(row_idx + 1, col_idx).text == cell_data
        
        # Verify grand total row
        assert table2.cell(4, 0).text == 'Grand Total'
        assert table2.cell(4, 3).text == '$54.00'
        
        # Verify cell merging (indirectly)
        assert table2.cell(4, 0) is table2.cell(4, 1)
    finally:
        # Clean up
        os.unlink(output_path)


if __name__ == '__main__':
    pytest.main(['-v', __file__]) 