"""
File Conversion tests for LlamaDocx.

This module contains tests for the file format conversion functionality of the LlamaDocx package.
"""

import os
import tempfile
from pathlib import Path
import pytest
import shutil

from llamadocx.file_conversion import (
    docx_to_pdf,
    docx_to_html,
    docx_to_txt,
    docx_to_markdown,
    pdf_to_docx,
    html_to_docx,
    txt_to_docx,
    markdown_to_docx,
    is_conversion_supported,
    get_supported_formats
)


@pytest.fixture
def sample_docx():
    """Create a sample DOCX document for testing conversion."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        # We'll create a docx file using Python-docx
        from docx import Document
        
        doc = Document()
        doc.add_heading('File Conversion Test Document', level=1)
        doc.add_paragraph('This is a simple document for testing file format conversion.')
        doc.add_paragraph('It includes multiple paragraphs of text to ensure the content is preserved.')
        doc.add_heading('Second Heading', level=2)
        doc.add_paragraph('This paragraph is under a second-level heading.')
        
        # Add a simple table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = 'Top Left'
        table.cell(0, 1).text = 'Top Right'
        table.cell(1, 0).text = 'Bottom Left'
        table.cell(1, 1).text = 'Bottom Right'
        
        doc.add_paragraph('Text after the table.')
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    if os.path.exists(tmp.name):
        os.unlink(tmp.name)


@pytest.fixture
def sample_txt():
    """Create a sample TXT file for testing conversion."""
    with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp:
        # Write content to the text file
        content = """# File Conversion Test Document

This is a simple document for testing file format conversion.
It includes multiple paragraphs of text to ensure the content is preserved.

## Second Heading

This paragraph is under a second-level heading.

Text after the content.
"""
        tmp.write(content.encode('utf-8'))
        
    yield tmp.name
    
    # Clean up
    if os.path.exists(tmp.name):
        os.unlink(tmp.name)


@pytest.fixture
def sample_html():
    """Create a sample HTML file for testing conversion."""
    with tempfile.NamedTemporaryFile(suffix='.html', delete=False) as tmp:
        # Write content to the HTML file
        content = """<!DOCTYPE html>
<html>
<head>
    <title>File Conversion Test Document</title>
</head>
<body>
    <h1>File Conversion Test Document</h1>
    <p>This is a simple document for testing file format conversion.</p>
    <p>It includes multiple paragraphs of text to ensure the content is preserved.</p>
    <h2>Second Heading</h2>
    <p>This paragraph is under a second-level heading.</p>
    <table>
        <tr>
            <td>Top Left</td>
            <td>Top Right</td>
        </tr>
        <tr>
            <td>Bottom Left</td>
            <td>Bottom Right</td>
        </tr>
    </table>
    <p>Text after the table.</p>
</body>
</html>
"""
        tmp.write(content.encode('utf-8'))
        
    yield tmp.name
    
    # Clean up
    if os.path.exists(tmp.name):
        os.unlink(tmp.name)


@pytest.fixture
def sample_markdown():
    """Create a sample Markdown file for testing conversion."""
    with tempfile.NamedTemporaryFile(suffix='.md', delete=False) as tmp:
        # Write content to the Markdown file
        content = """# File Conversion Test Document

This is a simple document for testing file format conversion.
It includes multiple paragraphs of text to ensure the content is preserved.

## Second Heading

This paragraph is under a second-level heading.

| Top Left | Top Right |
|----------|-----------|
| Bottom Left | Bottom Right |

Text after the table.
"""
        tmp.write(content.encode('utf-8'))
        
    yield tmp.name
    
    # Clean up
    if os.path.exists(tmp.name):
        os.unlink(tmp.name)


def test_docx_to_pdf(sample_docx):
    """Test converting DOCX to PDF."""
    # Skip if PDF conversion is not supported in this environment
    if not is_conversion_supported('docx', 'pdf'):
        pytest.skip("PDF conversion not supported in this environment")
    
    # Convert DOCX to PDF
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
        output_path = tmp.name
    
    try:
        result = docx_to_pdf(sample_docx, output_path)
        
        # Verify conversion was successful
        assert result is True
        assert os.path.exists(output_path)
        
        # Check file size is reasonable (not empty)
        assert os.path.getsize(output_path) > 100
    finally:
        # Clean up
        if os.path.exists(output_path):
            os.unlink(output_path)


def test_docx_to_html(sample_docx):
    """Test converting DOCX to HTML."""
    # Skip if HTML conversion is not supported in this environment
    if not is_conversion_supported('docx', 'html'):
        pytest.skip("HTML conversion not supported in this environment")
    
    # Convert DOCX to HTML
    with tempfile.NamedTemporaryFile(suffix='.html', delete=False) as tmp:
        output_path = tmp.name
    
    try:
        result = docx_to_html(sample_docx, output_path)
        
        # Verify conversion was successful
        assert result is True
        assert os.path.exists(output_path)
        
        # Check file size is reasonable (not empty)
        assert os.path.getsize(output_path) > 100
        
        # Check if the HTML file contains expected content
        with open(output_path, 'r', encoding='utf-8') as f:
            content = f.read()
            assert 'File Conversion Test Document' in content
            assert 'Second Heading' in content
    finally:
        # Clean up
        if os.path.exists(output_path):
            os.unlink(output_path)


def test_docx_to_txt(sample_docx):
    """Test converting DOCX to TXT."""
    # Skip if TXT conversion is not supported in this environment
    if not is_conversion_supported('docx', 'txt'):
        pytest.skip("TXT conversion not supported in this environment")
    
    # Convert DOCX to TXT
    with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp:
        output_path = tmp.name
    
    try:
        result = docx_to_txt(sample_docx, output_path)
        
        # Verify conversion was successful
        assert result is True
        assert os.path.exists(output_path)
        
        # Check file size is reasonable (not empty)
        assert os.path.getsize(output_path) > 10
        
        # Check if the TXT file contains expected content
        with open(output_path, 'r', encoding='utf-8') as f:
            content = f.read()
            assert 'File Conversion Test Document' in content
            assert 'Second Heading' in content
    finally:
        # Clean up
        if os.path.exists(output_path):
            os.unlink(output_path)


def test_docx_to_markdown(sample_docx):
    """Test converting DOCX to Markdown."""
    # Skip if Markdown conversion is not supported in this environment
    if not is_conversion_supported('docx', 'markdown'):
        pytest.skip("Markdown conversion not supported in this environment")
    
    # Convert DOCX to Markdown
    with tempfile.NamedTemporaryFile(suffix='.md', delete=False) as tmp:
        output_path = tmp.name
    
    try:
        result = docx_to_markdown(sample_docx, output_path)
        
        # Verify conversion was successful
        assert result is True
        assert os.path.exists(output_path)
        
        # Check file size is reasonable (not empty)
        assert os.path.getsize(output_path) > 10
        
        # Check if the Markdown file contains expected content
        with open(output_path, 'r', encoding='utf-8') as f:
            content = f.read()
            assert '# File Conversion Test Document' in content
            assert '## Second Heading' in content
    finally:
        # Clean up
        if os.path.exists(output_path):
            os.unlink(output_path)


def test_txt_to_docx(sample_txt):
    """Test converting TXT to DOCX."""
    # Skip if TXT to DOCX conversion is not supported in this environment
    if not is_conversion_supported('txt', 'docx'):
        pytest.skip("TXT to DOCX conversion not supported in this environment")
    
    # Convert TXT to DOCX
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
    
    try:
        result = txt_to_docx(sample_txt, output_path)
        
        # Verify conversion was successful
        assert result is True
        assert os.path.exists(output_path)
        
        # Check file size is reasonable (not empty)
        assert os.path.getsize(output_path) > 1000
        
        # Load the document and check content
        from docx import Document
        doc = Document(output_path)
        
        # Extract text from the document
        text = '\n'.join([para.text for para in doc.paragraphs])
        
        # Verify content was preserved
        assert 'File Conversion Test Document' in text
        assert 'Second Heading' in text
    finally:
        # Clean up
        if os.path.exists(output_path):
            os.unlink(output_path)


def test_html_to_docx(sample_html):
    """Test converting HTML to DOCX."""
    # Skip if HTML to DOCX conversion is not supported in this environment
    if not is_conversion_supported('html', 'docx'):
        pytest.skip("HTML to DOCX conversion not supported in this environment")
    
    # Convert HTML to DOCX
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
    
    try:
        result = html_to_docx(sample_html, output_path)
        
        # Verify conversion was successful
        assert result is True
        assert os.path.exists(output_path)
        
        # Check file size is reasonable (not empty)
        assert os.path.getsize(output_path) > 1000
        
        # Load the document and check content
        from docx import Document
        doc = Document(output_path)
        
        # Extract text from the document
        text = '\n'.join([para.text for para in doc.paragraphs])
        
        # Verify content was preserved
        assert 'File Conversion Test Document' in text
        assert 'Second Heading' in text
    finally:
        # Clean up
        if os.path.exists(output_path):
            os.unlink(output_path)


def test_markdown_to_docx(sample_markdown):
    """Test converting Markdown to DOCX."""
    # Skip if Markdown to DOCX conversion is not supported in this environment
    if not is_conversion_supported('markdown', 'docx'):
        pytest.skip("Markdown to DOCX conversion not supported in this environment")
    
    # Convert Markdown to DOCX
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
    
    try:
        result = markdown_to_docx(sample_markdown, output_path)
        
        # Verify conversion was successful
        assert result is True
        assert os.path.exists(output_path)
        
        # Check file size is reasonable (not empty)
        assert os.path.getsize(output_path) > 1000
        
        # Load the document and check content
        from docx import Document
        doc = Document(output_path)
        
        # Extract text from the document
        text = '\n'.join([para.text for para in doc.paragraphs])
        
        # Verify content was preserved
        assert 'File Conversion Test Document' in text
        assert 'Second Heading' in text
    finally:
        # Clean up
        if os.path.exists(output_path):
            os.unlink(output_path)


def test_pdf_to_docx(sample_docx):
    """Test converting PDF to DOCX."""
    # First convert DOCX to PDF, then PDF to DOCX
    # Skip if either conversion is not supported in this environment
    if not is_conversion_supported('docx', 'pdf') or not is_conversion_supported('pdf', 'docx'):
        pytest.skip("PDF conversion not supported in this environment")
    
    # Create temporary files for intermediate PDF and final DOCX
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as pdf_tmp:
        pdf_path = pdf_tmp.name
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as docx_tmp:
        final_docx_path = docx_tmp.name
    
    try:
        # First convert DOCX to PDF
        docx_to_pdf_result = docx_to_pdf(sample_docx, pdf_path)
        assert docx_to_pdf_result is True
        assert os.path.exists(pdf_path)
        
        # Then convert PDF back to DOCX
        pdf_to_docx_result = pdf_to_docx(pdf_path, final_docx_path)
        
        # Verify conversion was successful
        assert pdf_to_docx_result is True
        assert os.path.exists(final_docx_path)
        
        # Check file size is reasonable (not empty)
        assert os.path.getsize(final_docx_path) > 1000
        
        # Note: Content verification is challenging for PDF to DOCX conversions
        # as the content structure may change significantly. We just verify 
        # the conversion completed and produced a non-empty file.
    finally:
        # Clean up
        if os.path.exists(pdf_path):
            os.unlink(pdf_path)
        if os.path.exists(final_docx_path):
            os.unlink(final_docx_path)


def test_is_conversion_supported():
    """Test checking if a conversion is supported."""
    # These conversions should be supported in most environments
    assert is_conversion_supported('docx', 'txt') is True
    
    # Test an unsupported conversion
    assert is_conversion_supported('docx', 'unsupported_format') is False


def test_get_supported_formats():
    """Test getting the list of supported formats."""
    supported_formats = get_supported_formats()
    
    # Verify the result is a list or dictionary
    assert isinstance(supported_formats, (list, dict))
    
    # Verify common formats are included
    if isinstance(supported_formats, list):
        assert 'docx' in supported_formats
        assert 'txt' in supported_formats
    else:  # If it's a dictionary
        assert 'input' in supported_formats
        assert 'output' in supported_formats


def test_conversion_with_options(sample_docx):
    """Test conversion with additional options (if supported)."""
    # Skip if HTML conversion is not supported in this environment
    if not is_conversion_supported('docx', 'html'):
        pytest.skip("HTML conversion not supported in this environment")
    
    # Convert DOCX to HTML with additional options
    with tempfile.NamedTemporaryFile(suffix='.html', delete=False) as tmp:
        output_path = tmp.name
    
    try:
        # Example options: include images, use specific CSS, etc.
        options = {
            'include_images': True,
            'include_css': True
        }
        
        result = docx_to_html(sample_docx, output_path, **options)
        
        # Verify conversion was successful
        assert result is True
        assert os.path.exists(output_path)
        
        # Check file size is reasonable (not empty)
        assert os.path.getsize(output_path) > 100
    finally:
        # Clean up
        if os.path.exists(output_path):
            os.unlink(output_path)


def test_batch_conversion(sample_docx):
    """Test batch conversion of multiple files (if supported)."""
    # Skip if TXT conversion is not supported
    if not is_conversion_supported('docx', 'txt'):
        pytest.skip("TXT conversion not supported in this environment")
    
    # Create multiple copies of sample_docx
    temp_dir = tempfile.mkdtemp()
    try:
        # Create 3 copies of the sample file
        input_files = []
        for i in range(3):
            dst = os.path.join(temp_dir, f"input{i}.docx")
            shutil.copy(sample_docx, dst)
            input_files.append(dst)
        
        # Create output directory
        output_dir = os.path.join(temp_dir, "output")
        os.mkdir(output_dir)
        
        # Convert each file
        results = []
        for input_file in input_files:
            base_name = os.path.basename(input_file)
            name_without_ext = os.path.splitext(base_name)[0]
            output_path = os.path.join(output_dir, f"{name_without_ext}.txt")
            
            result = docx_to_txt(input_file, output_path)
            results.append(result)
        
        # Verify all conversions were successful
        assert all(results)
        
        # Verify all output files exist
        for i in range(3):
            output_path = os.path.join(output_dir, f"input{i}.txt")
            assert os.path.exists(output_path)
            assert os.path.getsize(output_path) > 10
    finally:
        # Clean up
        shutil.rmtree(temp_dir)


if __name__ == '__main__':
    pytest.main(['-v', __file__]) 