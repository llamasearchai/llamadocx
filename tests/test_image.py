"""
Image tests for LlamaDocx.

This module contains tests for the image handling functionality of the LlamaDocx package.
"""

import os
import tempfile
from pathlib import Path
import pytest
from docx import Document

from llamadocx.image import Image, ImageProcessor


# Create a very simple test image file
def create_test_image(path, size=(100, 100)):
    """Create a simple test image at the specified path."""
    try:
        from PIL import Image as PILImage
        
        # Create a simple RGB image
        img = PILImage.new('RGB', size, color=(73, 109, 137))
        img.save(path)
        return True
    except ImportError:
        # PIL not available, create an empty file
        with open(path, 'wb') as f:
            f.write(b'')
        return False


@pytest.fixture
def temp_image():
    """Create a temporary image file for testing."""
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
        has_pil = create_test_image(tmp.name)
        
    yield tmp.name, has_pil
    
    # Clean up
    os.unlink(tmp.name)


@pytest.fixture
def temp_docx_with_image(temp_image):
    """Create a temporary DOCX file with an embedded image for testing."""
    image_path, has_pil = temp_image
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        doc.add_heading('Document with Image', level=1)
        doc.add_paragraph('This is a test document with an embedded image.')
        
        # Add the image to the document
        if os.path.exists(image_path):
            doc.add_picture(image_path, width=2000000, height=2000000)  # ~2 inches
            
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name, has_pil
    
    # Clean up
    os.unlink(tmp.name)


def test_image_properties(temp_docx_with_image):
    """Test basic Image class properties."""
    docx_path, has_pil = temp_docx_with_image
    
    # Skip test if we couldn't create a proper image
    if not has_pil:
        pytest.skip("PIL not available for image testing")
    
    # Load the document and get the image
    doc = Document(docx_path)
    
    # Make sure we have an image
    assert len(doc.inline_shapes) > 0
    
    # Create an Image wrapper for the first inline shape
    image = Image(doc.inline_shapes[0])
    
    # Test basic properties
    assert image.width > 0
    assert image.height > 0
    assert image.aspect_ratio > 0


def test_image_resize(temp_docx_with_image):
    """Test Image resize functionality."""
    docx_path, has_pil = temp_docx_with_image
    
    # Skip test if we couldn't create a proper image
    if not has_pil:
        pytest.skip("PIL not available for image testing")
    
    # Load the document and get the image
    doc = Document(docx_path)
    
    # Make sure we have an image
    assert len(doc.inline_shapes) > 0
    
    # Create an Image wrapper for the first inline shape
    image = Image(doc.inline_shapes[0])
    
    # Save original dimensions
    original_width = image.width
    original_height = image.height
    original_ratio = image.aspect_ratio
    
    # Test resize_to_width
    new_width = original_width // 2
    image.resize_to_width(new_width)
    
    assert image.width == new_width
    assert round(image.aspect_ratio, 5) == round(original_ratio, 5)  # Aspect ratio should be preserved
    
    # Test resize_to_height
    new_height = original_height // 2
    image.resize_to_height(new_height)
    
    assert image.height == new_height
    assert round(image.aspect_ratio, 5) == round(original_ratio, 5)  # Aspect ratio should be preserved
    
    # Test resize
    image.resize(original_width, original_height)
    
    assert image.width == original_width
    assert image.height == original_height


@pytest.mark.skipif(not ImageProcessor.has_pil(), reason="PIL not available")
def test_image_processor_basic(temp_image):
    """Test basic ImageProcessor functionality."""
    image_path, _ = temp_image
    
    # Create a processor
    processor = ImageProcessor(image_path)
    
    # Test that the processor loaded the image
    assert processor.image is not None
    
    # Test resizing
    resized = processor.resize(50, 50)
    assert resized.width == 50
    assert resized.height == 50
    
    # Test rotating
    rotated = processor.rotate(90)
    assert rotated.width == resized.height
    assert rotated.height == resized.width
    
    # Test saving
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
        output_path = tmp.name
    
    try:
        processor.save(output_path)
        assert os.path.exists(output_path)
        assert os.path.getsize(output_path) > 0
    finally:
        if os.path.exists(output_path):
            os.unlink(output_path)


@pytest.mark.skipif(not ImageProcessor.has_pil(), reason="PIL not available")
def test_image_processor_methods(temp_image):
    """Test various ImageProcessor methods."""
    image_path, _ = temp_image
    
    # Create a processor
    processor = ImageProcessor(image_path)
    
    # Test crop
    cropped = processor.crop(10, 10, 90, 90)
    assert cropped.width == 80
    assert cropped.height == 80
    
    # Test grayscale conversion
    gray = processor.convert_to_grayscale()
    # We can't easily test the color conversion itself, but we can ensure it returns an image
    assert gray is not None
    
    # Test brightness adjustment
    brighter = processor.adjust_brightness(1.5)
    darker = processor.adjust_brightness(0.5)
    assert brighter is not None
    assert darker is not None
    
    # Test contrast adjustment
    higher_contrast = processor.adjust_contrast(1.5)
    lower_contrast = processor.adjust_contrast(0.5)
    assert higher_contrast is not None
    assert lower_contrast is not None


def test_image_in_document(temp_docx_with_image, temp_image):
    """Test adding and manipulating images in a document."""
    docx_path, has_pil = temp_docx_with_image
    image_path, _ = temp_image
    
    # Skip test if we couldn't create a proper image
    if not has_pil or not ImageProcessor.has_pil():
        pytest.skip("PIL not available for image testing")
    
    # Load the document
    doc = Document(docx_path)
    
    # Count initial images
    initial_image_count = len(doc.inline_shapes)
    
    # Add another image
    doc.add_picture(image_path)
    
    # Verify image was added
    assert len(doc.inline_shapes) == initial_image_count + 1
    
    # Create Image wrapper for the new image
    image = Image(doc.inline_shapes[-1])
    
    # Test setting width and height directly
    new_width = 3000000  # ~3 inches
    new_height = 2000000  # ~2 inches
    
    image.width = new_width
    image.height = new_height
    
    assert image.width == new_width
    assert image.height == new_height
    
    # Save the document with the modified image
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
    
    try:
        doc.save(output_path)
        
        # Reload the document and check image dimensions
        doc2 = Document(output_path)
        assert len(doc2.inline_shapes) == initial_image_count + 1
        
        image2 = Image(doc2.inline_shapes[-1])
        assert image2.width == new_width
        assert image2.height == new_height
    finally:
        if os.path.exists(output_path):
            os.unlink(output_path)


if __name__ == '__main__':
    pytest.main(['-v', __file__]) 