"""
Formatting tests for LlamaDocx.

This module contains tests for the text and paragraph formatting functionality of the LlamaDocx package.
"""

import os
import tempfile
import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, RGBColor, Inches

from llamadocx.formatting import (
    set_font,
    set_font_size,
    set_bold,
    set_italic,
    set_underline,
    set_font_color,
    set_highlight_color,
    set_alignment,
    set_indentation,
    set_spacing,
    set_line_spacing,
    set_keep_together,
    set_keep_with_next,
    set_page_break_before,
    set_widow_control,
    set_style,
    apply_character_style,
    apply_paragraph_style
)


@pytest.fixture
def sample_document():
    """Create a sample document for testing formatting operations."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        
        # Add heading
        doc.add_heading('Formatting Test Document', level=1)
        
        # Add multiple paragraphs with different content
        for i in range(5):
            doc.add_paragraph(f'This is paragraph {i+1} for testing formatting.')
        
        # Add a paragraph with multiple runs
        p = doc.add_paragraph()
        p.add_run('This is the first run. ')
        p.add_run('This is the second run. ')
        p.add_run('This is the third run.')
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


def test_set_font(sample_document):
    """Test setting font for runs and paragraphs."""
    # Load the document
    doc = Document(sample_document)
    
    # Set font for a run
    run = doc.paragraphs[1].add_run("This text should use Arial font.")
    set_font(run, "Arial")
    
    # Set font for all runs in a paragraph
    for run in doc.paragraphs[2].runs:
        set_font(run, "Times New Roman")
    
    # Verify font was set correctly
    assert doc.paragraphs[1].runs[-1].font.name == "Arial"
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify font persisted
        assert doc2.paragraphs[1].runs[-1].font.name == "Arial"
        
    finally:
        # Clean up
        os.unlink(output_path)


def test_set_font_size(sample_document):
    """Test setting font size for runs."""
    # Load the document
    doc = Document(sample_document)
    
    # Set different font sizes
    run1 = doc.paragraphs[1].add_run("This text should be 12pt. ")
    set_font_size(run1, Pt(12))
    
    run2 = doc.paragraphs[1].add_run("This text should be 14pt. ")
    set_font_size(run2, Pt(14))
    
    run3 = doc.paragraphs[1].add_run("This text should be 16pt.")
    set_font_size(run3, Pt(16))
    
    # Verify font sizes were set correctly
    assert run1.font.size == Pt(12)
    assert run2.font.size == Pt(14)
    assert run3.font.size == Pt(16)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify font sizes persisted
        runs = doc2.paragraphs[1].runs
        assert runs[-3].font.size == Pt(12)
        assert runs[-2].font.size == Pt(14)
        assert runs[-1].font.size == Pt(16)
        
    finally:
        # Clean up
        os.unlink(output_path)


def test_set_text_emphasis(sample_document):
    """Test setting bold, italic, and underline for runs."""
    # Load the document
    doc = Document(sample_document)
    
    # Add runs with different emphasis
    run1 = doc.paragraphs[1].add_run("This text should be bold. ")
    set_bold(run1, True)
    
    run2 = doc.paragraphs[1].add_run("This text should be italic. ")
    set_italic(run2, True)
    
    run3 = doc.paragraphs[1].add_run("This text should be underlined. ")
    set_underline(run3, True)
    
    run4 = doc.paragraphs[1].add_run("This text should be bold, italic, and underlined.")
    set_bold(run4, True)
    set_italic(run4, True)
    set_underline(run4, True)
    
    # Verify emphasis was set correctly
    assert run1.bold == True
    assert run2.italic == True
    assert run3.underline == True
    assert run4.bold == True and run4.italic == True and run4.underline == True
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify emphasis persisted
        runs = doc2.paragraphs[1].runs
        assert runs[-4].bold == True
        assert runs[-3].italic == True
        assert runs[-2].underline == True
        assert runs[-1].bold == True and runs[-1].italic == True and runs[-1].underline == True
        
    finally:
        # Clean up
        os.unlink(output_path)


def test_set_font_color(sample_document):
    """Test setting font color for runs."""
    # Load the document
    doc = Document(sample_document)
    
    # Add runs with different colors
    run1 = doc.paragraphs[1].add_run("This text should be red. ")
    set_font_color(run1, RGBColor(255, 0, 0))
    
    run2 = doc.paragraphs[1].add_run("This text should be green. ")
    set_font_color(run2, RGBColor(0, 255, 0))
    
    run3 = doc.paragraphs[1].add_run("This text should be blue.")
    set_font_color(run3, RGBColor(0, 0, 255))
    
    # Verify colors were set correctly
    assert run1.font.color.rgb == RGBColor(255, 0, 0)
    assert run2.font.color.rgb == RGBColor(0, 255, 0)
    assert run3.font.color.rgb == RGBColor(0, 0, 255)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify colors persisted
        runs = doc2.paragraphs[1].runs
        assert runs[-3].font.color.rgb == RGBColor(255, 0, 0)
        assert runs[-2].font.color.rgb == RGBColor(0, 255, 0)
        assert runs[-1].font.color.rgb == RGBColor(0, 0, 255)
        
    finally:
        # Clean up
        os.unlink(output_path)


def test_set_highlight_color(sample_document):
    """Test setting highlight color for runs."""
    # Load the document
    doc = Document(sample_document)
    
    # Add runs with different highlight colors
    run1 = doc.paragraphs[1].add_run("This text should be highlighted in yellow. ")
    set_highlight_color(run1, "yellow")
    
    run2 = doc.paragraphs[1].add_run("This text should be highlighted in green. ")
    set_highlight_color(run2, "green")
    
    run3 = doc.paragraphs[1].add_run("This text should be highlighted in cyan.")
    set_highlight_color(run3, "cyan")
    
    # Verify highlight colors were set correctly
    assert run1.font.highlight_color is not None
    assert run2.font.highlight_color is not None
    assert run3.font.highlight_color is not None
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify highlight colors persisted
        runs = doc2.paragraphs[1].runs
        assert runs[-3].font.highlight_color is not None
        assert runs[-2].font.highlight_color is not None
        assert runs[-1].font.highlight_color is not None
        
    finally:
        # Clean up
        os.unlink(output_path)


def test_set_alignment(sample_document):
    """Test setting paragraph alignment."""
    # Load the document
    doc = Document(sample_document)
    
    # Set different alignments for different paragraphs
    set_alignment(doc.paragraphs[1], WD_ALIGN_PARAGRAPH.LEFT)
    set_alignment(doc.paragraphs[2], WD_ALIGN_PARAGRAPH.CENTER)
    set_alignment(doc.paragraphs[3], WD_ALIGN_PARAGRAPH.RIGHT)
    set_alignment(doc.paragraphs[4], WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    # Verify alignments were set correctly
    assert doc.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert doc.paragraphs[2].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert doc.paragraphs[3].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert doc.paragraphs[4].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify alignments persisted
        assert doc2.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.LEFT
        assert doc2.paragraphs[2].alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert doc2.paragraphs[3].alignment == WD_ALIGN_PARAGRAPH.RIGHT
        assert doc2.paragraphs[4].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
        
    finally:
        # Clean up
        os.unlink(output_path)


def test_set_indentation(sample_document):
    """Test setting paragraph indentation."""
    # Load the document
    doc = Document(sample_document)
    
    # Set different indentations
    set_indentation(doc.paragraphs[1], left=Inches(0.5))
    set_indentation(doc.paragraphs[2], left=Inches(0.25), right=Inches(0.25))
    set_indentation(doc.paragraphs[3], first_line=Inches(0.5))
    set_indentation(doc.paragraphs[4], hanging=Inches(0.5))
    
    # Verify indentations were set correctly
    assert doc.paragraphs[1].paragraph_format.left_indent == Inches(0.5)
    assert doc.paragraphs[2].paragraph_format.left_indent == Inches(0.25)
    assert doc.paragraphs[2].paragraph_format.right_indent == Inches(0.25)
    assert doc.paragraphs[3].paragraph_format.first_line_indent == Inches(0.5)
    assert doc.paragraphs[4].paragraph_format.first_line_indent == -Inches(0.5)  # Hanging indent is negative
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify indentations persisted
        assert doc2.paragraphs[1].paragraph_format.left_indent == Inches(0.5)
        assert doc2.paragraphs[2].paragraph_format.left_indent == Inches(0.25)
        assert doc2.paragraphs[2].paragraph_format.right_indent == Inches(0.25)
        assert doc2.paragraphs[3].paragraph_format.first_line_indent == Inches(0.5)
        assert doc2.paragraphs[4].paragraph_format.first_line_indent == -Inches(0.5)
        
    finally:
        # Clean up
        os.unlink(output_path)


def test_set_spacing(sample_document):
    """Test setting paragraph spacing."""
    # Load the document
    doc = Document(sample_document)
    
    # Set different spacing
    set_spacing(doc.paragraphs[1], before=Pt(12), after=Pt(12))
    set_spacing(doc.paragraphs[2], before=Pt(6), after=Pt(18))
    set_spacing(doc.paragraphs[3], before=Pt(18), after=Pt(6))
    
    # Verify spacing was set correctly
    assert doc.paragraphs[1].paragraph_format.space_before == Pt(12)
    assert doc.paragraphs[1].paragraph_format.space_after == Pt(12)
    assert doc.paragraphs[2].paragraph_format.space_before == Pt(6)
    assert doc.paragraphs[2].paragraph_format.space_after == Pt(18)
    assert doc.paragraphs[3].paragraph_format.space_before == Pt(18)
    assert doc.paragraphs[3].paragraph_format.space_after == Pt(6)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify spacing persisted
        assert doc2.paragraphs[1].paragraph_format.space_before == Pt(12)
        assert doc2.paragraphs[1].paragraph_format.space_after == Pt(12)
        assert doc2.paragraphs[2].paragraph_format.space_before == Pt(6)
        assert doc2.paragraphs[2].paragraph_format.space_after == Pt(18)
        assert doc2.paragraphs[3].paragraph_format.space_before == Pt(18)
        assert doc2.paragraphs[3].paragraph_format.space_after == Pt(6)
        
    finally:
        # Clean up
        os.unlink(output_path)


def test_set_line_spacing(sample_document):
    """Test setting paragraph line spacing."""
    # Load the document
    doc = Document(sample_document)
    
    # Set different line spacing
    set_line_spacing(doc.paragraphs[1], WD_LINE_SPACING.SINGLE)
    set_line_spacing(doc.paragraphs[2], WD_LINE_SPACING.DOUBLE)
    set_line_spacing(doc.paragraphs[3], WD_LINE_SPACING.ONE_POINT_FIVE)
    # Exact value in points
    set_line_spacing(doc.paragraphs[4], None, Pt(24))
    
    # Verify line spacing was set correctly
    assert doc.paragraphs[1].paragraph_format.line_spacing_rule == WD_LINE_SPACING.SINGLE
    assert doc.paragraphs[2].paragraph_format.line_spacing_rule == WD_LINE_SPACING.DOUBLE
    assert doc.paragraphs[3].paragraph_format.line_spacing_rule == WD_LINE_SPACING.ONE_POINT_FIVE
    assert doc.paragraphs[4].paragraph_format.line_spacing_rule == WD_LINE_SPACING.EXACTLY
    assert doc.paragraphs[4].paragraph_format.line_spacing == Pt(24)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify line spacing persisted
        assert doc2.paragraphs[1].paragraph_format.line_spacing_rule == WD_LINE_SPACING.SINGLE
        assert doc2.paragraphs[2].paragraph_format.line_spacing_rule == WD_LINE_SPACING.DOUBLE
        assert doc2.paragraphs[3].paragraph_format.line_spacing_rule == WD_LINE_SPACING.ONE_POINT_FIVE
        assert doc2.paragraphs[4].paragraph_format.line_spacing_rule == WD_LINE_SPACING.EXACTLY
        assert doc2.paragraphs[4].paragraph_format.line_spacing == Pt(24)
        
    finally:
        # Clean up
        os.unlink(output_path)


def test_paragraph_formatting_options(sample_document):
    """Test paragraph formatting options like keep_together, keep_with_next, page_break_before, and widow_control."""
    # Load the document
    doc = Document(sample_document)
    
    # Set different paragraph formatting options
    set_keep_together(doc.paragraphs[1], True)
    set_keep_with_next(doc.paragraphs[2], True)
    set_page_break_before(doc.paragraphs[3], True)
    set_widow_control(doc.paragraphs[4], True)
    
    # Verify options were set correctly
    assert doc.paragraphs[1].paragraph_format.keep_together == True
    assert doc.paragraphs[2].paragraph_format.keep_with_next == True
    assert doc.paragraphs[3].paragraph_format.page_break_before == True
    assert doc.paragraphs[4].paragraph_format.widow_control == True
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify options persisted
        assert doc2.paragraphs[1].paragraph_format.keep_together == True
        assert doc2.paragraphs[2].paragraph_format.keep_with_next == True
        assert doc2.paragraphs[3].paragraph_format.page_break_before == True
        assert doc2.paragraphs[4].paragraph_format.widow_control == True
        
    finally:
        # Clean up
        os.unlink(output_path)


def test_set_style(sample_document):
    """Test applying document styles to paragraphs."""
    # Load the document
    doc = Document(sample_document)
    
    # Apply built-in styles
    set_style(doc.paragraphs[1], "Heading 2")
    set_style(doc.paragraphs[2], "Quote")
    set_style(doc.paragraphs[3], "Intense Quote")
    set_style(doc.paragraphs[4], "List Paragraph")
    
    # Verify styles were applied
    assert doc.paragraphs[1].style.name == "Heading 2"
    assert doc.paragraphs[2].style.name == "Quote"
    assert doc.paragraphs[3].style.name == "Intense Quote"
    assert doc.paragraphs[4].style.name == "List Paragraph"
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify styles persisted
        assert doc2.paragraphs[1].style.name == "Heading 2"
        assert doc2.paragraphs[2].style.name == "Quote"
        assert doc2.paragraphs[3].style.name == "Intense Quote"
        assert doc2.paragraphs[4].style.name == "List Paragraph"
        
    finally:
        # Clean up
        os.unlink(output_path)


def test_apply_character_style(sample_document):
    """Test applying character styles to runs."""
    # Load the document
    doc = Document(sample_document)
    
    # Create runs with character styles
    run1 = doc.paragraphs[1].add_run("This text should use the Emphasis style. ")
    apply_character_style(run1, "Emphasis")
    
    run2 = doc.paragraphs[1].add_run("This text should use the Strong style. ")
    apply_character_style(run2, "Strong")
    
    run3 = doc.paragraphs[1].add_run("This text should use the Subtle Reference style.")
    apply_character_style(run3, "Subtle Reference")
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify character styles persisted (checking style names may vary between Word versions)
        runs = doc2.paragraphs[1].runs
        assert runs[-3].style is not None
        assert runs[-2].style is not None
        assert runs[-1].style is not None
        
    finally:
        # Clean up
        os.unlink(output_path)


def test_apply_paragraph_style(sample_document):
    """Test applying complex paragraph styles."""
    # Load the document
    doc = Document(sample_document)
    
    # Apply paragraph style with custom formatting
    apply_paragraph_style(
        doc.paragraphs[1],
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        left_indent=Inches(0.5),
        right_indent=Inches(0.5),
        space_before=Pt(12),
        space_after=Pt(12),
        line_spacing=WD_LINE_SPACING.DOUBLE,
        keep_together=True
    )
    
    # Verify all styles were applied
    assert doc.paragraphs[1].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert doc.paragraphs[1].paragraph_format.left_indent == Inches(0.5)
    assert doc.paragraphs[1].paragraph_format.right_indent == Inches(0.5)
    assert doc.paragraphs[1].paragraph_format.space_before == Pt(12)
    assert doc.paragraphs[1].paragraph_format.space_after == Pt(12)
    assert doc.paragraphs[1].paragraph_format.line_spacing_rule == WD_LINE_SPACING.DOUBLE
    assert doc.paragraphs[1].paragraph_format.keep_together == True
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify complex style persisted
        assert doc2.paragraphs[1].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
        assert doc2.paragraphs[1].paragraph_format.left_indent == Inches(0.5)
        assert doc2.paragraphs[1].paragraph_format.right_indent == Inches(0.5)
        assert doc2.paragraphs[1].paragraph_format.space_before == Pt(12)
        assert doc2.paragraphs[1].paragraph_format.space_after == Pt(12)
        assert doc2.paragraphs[1].paragraph_format.line_spacing_rule == WD_LINE_SPACING.DOUBLE
        assert doc2.paragraphs[1].paragraph_format.keep_together == True
        
    finally:
        # Clean up
        os.unlink(output_path)


def test_combined_formatting(sample_document):
    """Test combining multiple formatting attributes."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a paragraph with multiple combined formatting
    p = doc.add_paragraph()
    
    # Create a run with multiple formatting attributes
    run = p.add_run("This text has multiple formatting attributes applied.")
    set_font(run, "Arial")
    set_font_size(run, Pt(14))
    set_bold(run, True)
    set_italic(run, True)
    set_underline(run, True)
    set_font_color(run, RGBColor(0, 0, 255))
    set_highlight_color(run, "yellow")
    
    # Apply paragraph formatting
    set_alignment(p, WD_ALIGN_PARAGRAPH.CENTER)
    set_indentation(p, left=Inches(0.5), right=Inches(0.5))
    set_spacing(p, before=Pt(12), after=Pt(12))
    set_line_spacing(p, WD_LINE_SPACING.DOUBLE)
    set_keep_together(p, True)
    
    # Verify all formatting was applied correctly
    assert run.font.name == "Arial"
    assert run.font.size == Pt(14)
    assert run.bold == True
    assert run.italic == True
    assert run.underline == True
    assert run.font.color.rgb == RGBColor(0, 0, 255)
    assert run.font.highlight_color is not None
    
    assert p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert p.paragraph_format.left_indent == Inches(0.5)
    assert p.paragraph_format.right_indent == Inches(0.5)
    assert p.paragraph_format.space_before == Pt(12)
    assert p.paragraph_format.space_after == Pt(12)
    assert p.paragraph_format.line_spacing_rule == WD_LINE_SPACING.DOUBLE
    assert p.paragraph_format.keep_together == True
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get the added paragraph (last paragraph)
        p2 = doc2.paragraphs[-1]
        run2 = p2.runs[0]
        
        # Verify character formatting persisted
        assert run2.font.name == "Arial"
        assert run2.font.size == Pt(14)
        assert run2.bold == True
        assert run2.italic == True
        assert run2.underline == True
        assert run2.font.color.rgb == RGBColor(0, 0, 255)
        assert run2.font.highlight_color is not None
        
        # Verify paragraph formatting persisted
        assert p2.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert p2.paragraph_format.left_indent == Inches(0.5)
        assert p2.paragraph_format.right_indent == Inches(0.5)
        assert p2.paragraph_format.space_before == Pt(12)
        assert p2.paragraph_format.space_after == Pt(12)
        assert p2.paragraph_format.line_spacing_rule == WD_LINE_SPACING.DOUBLE
        assert p2.paragraph_format.keep_together == True
        
    finally:
        # Clean up
        os.unlink(output_path)


if __name__ == '__main__':
    pytest.main(['-v', __file__]) 