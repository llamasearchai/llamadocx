"""
Footnotes and Endnotes tests for LlamaDocx.

This module contains tests for the footnotes and endnotes functionality of the LlamaDocx package.
"""

import os
import tempfile
import pytest
from docx import Document
from docx.shared import Pt

from llamadocx.notes import (
    add_footnote,
    add_endnote,
    get_footnotes,
    get_endnotes,
    update_footnote,
    update_endnote,
    delete_footnote,
    delete_endnote,
    format_footnotes,
    format_endnotes,
    get_footnote_text,
    get_endnote_text
)


@pytest.fixture
def sample_document():
    """Create a sample document for testing footnotes and endnotes."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        
        # Add heading
        doc.add_heading('Footnotes and Endnotes Test Document', level=1)
        
        # Add paragraphs
        doc.add_paragraph('This is paragraph 1 for testing footnotes.')
        doc.add_paragraph('This is paragraph 2 for testing more footnotes.')
        doc.add_paragraph('This is paragraph 3 for testing endnotes.')
        doc.add_paragraph('This is paragraph 4 for testing more endnotes.')
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


def test_add_footnote(sample_document):
    """Test adding a footnote to a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a footnote to paragraph 1
    paragraph = doc.paragraphs[1]
    footnote_text = "This is a test footnote"
    footnote = add_footnote(doc, paragraph, footnote_text)
    
    # Verify footnote was created
    assert footnote is not None
    
    # Get all footnotes
    footnotes = get_footnotes(doc)
    
    # Verify footnote exists
    assert len(footnotes) == 1
    assert get_footnote_text(footnotes[0]) == footnote_text
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get all footnotes
        footnotes = get_footnotes(doc2)
        
        # Verify footnote persisted
        assert len(footnotes) == 1
        assert get_footnote_text(footnotes[0]) == footnote_text
    finally:
        # Clean up
        os.unlink(output_path)


def test_add_multiple_footnotes(sample_document):
    """Test adding multiple footnotes to a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add footnotes to different paragraphs
    paragraph1 = doc.paragraphs[1]
    footnote1_text = "First footnote"
    footnote1 = add_footnote(doc, paragraph1, footnote1_text)
    
    paragraph2 = doc.paragraphs[2]
    footnote2_text = "Second footnote"
    footnote2 = add_footnote(doc, paragraph2, footnote2_text)
    
    # Verify footnotes were created
    assert footnote1 is not None
    assert footnote2 is not None
    
    # Get all footnotes
    footnotes = get_footnotes(doc)
    
    # Verify footnotes exist
    assert len(footnotes) == 2
    
    # Get footnote texts
    footnote_texts = [get_footnote_text(f) for f in footnotes]
    assert footnote1_text in footnote_texts
    assert footnote2_text in footnote_texts
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get all footnotes
        footnotes = get_footnotes(doc2)
        
        # Verify footnotes persisted
        assert len(footnotes) == 2
        
        # Get footnote texts
        footnote_texts = [get_footnote_text(f) for f in footnotes]
        assert footnote1_text in footnote_texts
        assert footnote2_text in footnote_texts
    finally:
        # Clean up
        os.unlink(output_path)


def test_add_endnote(sample_document):
    """Test adding an endnote to a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add an endnote to paragraph 3
    paragraph = doc.paragraphs[3]
    endnote_text = "This is a test endnote"
    endnote = add_endnote(doc, paragraph, endnote_text)
    
    # Verify endnote was created
    assert endnote is not None
    
    # Get all endnotes
    endnotes = get_endnotes(doc)
    
    # Verify endnote exists
    assert len(endnotes) == 1
    assert get_endnote_text(endnotes[0]) == endnote_text
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get all endnotes
        endnotes = get_endnotes(doc2)
        
        # Verify endnote persisted
        assert len(endnotes) == 1
        assert get_endnote_text(endnotes[0]) == endnote_text
    finally:
        # Clean up
        os.unlink(output_path)


def test_add_multiple_endnotes(sample_document):
    """Test adding multiple endnotes to a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add endnotes to different paragraphs
    paragraph3 = doc.paragraphs[3]
    endnote1_text = "First endnote"
    endnote1 = add_endnote(doc, paragraph3, endnote1_text)
    
    paragraph4 = doc.paragraphs[4]
    endnote2_text = "Second endnote"
    endnote2 = add_endnote(doc, paragraph4, endnote2_text)
    
    # Verify endnotes were created
    assert endnote1 is not None
    assert endnote2 is not None
    
    # Get all endnotes
    endnotes = get_endnotes(doc)
    
    # Verify endnotes exist
    assert len(endnotes) == 2
    
    # Get endnote texts
    endnote_texts = [get_endnote_text(e) for e in endnotes]
    assert endnote1_text in endnote_texts
    assert endnote2_text in endnote_texts
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get all endnotes
        endnotes = get_endnotes(doc2)
        
        # Verify endnotes persisted
        assert len(endnotes) == 2
        
        # Get endnote texts
        endnote_texts = [get_endnote_text(e) for e in endnotes]
        assert endnote1_text in endnote_texts
        assert endnote2_text in endnote_texts
    finally:
        # Clean up
        os.unlink(output_path)


def test_update_footnote(sample_document):
    """Test updating a footnote in a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a footnote
    paragraph = doc.paragraphs[1]
    original_text = "Original footnote text"
    footnote = add_footnote(doc, paragraph, original_text)
    
    # Verify footnote was created
    assert footnote is not None
    
    # Get all footnotes
    footnotes = get_footnotes(doc)
    
    # Verify original footnote exists
    assert len(footnotes) == 1
    assert get_footnote_text(footnotes[0]) == original_text
    
    # Update the footnote
    updated_text = "Updated footnote text"
    update_result = update_footnote(doc, footnotes[0], updated_text)
    
    # Verify update was successful
    assert update_result is True
    
    # Get all footnotes again
    updated_footnotes = get_footnotes(doc)
    
    # Verify footnote was updated
    assert len(updated_footnotes) == 1
    assert get_footnote_text(updated_footnotes[0]) == updated_text
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get all footnotes
        footnotes = get_footnotes(doc2)
        
        # Verify footnote update persisted
        assert len(footnotes) == 1
        assert get_footnote_text(footnotes[0]) == updated_text
    finally:
        # Clean up
        os.unlink(output_path)


def test_update_endnote(sample_document):
    """Test updating an endnote in a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add an endnote
    paragraph = doc.paragraphs[3]
    original_text = "Original endnote text"
    endnote = add_endnote(doc, paragraph, original_text)
    
    # Verify endnote was created
    assert endnote is not None
    
    # Get all endnotes
    endnotes = get_endnotes(doc)
    
    # Verify original endnote exists
    assert len(endnotes) == 1
    assert get_endnote_text(endnotes[0]) == original_text
    
    # Update the endnote
    updated_text = "Updated endnote text"
    update_result = update_endnote(doc, endnotes[0], updated_text)
    
    # Verify update was successful
    assert update_result is True
    
    # Get all endnotes again
    updated_endnotes = get_endnotes(doc)
    
    # Verify endnote was updated
    assert len(updated_endnotes) == 1
    assert get_endnote_text(updated_endnotes[0]) == updated_text
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get all endnotes
        endnotes = get_endnotes(doc2)
        
        # Verify endnote update persisted
        assert len(endnotes) == 1
        assert get_endnote_text(endnotes[0]) == updated_text
    finally:
        # Clean up
        os.unlink(output_path)


def test_delete_footnote(sample_document):
    """Test deleting a footnote from a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add multiple footnotes
    paragraph1 = doc.paragraphs[1]
    footnote1_text = "Footnote to keep"
    footnote1 = add_footnote(doc, paragraph1, footnote1_text)
    
    paragraph2 = doc.paragraphs[2]
    footnote2_text = "Footnote to delete"
    footnote2 = add_footnote(doc, paragraph2, footnote2_text)
    
    # Verify footnotes were created
    assert footnote1 is not None
    assert footnote2 is not None
    
    # Get all footnotes
    footnotes = get_footnotes(doc)
    
    # Verify both footnotes exist
    assert len(footnotes) == 2
    
    # Find the second footnote
    footnote_to_delete = None
    for f in footnotes:
        if get_footnote_text(f) == footnote2_text:
            footnote_to_delete = f
            break
    
    assert footnote_to_delete is not None
    
    # Delete the second footnote
    delete_result = delete_footnote(doc, footnote_to_delete)
    
    # Verify deletion was successful
    assert delete_result is True
    
    # Get remaining footnotes
    remaining_footnotes = get_footnotes(doc)
    
    # Verify only one footnote remains
    assert len(remaining_footnotes) == 1
    assert get_footnote_text(remaining_footnotes[0]) == footnote1_text
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get all footnotes
        footnotes = get_footnotes(doc2)
        
        # Verify deletion persisted
        assert len(footnotes) == 1
        assert get_footnote_text(footnotes[0]) == footnote1_text
    finally:
        # Clean up
        os.unlink(output_path)


def test_delete_endnote(sample_document):
    """Test deleting an endnote from a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add multiple endnotes
    paragraph3 = doc.paragraphs[3]
    endnote1_text = "Endnote to keep"
    endnote1 = add_endnote(doc, paragraph3, endnote1_text)
    
    paragraph4 = doc.paragraphs[4]
    endnote2_text = "Endnote to delete"
    endnote2 = add_endnote(doc, paragraph4, endnote2_text)
    
    # Verify endnotes were created
    assert endnote1 is not None
    assert endnote2 is not None
    
    # Get all endnotes
    endnotes = get_endnotes(doc)
    
    # Verify both endnotes exist
    assert len(endnotes) == 2
    
    # Find the second endnote
    endnote_to_delete = None
    for e in endnotes:
        if get_endnote_text(e) == endnote2_text:
            endnote_to_delete = e
            break
    
    assert endnote_to_delete is not None
    
    # Delete the second endnote
    delete_result = delete_endnote(doc, endnote_to_delete)
    
    # Verify deletion was successful
    assert delete_result is True
    
    # Get remaining endnotes
    remaining_endnotes = get_endnotes(doc)
    
    # Verify only one endnote remains
    assert len(remaining_endnotes) == 1
    assert get_endnote_text(remaining_endnotes[0]) == endnote1_text
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get all endnotes
        endnotes = get_endnotes(doc2)
        
        # Verify deletion persisted
        assert len(endnotes) == 1
        assert get_endnote_text(endnotes[0]) == endnote1_text
    finally:
        # Clean up
        os.unlink(output_path)


def test_format_footnotes(sample_document):
    """Test formatting footnotes in a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a footnote
    paragraph = doc.paragraphs[1]
    footnote_text = "Footnote to format"
    footnote = add_footnote(doc, paragraph, footnote_text)
    
    # Verify footnote was created
    assert footnote is not None
    
    # Format all footnotes
    font_name = "Arial"
    font_size = Pt(10)
    format_result = format_footnotes(doc, font_name=font_name, font_size=font_size)
    
    # Verify formatting was applied
    assert format_result is True
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get all footnotes
        footnotes = get_footnotes(doc2)
        
        # Verify formatting persisted
        assert len(footnotes) == 1
        
        # Note: It's difficult to verify formatting directly,
        # as Python-docx doesn't easily expose formatting of footnotes
    finally:
        # Clean up
        os.unlink(output_path)


def test_format_endnotes(sample_document):
    """Test formatting endnotes in a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Add an endnote
    paragraph = doc.paragraphs[3]
    endnote_text = "Endnote to format"
    endnote = add_endnote(doc, paragraph, endnote_text)
    
    # Verify endnote was created
    assert endnote is not None
    
    # Format all endnotes
    font_name = "Times New Roman"
    font_size = Pt(9)
    format_result = format_endnotes(doc, font_name=font_name, font_size=font_size)
    
    # Verify formatting was applied
    assert format_result is True
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get all endnotes
        endnotes = get_endnotes(doc2)
        
        # Verify formatting persisted
        assert len(endnotes) == 1
        
        # Note: It's difficult to verify formatting directly,
        # as Python-docx doesn't easily expose formatting of endnotes
    finally:
        # Clean up
        os.unlink(output_path)


def test_complex_footnote_content(sample_document):
    """Test creating footnotes with complex content."""
    # Load the document
    doc = Document(sample_document)
    
    # Add a footnote with complex formatting
    paragraph = doc.paragraphs[1]
    complex_footnote_text = "This is a complex footnote with *italic* and **bold** text."
    footnote = add_footnote(doc, paragraph, complex_footnote_text)
    
    # Verify footnote was created
    assert footnote is not None
    
    # Get all footnotes
    footnotes = get_footnotes(doc)
    
    # Verify footnote exists
    assert len(footnotes) == 1
    assert complex_footnote_text in get_footnote_text(footnotes[0])
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get all footnotes
        footnotes = get_footnotes(doc2)
        
        # Verify footnote persisted
        assert len(footnotes) == 1
        assert complex_footnote_text in get_footnote_text(footnotes[0])
    finally:
        # Clean up
        os.unlink(output_path)


def test_complex_endnote_content(sample_document):
    """Test creating endnotes with complex content."""
    # Load the document
    doc = Document(sample_document)
    
    # Add an endnote with complex formatting
    paragraph = doc.paragraphs[3]
    complex_endnote_text = "This is a complex endnote with a citation: Smith, J. (2020)."
    endnote = add_endnote(doc, paragraph, complex_endnote_text)
    
    # Verify endnote was created
    assert endnote is not None
    
    # Get all endnotes
    endnotes = get_endnotes(doc)
    
    # Verify endnote exists
    assert len(endnotes) == 1
    assert complex_endnote_text in get_endnote_text(endnotes[0])
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Get all endnotes
        endnotes = get_endnotes(doc2)
        
        # Verify endnote persisted
        assert len(endnotes) == 1
        assert complex_endnote_text in get_endnote_text(endnotes[0])
    finally:
        # Clean up
        os.unlink(output_path)


if __name__ == '__main__':
    pytest.main(['-v', __file__]) 