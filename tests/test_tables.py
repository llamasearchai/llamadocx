"""
Tables tests for LlamaDocx.

This module contains tests for the table handling functionality of the LlamaDocx package.
"""

import os
import tempfile
import pytest
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from llamadocx.tables import (
    create_table,
    add_row,
    add_column,
    set_table_style,
    set_cell_text,
    set_cell_formatting,
    merge_cells,
    set_column_width,
    set_table_alignment,
    set_cell_vertical_alignment,
    set_cell_background,
    set_table_borders,
    set_cell_borders,
    delete_row,
    delete_column
)


@pytest.fixture
def sample_document():
    """Create a sample document for testing table operations."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        
        # Add heading
        doc.add_heading('Table Handling Test Document', level=1)
        
        # Add some content
        doc.add_paragraph('This document is used for testing table functionality.')
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


def test_create_table(sample_document):
    """Test creating a table."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a simple table
    table = create_table(doc, rows=3, cols=3)
    
    # Verify table was created with correct dimensions
    assert len(table.rows) == 3
    assert len(table.columns) == 3
    
    # Fill in some cells to verify they're accessible
    set_cell_text(table.cell(0, 0), "Cell 0,0")
    set_cell_text(table.cell(1, 1), "Cell 1,1")
    set_cell_text(table.cell(2, 2), "Cell 2,2")
    
    # Verify cell contents
    assert table.cell(0, 0).text == "Cell 0,0"
    assert table.cell(1, 1).text == "Cell 1,1"
    assert table.cell(2, 2).text == "Cell 2,2"
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify table persisted
        assert len(doc2.tables) == 1
        table2 = doc2.tables[0]
        assert len(table2.rows) == 3
        assert len(table2.columns) == 3
        assert table2.cell(0, 0).text == "Cell 0,0"
        assert table2.cell(1, 1).text == "Cell 1,1"
        assert table2.cell(2, 2).text == "Cell 2,2"
    finally:
        # Clean up
        os.unlink(output_path)


def test_add_row_column(sample_document):
    """Test adding rows and columns to a table."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a table
    table = create_table(doc, rows=2, cols=2)
    
    # Add content to initial cells for reference
    set_cell_text(table.cell(0, 0), "Top Left")
    set_cell_text(table.cell(0, 1), "Top Right")
    set_cell_text(table.cell(1, 0), "Bottom Left")
    set_cell_text(table.cell(1, 1), "Bottom Right")
    
    # Add a new row
    new_row = add_row(table)
    set_cell_text(new_row.cells[0], "New Row Left")
    set_cell_text(new_row.cells[1], "New Row Right")
    
    # Add a new column
    add_column(table)
    set_cell_text(table.cell(0, 2), "Top New Column")
    set_cell_text(table.cell(1, 2), "Middle New Column")
    set_cell_text(table.cell(2, 2), "Bottom New Column")
    
    # Verify table dimensions
    assert len(table.rows) == 3
    assert len(table.columns) == 3
    
    # Verify cell contents
    assert table.cell(0, 0).text == "Top Left"
    assert table.cell(0, 1).text == "Top Right"
    assert table.cell(0, 2).text == "Top New Column"
    assert table.cell(1, 0).text == "Bottom Left"
    assert table.cell(1, 1).text == "Bottom Right"
    assert table.cell(1, 2).text == "Middle New Column"
    assert table.cell(2, 0).text == "New Row Left"
    assert table.cell(2, 1).text == "New Row Right"
    assert table.cell(2, 2).text == "Bottom New Column"
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify table persisted with correct dimensions
        table2 = doc2.tables[0]
        assert len(table2.rows) == 3
        assert len(table2.columns) == 3
        
        # Verify cell contents persisted
        assert table2.cell(0, 0).text == "Top Left"
        assert table2.cell(0, 1).text == "Top Right"
        assert table2.cell(0, 2).text == "Top New Column"
        assert table2.cell(1, 0).text == "Bottom Left"
        assert table2.cell(1, 1).text == "Bottom Right"
        assert table2.cell(1, 2).text == "Middle New Column"
        assert table2.cell(2, 0).text == "New Row Left"
        assert table2.cell(2, 1).text == "New Row Right"
        assert table2.cell(2, 2).text == "Bottom New Column"
    finally:
        # Clean up
        os.unlink(output_path)


def test_delete_row_column(sample_document):
    """Test deleting rows and columns from a table."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a table with content
    table = create_table(doc, rows=3, cols=3)
    
    # Add content to cells
    for row in range(3):
        for col in range(3):
            set_cell_text(table.cell(row, col), f"Cell {row},{col}")
    
    # Delete a row and column
    delete_row(table, 1)  # Delete middle row
    delete_column(table, 1)  # Delete middle column
    
    # Verify table dimensions
    assert len(table.rows) == 2
    assert len(table.columns) == 2
    
    # Verify remaining cell contents (original positions)
    assert table.cell(0, 0).text == "Cell 0,0"
    assert table.cell(0, 1).text == "Cell 0,2"  # Column 1 was deleted
    assert table.cell(1, 0).text == "Cell 2,0"  # Row 1 was deleted
    assert table.cell(1, 1).text == "Cell 2,2"  # Both row 1 and column 1 were deleted
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify table persisted with correct dimensions
        table2 = doc2.tables[0]
        assert len(table2.rows) == 2
        assert len(table2.columns) == 2
        
        # Verify cell contents persisted
        assert table2.cell(0, 0).text == "Cell 0,0"
        assert table2.cell(0, 1).text == "Cell 0,2"
        assert table2.cell(1, 0).text == "Cell 2,0"
        assert table2.cell(1, 1).text == "Cell 2,2"
    finally:
        # Clean up
        os.unlink(output_path)


def test_set_table_style(sample_document):
    """Test setting table style."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a table
    table = create_table(doc, rows=3, cols=3)
    
    # Add content to cells
    for row in range(3):
        for col in range(3):
            set_cell_text(table.cell(row, col), f"Cell {row},{col}")
    
    # Apply a table style
    set_table_style(table, "LightShading-Accent1")
    
    # Verify style was applied
    assert table.style.name == "LightShading-Accent1"
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify style persisted
        table2 = doc2.tables[0]
        assert table2.style.name == "LightShading-Accent1"
    finally:
        # Clean up
        os.unlink(output_path)


def test_cell_formatting(sample_document):
    """Test cell text formatting."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a table
    table = create_table(doc, rows=2, cols=2)
    
    # Add content to cells with formatting
    cell = table.cell(0, 0)
    set_cell_text(cell, "Bold Text")
    set_cell_formatting(cell, bold=True)
    
    cell = table.cell(0, 1)
    set_cell_text(cell, "Italic Text")
    set_cell_formatting(cell, italic=True)
    
    cell = table.cell(1, 0)
    set_cell_text(cell, "Underlined Text")
    set_cell_formatting(cell, underline=True)
    
    cell = table.cell(1, 1)
    set_cell_text(cell, "Colored Text")
    set_cell_formatting(cell, font="Arial", font_size=Pt(14), color=RGBColor(255, 0, 0))
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify formatting persisted
        table2 = doc2.tables[0]
        
        # Check formatting by examining the runs
        assert table2.cell(0, 0).paragraphs[0].runs[0].bold == True
        assert table2.cell(0, 1).paragraphs[0].runs[0].italic == True
        assert table2.cell(1, 0).paragraphs[0].runs[0].underline == True
        
        # Check font properties
        run = table2.cell(1, 1).paragraphs[0].runs[0]
        assert run.font.name == "Arial"
        assert run.font.size == Pt(14)
        assert run.font.color.rgb == RGBColor(255, 0, 0)
    finally:
        # Clean up
        os.unlink(output_path)


def test_merge_cells(sample_document):
    """Test merging cells."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a table
    table = create_table(doc, rows=3, cols=3)
    
    # Add content to cells
    for row in range(3):
        for col in range(3):
            set_cell_text(table.cell(row, col), f"Cell {row},{col}")
    
    # Merge cells horizontally in first row
    merge_cells(table, 0, 0, 0, 2)
    set_cell_text(table.cell(0, 0), "Merged Row 1")
    
    # Merge cells vertically in last column
    merge_cells(table, 1, 2, 2, 2)
    set_cell_text(table.cell(1, 2), "Merged Column 3")
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify table and merged cells
        table2 = doc2.tables[0]
        
        # Check merged cells content
        assert table2.cell(0, 0).text == "Merged Row 1"
        assert table2.cell(1, 2).text == "Merged Column 3"
        
        # Cell references that are part of merged cells will point to the
        # first cell of the merged region, so we can verify by comparing text
        assert table2.cell(0, 1).text == "Merged Row 1"
        assert table2.cell(0, 2).text == "Merged Row 1"
        assert table2.cell(2, 2).text == "Merged Column 3"
    finally:
        # Clean up
        os.unlink(output_path)


def test_column_width(sample_document):
    """Test setting column width."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a table
    table = create_table(doc, rows=2, cols=3)
    
    # Set different column widths
    set_column_width(table.columns[0], Inches(1.0))
    set_column_width(table.columns[1], Inches(2.0))
    set_column_width(table.columns[2], Inches(1.5))
    
    # Add content to cells
    for row in range(2):
        for col in range(3):
            set_cell_text(table.cell(row, col), f"Cell {row},{col}")
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify column widths
        table2 = doc2.tables[0]
        assert table2.columns[0].width == Inches(1.0)
        assert table2.columns[1].width == Inches(2.0)
        assert table2.columns[2].width == Inches(1.5)
    finally:
        # Clean up
        os.unlink(output_path)


def test_table_alignment(sample_document):
    """Test setting table alignment."""
    # Load the document
    doc = Document(sample_document)
    
    # Create tables with different alignments
    table1 = create_table(doc, rows=2, cols=2)
    set_table_alignment(table1, WD_TABLE_ALIGNMENT.LEFT)
    set_cell_text(table1.cell(0, 0), "Left-aligned table")
    
    doc.add_paragraph() # Add some space between tables
    
    table2 = create_table(doc, rows=2, cols=2)
    set_table_alignment(table2, WD_TABLE_ALIGNMENT.CENTER)
    set_cell_text(table2.cell(0, 0), "Center-aligned table")
    
    doc.add_paragraph() # Add some space between tables
    
    table3 = create_table(doc, rows=2, cols=2)
    set_table_alignment(table3, WD_TABLE_ALIGNMENT.RIGHT)
    set_cell_text(table3.cell(0, 0), "Right-aligned table")
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify table alignments
        tables = doc2.tables
        assert tables[0].alignment == WD_TABLE_ALIGNMENT.LEFT
        assert tables[1].alignment == WD_TABLE_ALIGNMENT.CENTER
        assert tables[2].alignment == WD_TABLE_ALIGNMENT.RIGHT
    finally:
        # Clean up
        os.unlink(output_path)


def test_cell_vertical_alignment(sample_document):
    """Test setting cell vertical alignment."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a table
    table = create_table(doc, rows=1, cols=3)
    
    # Set different vertical alignments
    set_cell_vertical_alignment(table.cell(0, 0), WD_CELL_VERTICAL_ALIGNMENT.TOP)
    set_cell_text(table.cell(0, 0), "Top aligned")
    
    set_cell_vertical_alignment(table.cell(0, 1), WD_CELL_VERTICAL_ALIGNMENT.CENTER)
    set_cell_text(table.cell(0, 1), "Center aligned")
    
    set_cell_vertical_alignment(table.cell(0, 2), WD_CELL_VERTICAL_ALIGNMENT.BOTTOM)
    set_cell_text(table.cell(0, 2), "Bottom aligned")
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify vertical alignments
        table2 = doc2.tables[0]
        assert table2.cell(0, 0).vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.TOP
        assert table2.cell(0, 1).vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER
        assert table2.cell(0, 2).vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    finally:
        # Clean up
        os.unlink(output_path)


def test_cell_background(sample_document):
    """Test setting cell background color."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a table
    table = create_table(doc, rows=1, cols=3)
    
    # Set different background colors
    set_cell_background(table.cell(0, 0), RGBColor(255, 0, 0))  # Red
    set_cell_text(table.cell(0, 0), "Red background")
    
    set_cell_background(table.cell(0, 1), RGBColor(0, 255, 0))  # Green
    set_cell_text(table.cell(0, 1), "Green background")
    
    set_cell_background(table.cell(0, 2), RGBColor(0, 0, 255))  # Blue
    set_cell_text(table.cell(0, 2), "Blue background")
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify cell content (color properties are harder to check)
        table2 = doc2.tables[0]
        assert table2.cell(0, 0).text == "Red background"
        assert table2.cell(0, 1).text == "Green background"
        assert table2.cell(0, 2).text == "Blue background"
        
        # Verify cell background colors (implementation may vary)
        # Note: Some properties like cell colors require special handling to verify
    finally:
        # Clean up
        os.unlink(output_path)


def test_table_borders(sample_document):
    """Test setting table borders."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a table
    table = create_table(doc, rows=3, cols=3)
    
    # Add content to cells
    for row in range(3):
        for col in range(3):
            set_cell_text(table.cell(row, col), f"Cell {row},{col}")
    
    # Set table borders (all around)
    set_table_borders(table, width=2)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify table content (border properties are harder to check)
        table2 = doc2.tables[0]
        for row in range(3):
            for col in range(3):
                assert table2.cell(row, col).text == f"Cell {row},{col}"
    finally:
        # Clean up
        os.unlink(output_path)


def test_cell_borders(sample_document):
    """Test setting individual cell borders."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a table
    table = create_table(doc, rows=2, cols=2)
    
    # Add content to cells
    for row in range(2):
        for col in range(2):
            set_cell_text(table.cell(row, col), f"Cell {row},{col}")
    
    # Set border for specific cell
    set_cell_borders(table.cell(0, 0), top=True, left=True, bottom=True, right=True, width=4)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify table content (border properties are harder to check)
        table2 = doc2.tables[0]
        assert table2.cell(0, 0).text == "Cell 0,0"
    finally:
        # Clean up
        os.unlink(output_path)


def test_complex_table(sample_document):
    """Test creating a complex table with various features."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a table with multiple features
    table = create_table(doc, rows=4, cols=3)
    
    # Set table alignment
    set_table_alignment(table, WD_TABLE_ALIGNMENT.CENTER)
    
    # Set column widths
    set_column_width(table.columns[0], Inches(1.5))
    set_column_width(table.columns[1], Inches(2.0))
    set_column_width(table.columns[2], Inches(1.5))
    
    # Add header row with merged cells and formatting
    merge_cells(table, 0, 0, 0, 2)
    header_cell = table.cell(0, 0)
    set_cell_text(header_cell, "Complex Table Example")
    set_cell_formatting(header_cell, bold=True, font_size=Pt(14), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_background(header_cell, RGBColor(200, 200, 200))
    
    # Add subheaders
    set_cell_text(table.cell(1, 0), "Column 1")
    set_cell_formatting(table.cell(1, 0), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    set_cell_text(table.cell(1, 1), "Column 2")
    set_cell_formatting(table.cell(1, 1), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    set_cell_text(table.cell(1, 2), "Column 3")
    set_cell_formatting(table.cell(1, 2), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Add data with different alignments and formatting
    set_cell_text(table.cell(2, 0), "Left Data")
    set_cell_formatting(table.cell(2, 0), alignment=WD_ALIGN_PARAGRAPH.LEFT)
    
    set_cell_text(table.cell(2, 1), "Center Data")
    set_cell_formatting(table.cell(2, 1), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_vertical_alignment(table.cell(2, 1), WD_CELL_VERTICAL_ALIGNMENT.CENTER)
    
    set_cell_text(table.cell(2, 2), "Right Data")
    set_cell_formatting(table.cell(2, 2), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    
    # Merge cells in the bottom row
    merge_cells(table, 3, 0, 3, 2)
    footer_cell = table.cell(3, 0)
    set_cell_text(footer_cell, "Table Footer")
    set_cell_formatting(footer_cell, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Add borders
    set_table_borders(table, width=1)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify table properties
        table2 = doc2.tables[0]
        assert len(table2.rows) == 4
        assert len(table2.columns) == 3
        
        # Verify column widths
        assert table2.columns[0].width == Inches(1.5)
        assert table2.columns[1].width == Inches(2.0)
        assert table2.columns[2].width == Inches(1.5)
        
        # Verify content
        assert table2.cell(0, 0).text == "Complex Table Example"
        assert table2.cell(1, 0).text == "Column 1"
        assert table2.cell(1, 1).text == "Column 2"
        assert table2.cell(1, 2).text == "Column 3"
        assert table2.cell(2, 0).text == "Left Data"
        assert table2.cell(2, 1).text == "Center Data"
        assert table2.cell(2, 2).text == "Right Data"
        assert table2.cell(3, 0).text == "Table Footer"
        
        # Verify merged cells
        assert table2.cell(0, 1).text == "Complex Table Example"
        assert table2.cell(0, 2).text == "Complex Table Example"
        assert table2.cell(3, 1).text == "Table Footer"
        assert table2.cell(3, 2).text == "Table Footer"
    finally:
        # Clean up
        os.unlink(output_path)


if __name__ == '__main__':
    pytest.main(['-v', __file__]) 