"""
Template tests for LlamaDocx.

This module contains tests for the template processing functionality of the LlamaDocx package.
"""

import os
import tempfile
from pathlib import Path
import pytest
from docx import Document

from llamadocx.template import Template


@pytest.fixture
def basic_template():
    """Create a basic template document with simple fields."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        doc.add_heading('{{title}}', level=1)
        doc.add_paragraph('Name: {{name}}')
        doc.add_paragraph('Date: {{date}}')
        doc.add_paragraph('Description: {{description}}')
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


@pytest.fixture
def table_template():
    """Create a template document with tables."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        doc.add_heading('{{report_title}}', level=1)
        doc.add_paragraph('Prepared by: {{author}}')
        
        # Add a table with template fields
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        # Header row
        table.cell(0, 0).text = 'Item'
        table.cell(0, 1).text = 'Value'
        
        # Data rows
        table.cell(1, 0).text = 'Project'
        table.cell(1, 1).text = '{{project_name}}'
        
        table.cell(2, 0).text = 'Client'
        table.cell(2, 1).text = '{{client_name}}'
        
        table.cell(3, 0).text = 'Date'
        table.cell(3, 1).text = '{{report_date}}'
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


@pytest.fixture
def repeating_section_template():
    """Create a template document with a repeating section."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        doc.add_heading('{{company}} Invoice', level=1)
        doc.add_paragraph('Invoice #: {{invoice_number}}')
        doc.add_paragraph('Date: {{invoice_date}}')
        
        # Add header for items section
        doc.add_heading('Items', level=2)
        
        # Add a repeating section for invoice items
        # First, we need to add a marker paragraph
        doc.add_paragraph('{{#items}}')
        
        # Then add the template for each item
        doc.add_paragraph('Item: {{name}}')
        doc.add_paragraph('Quantity: {{quantity}}')
        doc.add_paragraph('Price: ${{price}}')
        doc.add_paragraph('Subtotal: ${{subtotal}}')
        
        # End the repeating section
        doc.add_paragraph('{{/items}}')
        
        # Add total
        doc.add_paragraph('Total: ${{total}}')
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


def test_template_initialization(basic_template):
    """Test template initialization and field detection."""
    # Create a template instance
    template = Template(basic_template)
    
    # Test field detection
    fields = template.get_fields()
    
    # Check that all fields were detected
    assert 'title' in fields
    assert 'name' in fields
    assert 'date' in fields
    assert 'description' in fields
    
    # Check field count
    assert len(fields) == 4


def test_template_custom_delimiters(basic_template):
    """Test template with custom field delimiters."""
    # First, create a document with different delimiters
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        doc.add_heading('<<title>>', level=1)
        doc.add_paragraph('Name: <<name>>')
        doc.add_paragraph('Date: <<date>>')
        
        # Save the document
        doc.save(tmp.name)
    
    try:
        # Create a template instance with custom delimiters
        template = Template(tmp.name, field_delimiters=('<<', '>>'))
        
        # Test field detection
        fields = template.get_fields()
        
        # Check that all fields were detected
        assert 'title' in fields
        assert 'name' in fields
        assert 'date' in fields
        
        # Check field count
        assert len(fields) == 3
    finally:
        # Clean up
        os.unlink(tmp.name)


def test_basic_field_merging(basic_template):
    """Test basic field merging functionality."""
    # Create test data
    data = {
        'title': 'Test Document',
        'name': 'John Doe',
        'date': '2023-06-15',
        'description': 'This is a test document generated from a template.'
    }
    
    # Create template and merge fields
    template = Template(basic_template)
    template.merge_fields(data)
    
    # Save the processed document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        template.save(output_path)
    
    try:
        # Load the processed document and verify content
        doc = Document(output_path)
        
        # Check heading
        assert doc.paragraphs[0].text == data['title']
        
        # Check paragraphs
        assert f"Name: {data['name']}" in doc.paragraphs[1].text
        assert f"Date: {data['date']}" in doc.paragraphs[2].text
        assert f"Description: {data['description']}" in doc.paragraphs[3].text
    finally:
        # Clean up
        os.unlink(output_path)


def test_table_field_merging(table_template):
    """Test field merging in tables."""
    # Create test data
    data = {
        'report_title': 'Quarterly Report',
        'author': 'Jane Smith',
        'project_name': 'Project Alpha',
        'client_name': 'Acme Corporation',
        'report_date': '2023-06-30'
    }
    
    # Create template and merge fields
    template = Template(table_template)
    template.merge_fields(data)
    
    # Save the processed document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        template.save(output_path)
    
    try:
        # Load the processed document and verify content
        doc = Document(output_path)
        
        # Check heading and paragraph
        assert doc.paragraphs[0].text == data['report_title']
        assert f"Prepared by: {data['author']}" in doc.paragraphs[1].text
        
        # Check table content
        table = doc.tables[0]
        assert table.cell(1, 1).text == data['project_name']
        assert table.cell(2, 1).text == data['client_name']
        assert table.cell(3, 1).text == data['report_date']
    finally:
        # Clean up
        os.unlink(output_path)


def test_repeating_section(repeating_section_template):
    """Test repeating sections in templates."""
    # Create test data with repeating items
    data = {
        'company': 'ABC Company',
        'invoice_number': 'INV-12345',
        'invoice_date': '2023-07-01',
        'items': [
            {
                'name': 'Widget A',
                'quantity': '2',
                'price': '10.00',
                'subtotal': '20.00'
            },
            {
                'name': 'Widget B',
                'quantity': '3',
                'price': '15.00',
                'subtotal': '45.00'
            },
            {
                'name': 'Widget C',
                'quantity': '1',
                'price': '25.00',
                'subtotal': '25.00'
            }
        ],
        'total': '90.00'
    }
    
    # Create template and merge fields
    template = Template(repeating_section_template)
    template.merge_fields(data)
    
    # Save the processed document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        template.save(output_path)
    
    try:
        # Load the processed document and verify content
        doc = Document(output_path)
        
        # Convert all paragraphs to text for easier checking
        paragraphs = [p.text for p in doc.paragraphs]
        
        # Check basic content
        assert f"{data['company']} Invoice" in paragraphs
        assert f"Invoice #: {data['invoice_number']}" in paragraphs
        assert f"Invoice #: {data['invoice_number']}" in paragraphs
        
        # Check that section markers are removed
        assert "{{#items}}" not in paragraphs
        assert "{{/items}}" not in paragraphs
        
        # Check that all items are in the document
        for item in data['items']:
            assert f"Item: {item['name']}" in paragraphs
            assert f"Quantity: {item['quantity']}" in paragraphs
            assert f"Price: ${item['price']}" in paragraphs
            assert f"Subtotal: ${item['subtotal']}" in paragraphs
        
        # Check total
        assert f"Total: ${data['total']}" in paragraphs
    finally:
        # Clean up
        os.unlink(output_path)


def test_missing_field_handling(basic_template):
    """Test how the template handles missing fields."""
    # Create test data with missing fields
    data = {
        'title': 'Incomplete Data Test',
        'name': 'John Doe'
        # 'date' and 'description' are missing
    }
    
    # Create template with remove_empty_fields=True
    template = Template(basic_template)
    template.merge_fields(data, remove_empty_fields=True)
    
    # Save the processed document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        template.save(output_path)
    
    try:
        # Load the processed document and verify content
        doc = Document(output_path)
        
        # Convert all paragraphs to text for easier checking
        paragraphs = [p.text for p in doc.paragraphs]
        
        # Check that present fields are merged
        assert data['title'] in paragraphs
        assert f"Name: {data['name']}" in paragraphs
        
        # Check that missing fields are removed
        assert "Date: {{date}}" not in paragraphs
        assert "Description: {{description}}" not in paragraphs
        
        # Check that the empty field placeholders are gone
        assert "{{date}}" not in " ".join(paragraphs)
        assert "{{description}}" not in " ".join(paragraphs)
    finally:
        # Clean up
        os.unlink(output_path)


def test_field_adding(basic_template):
    """Test adding a field to a document."""
    # Create a template instance
    template = Template(basic_template)
    
    # Add a new field
    template.add_field(paragraph_index=4, field_name="new_field", text_before="New Field: ")
    
    # Save the modified template
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        template.save(output_path)
    
    try:
        # Create a new template from the modified document
        new_template = Template(output_path)
        
        # Check that the new field was added
        fields = new_template.get_fields()
        assert "new_field" in fields
        
        # Test merging with the new field
        data = {
            'title': 'Field Addition Test',
            'name': 'John Doe',
            'date': '2023-06-15',
            'description': 'Testing field addition.',
            'new_field': 'This is a new field value'
        }
        
        new_template.merge_fields(data)
        new_template.save(output_path)
        
        # Verify that the new field was properly merged
        doc = Document(output_path)
        paragraphs = [p.text for p in doc.paragraphs]
        
        assert "New Field: This is a new field value" in paragraphs
    finally:
        # Clean up
        os.unlink(output_path)


if __name__ == '__main__':
    pytest.main(['-v', __file__]) 