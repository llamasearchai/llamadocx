"""
Metadata tests for LlamaDocx.

This module contains tests for the metadata handling functionality of the LlamaDocx package.
"""

import os
import json
import tempfile
from datetime import datetime
from pathlib import Path
import pytest
from docx import Document

from llamadocx.metadata import (
    get_metadata, set_metadata, set_title, set_author, set_subject,
    set_keywords, set_category, set_comments, set_created_time,
    set_last_modified_time, update_metadata_from_file, extract_metadata_to_file
)


@pytest.fixture
def temp_docx():
    """Create a temporary DOCX file for testing metadata."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        doc.add_heading('Metadata Test Document', level=1)
        doc.add_paragraph('This is a test document for metadata operations.')
        
        # Set some initial metadata
        core_props = doc.core_properties
        core_props.title = "Initial Title"
        core_props.author = "Initial Author"
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


def test_get_metadata(temp_docx):
    """Test retrieving document metadata."""
    # Load the document
    doc = Document(temp_docx)
    
    # Get metadata
    metadata = get_metadata(doc)
    
    # Verify initial metadata is retrieved correctly
    assert metadata['title'] == "Initial Title"
    assert metadata['author'] == "Initial Author"
    
    # Verify other fields exist but may be None or empty
    assert 'subject' in metadata
    assert 'keywords' in metadata
    assert 'category' in metadata
    assert 'created' in metadata
    assert 'modified' in metadata


def test_set_metadata(temp_docx):
    """Test setting document metadata with a dictionary."""
    # Prepare test metadata
    test_metadata = {
        'title': 'Updated Title',
        'author': 'Updated Author',
        'subject': 'Metadata Testing',
        'keywords': 'test,metadata,python',
        'category': 'Testing',
        'comments': 'This is a test comment'
    }
    
    # Load the document
    doc = Document(temp_docx)
    
    # Set metadata
    set_metadata(doc, test_metadata)
    doc.save(temp_docx)
    
    # Reload the document and check metadata
    doc = Document(temp_docx)
    metadata = get_metadata(doc)
    
    # Verify all fields were updated
    for key, value in test_metadata.items():
        assert metadata[key] == value


def test_individual_metadata_setters(temp_docx):
    """Test individual metadata setter functions."""
    # Load the document
    doc = Document(temp_docx)
    
    # Set metadata using individual functions
    set_title(doc, "Function Set Title")
    set_author(doc, "Function Set Author")
    set_subject(doc, "Function Set Subject")
    set_keywords(doc, "function,set,keywords")
    set_category(doc, "Function Tests")
    set_comments(doc, "Comment set via function")
    
    # Save the document
    doc.save(temp_docx)
    
    # Reload and verify
    doc = Document(temp_docx)
    metadata = get_metadata(doc)
    
    assert metadata['title'] == "Function Set Title"
    assert metadata['author'] == "Function Set Author"
    assert metadata['subject'] == "Function Set Subject"
    assert metadata['keywords'] == "function,set,keywords"
    assert metadata['category'] == "Function Tests"
    assert metadata['comments'] == "Comment set via function"


def test_datetime_metadata(temp_docx):
    """Test setting datetime metadata."""
    # Load the document
    doc = Document(temp_docx)
    
    # Set datetime values
    now = datetime.now()
    yesterday = datetime(now.year, now.month, now.day - 1, 12, 0, 0)
    
    set_created_time(doc, yesterday)
    set_last_modified_time(doc, now)
    
    # Save the document
    doc.save(temp_docx)
    
    # Reload and verify
    doc = Document(temp_docx)
    metadata = get_metadata(doc)
    
    # Datetime comparison can be tricky due to serialization/deserialization
    # We'll check that the dates are close to what we set
    created_time = datetime.fromisoformat(metadata['created'].replace('Z', '+00:00'))
    modified_time = datetime.fromisoformat(metadata['modified'].replace('Z', '+00:00'))
    
    assert created_time.date() == yesterday.date()
    assert modified_time.date() == now.date()


def test_metadata_file_operations(temp_docx):
    """Test extracting and updating metadata from/to files."""
    # First set some metadata to extract
    doc = Document(temp_docx)
    
    test_metadata = {
        'title': 'File Operation Test',
        'author': 'File Operation Author',
        'subject': 'File Operations',
        'keywords': 'file,json,metadata',
        'category': 'File Tests'
    }
    
    set_metadata(doc, test_metadata)
    doc.save(temp_docx)
    
    # Extract metadata to a file
    with tempfile.NamedTemporaryFile(suffix='.json', delete=False) as tmp:
        json_path = tmp.name
    
    try:
        # Extract to file
        extract_metadata_to_file(doc, json_path)
        
        # Verify the file was created and contains valid JSON
        assert os.path.exists(json_path)
        
        with open(json_path, 'r') as f:
            extracted_data = json.load(f)
        
        # Verify extracted data matches what we set
        for key, value in test_metadata.items():
            assert extracted_data[key] == value
            
        # Now modify the JSON and update back to the document
        extracted_data['title'] = 'Updated via JSON'
        extracted_data['author'] = 'JSON Updater'
        
        with open(json_path, 'w') as f:
            json.dump(extracted_data, f)
            
        # Update document from the modified JSON
        update_metadata_from_file(doc, json_path)
        doc.save(temp_docx)
        
        # Reload and verify
        doc = Document(temp_docx)
        metadata = get_metadata(doc)
        
        assert metadata['title'] == 'Updated via JSON'
        assert metadata['author'] == 'JSON Updater'
        assert metadata['subject'] == 'File Operations'  # Should be unchanged
    finally:
        # Clean up
        if os.path.exists(json_path):
            os.unlink(json_path)


def test_metadata_overwrite(temp_docx):
    """Test that metadata updates correctly overwrite existing values."""
    # Set initial metadata
    doc = Document(temp_docx)
    
    initial_metadata = {
        'title': 'Initial Title',
        'author': 'Initial Author',
        'subject': 'Initial Subject',
        'keywords': 'initial,keywords',
        'category': 'Initial Category',
        'comments': 'Initial comments'
    }
    
    set_metadata(doc, initial_metadata)
    doc.save(temp_docx)
    
    # Now update only some fields
    doc = Document(temp_docx)
    
    update_metadata = {
        'title': 'Updated Title',
        'author': 'Updated Author'
    }
    
    set_metadata(doc, update_metadata)
    doc.save(temp_docx)
    
    # Reload and verify
    doc = Document(temp_docx)
    metadata = get_metadata(doc)
    
    # Updated fields should have new values
    assert metadata['title'] == 'Updated Title'
    assert metadata['author'] == 'Updated Author'
    
    # Non-updated fields should retain their original values
    assert metadata['subject'] == 'Initial Subject'
    assert metadata['keywords'] == 'initial,keywords'
    assert metadata['category'] == 'Initial Category'
    assert metadata['comments'] == 'Initial comments'


if __name__ == '__main__':
    pytest.main(['-v', __file__]) 