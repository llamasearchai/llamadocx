"""
Document comparison tests for LlamaDocx.

This module contains tests for the document comparison functionality of the LlamaDocx package.
"""

import os
import tempfile
import pytest
from docx import Document
from datetime import datetime

from llamadocx.compare import (
    compare_documents,
    get_revisions,
    get_revision_details,
    get_deleted_text,
    get_inserted_text,
    accept_revisions,
    reject_revisions,
    merge_documents,
    create_redline_document
)


@pytest.fixture
def original_document():
    """Create a sample original document for comparison testing."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        
        # Add heading
        doc.add_heading('Document Comparison Test - Original', level=1)
        
        # Add content
        doc.add_paragraph('This is the first paragraph in the original document.')
        doc.add_paragraph('This is the second paragraph that will remain unchanged.')
        doc.add_paragraph('This is the third paragraph that will be modified.')
        doc.add_paragraph('This is the fourth paragraph that will be deleted.')
        doc.add_paragraph('This paragraph has some text that will be partially modified.')
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


@pytest.fixture
def modified_document(original_document):
    """Create a modified version of the original document for comparison testing."""
    # Load the original document
    doc = Document(original_document)
    
    # Make modifications
    # Keep the heading and first paragraph
    
    # Modify the third paragraph (index 2 after heading)
    doc.paragraphs[2].text = 'This is the second paragraph that will remain unchanged.'
    
    # Replace the third paragraph with new text
    doc.paragraphs[3].text = 'This is the third paragraph with modified content.'
    
    # Delete the fourth paragraph
    p = doc.paragraphs[4]
    p._element.getparent().remove(p._element)
    
    # Modify part of the text in the fifth paragraph
    doc.paragraphs[4].text = 'This paragraph has some text that was partially changed.'
    
    # Add a new paragraph
    doc.add_paragraph('This is a new paragraph added to the modified document.')
    
    # Save the modified document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        modified_path = tmp.name
    
    yield modified_path
    
    # Clean up
    os.unlink(modified_path)


def test_compare_documents(original_document, modified_document):
    """Test comparing two documents."""
    # Compare the documents
    comparison_result = compare_documents(original_document, modified_document)
    
    # Verify comparison result contains expected information
    assert comparison_result is not None
    assert 'diff_doc' in comparison_result
    assert 'summary' in comparison_result
    
    # Check summary contains expected counts
    summary = comparison_result['summary']
    assert summary['insertions'] > 0
    assert summary['deletions'] > 0
    assert summary['moves'] >= 0
    assert summary['formatting_changes'] >= 0
    
    # Check diff document exists
    diff_doc = comparison_result['diff_doc']
    assert diff_doc is not None


def test_get_revisions(original_document, modified_document):
    """Test retrieving revisions after document comparison."""
    # Compare the documents
    comparison_result = compare_documents(original_document, modified_document)
    diff_doc = comparison_result['diff_doc']
    
    # Get revisions
    revisions = get_revisions(diff_doc)
    
    # Verify revisions list
    assert revisions is not None
    assert len(revisions) > 0
    
    # Check revision types
    revision_types = set(rev.type for rev in revisions)
    assert 'insertion' in revision_types or 'deletion' in revision_types


def test_get_revision_details(original_document, modified_document):
    """Test retrieving details of revisions."""
    # Compare the documents
    comparison_result = compare_documents(original_document, modified_document)
    diff_doc = comparison_result['diff_doc']
    
    # Get revisions
    revisions = get_revisions(diff_doc)
    
    # Check details of first revision
    if revisions:
        details = get_revision_details(revisions[0])
        assert details is not None
        assert 'author' in details
        assert 'date' in details
        assert 'type' in details
        assert 'text' in details


def test_get_deleted_text(original_document, modified_document):
    """Test retrieving deleted text from comparison."""
    # Compare the documents
    comparison_result = compare_documents(original_document, modified_document)
    diff_doc = comparison_result['diff_doc']
    
    # Get deleted text
    deleted_text = get_deleted_text(diff_doc)
    
    # Verify deleted text includes expected content
    assert deleted_text is not None
    assert len(deleted_text) > 0
    assert any('fourth paragraph' in text for text in deleted_text)


def test_get_inserted_text(original_document, modified_document):
    """Test retrieving inserted text from comparison."""
    # Compare the documents
    comparison_result = compare_documents(original_document, modified_document)
    diff_doc = comparison_result['diff_doc']
    
    # Get inserted text
    inserted_text = get_inserted_text(diff_doc)
    
    # Verify inserted text includes expected content
    assert inserted_text is not None
    assert len(inserted_text) > 0
    assert any('new paragraph added' in text for text in inserted_text)


def test_accept_revisions(original_document, modified_document):
    """Test accepting all revisions in a comparison document."""
    # Compare the documents
    comparison_result = compare_documents(original_document, modified_document)
    diff_doc = comparison_result['diff_doc']
    
    # Accept all revisions
    accepted_doc = accept_revisions(diff_doc)
    
    # Save the accepted document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        accepted_doc.save(tmp.name)
        accepted_path = tmp.name
    
    try:
        # Verify all revisions were accepted
        accepted_doc_loaded = Document(accepted_path)
        remaining_revisions = get_revisions(accepted_doc_loaded)
        assert len(remaining_revisions) == 0
        
        # Verify content matches the modified document
        modified_doc = Document(modified_document)
        
        # Check for new paragraph that was added
        found_new_paragraph = False
        for paragraph in accepted_doc_loaded.paragraphs:
            if 'new paragraph added' in paragraph.text:
                found_new_paragraph = True
                break
                
        assert found_new_paragraph
    finally:
        # Clean up
        os.unlink(accepted_path)


def test_reject_revisions(original_document, modified_document):
    """Test rejecting all revisions in a comparison document."""
    # Compare the documents
    comparison_result = compare_documents(original_document, modified_document)
    diff_doc = comparison_result['diff_doc']
    
    # Reject all revisions
    rejected_doc = reject_revisions(diff_doc)
    
    # Save the rejected document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        rejected_doc.save(tmp.name)
        rejected_path = tmp.name
    
    try:
        # Verify all revisions were rejected
        rejected_doc_loaded = Document(rejected_path)
        remaining_revisions = get_revisions(rejected_doc_loaded)
        assert len(remaining_revisions) == 0
        
        # Verify content is similar to original document
        original_doc = Document(original_document)
        
        # Check that fourth paragraph is still there (not deleted)
        found_fourth_paragraph = False
        for paragraph in rejected_doc_loaded.paragraphs:
            if 'fourth paragraph' in paragraph.text:
                found_fourth_paragraph = True
                break
                
        assert found_fourth_paragraph
        
        # Check new paragraph is not there
        found_new_paragraph = False
        for paragraph in rejected_doc_loaded.paragraphs:
            if 'new paragraph added' in paragraph.text:
                found_new_paragraph = True
                break
                
        assert not found_new_paragraph
    finally:
        # Clean up
        os.unlink(rejected_path)


def test_merge_documents(original_document):
    """Test merging multiple documents."""
    # Create additional documents for merging
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp1:
        doc1 = Document()
        doc1.add_heading('Document to Merge 1', level=1)
        doc1.add_paragraph('Content from the first document to merge.')
        doc1.save(tmp1.name)
        merge_doc1 = tmp1.name
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp2:
        doc2 = Document()
        doc2.add_heading('Document to Merge 2', level=1)
        doc2.add_paragraph('Content from the second document to merge.')
        doc2.save(tmp2.name)
        merge_doc2 = tmp2.name
    
    try:
        # Merge the documents
        docs_to_merge = [original_document, merge_doc1, merge_doc2]
        merged_doc = merge_documents(docs_to_merge)
        
        # Save the merged document
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_merged:
            merged_doc.save(tmp_merged.name)
            merged_path = tmp_merged.name
        
        try:
            # Verify merged document contains content from all source documents
            merged_doc_loaded = Document(merged_path)
            
            # Extract all text
            full_text = "\n".join([p.text for p in merged_doc_loaded.paragraphs])
            
            # Check for content from each source document
            assert 'Document Comparison Test - Original' in full_text
            assert 'Document to Merge 1' in full_text
            assert 'Content from the first document to merge' in full_text
            assert 'Document to Merge 2' in full_text
            assert 'Content from the second document to merge' in full_text
        finally:
            # Clean up
            os.unlink(merged_path)
    finally:
        # Clean up
        os.unlink(merge_doc1)
        os.unlink(merge_doc2)


def test_create_redline_document(original_document, modified_document):
    """Test creating a redline document showing changes."""
    # Create redline document
    redline_doc = create_redline_document(original_document, modified_document)
    
    # Save the redline document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        redline_doc.save(tmp.name)
        redline_path = tmp.name
    
    try:
        # Verify redline document contains formatted revision marks
        redline_doc_loaded = Document(redline_path)
        
        # Extract full text
        full_text = "\n".join([p.text for p in redline_doc_loaded.paragraphs])
        
        # Check for expected content
        assert 'Document Comparison Test - Original' in full_text
        assert 'This is the first paragraph' in full_text
        
        # Check for revisions
        revisions = get_revisions(redline_doc_loaded)
        assert len(revisions) > 0
        
        # Check revision types for strikethrough and insertion formatting
        revision_types = set(rev.type for rev in revisions)
        assert 'insertion' in revision_types or 'deletion' in revision_types
    finally:
        # Clean up
        os.unlink(redline_path)


def test_partial_document_comparison(original_document):
    """Test comparing specific sections of documents."""
    # Create a document with partial changes for comparison
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document(original_document)
        
        # Only modify the second paragraph
        doc.paragraphs[2].text = 'This is the second paragraph with a small change.'
        
        doc.save(tmp.name)
        partial_modified_doc = tmp.name
    
    try:
        # Compare the documents
        comparison_result = compare_documents(original_document, partial_modified_doc)
        diff_doc = comparison_result['diff_doc']
        
        # Get revisions
        revisions = get_revisions(diff_doc)
        
        # Verify only the expected content was changed
        assert any('small change' in rev.text for rev in revisions if rev.type == 'insertion')
        assert any('remain unchanged' in rev.text for rev in revisions if rev.type == 'deletion')
        
        # Verify other paragraphs were not affected
        deleted_text = get_deleted_text(diff_doc)
        assert not any('first paragraph' in text for text in deleted_text)
        assert not any('third paragraph' in text for text in deleted_text)
    finally:
        # Clean up
        os.unlink(partial_modified_doc)


if __name__ == '__main__':
    pytest.main(['-v', __file__]) 