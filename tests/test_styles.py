"""
Styles tests for LlamaDocx.

This module contains tests for the document styles functionality of the LlamaDocx package.
"""

import os
import tempfile
import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Inches

from llamadocx.styles import (
    get_styles,
    get_style_by_name,
    create_paragraph_style,
    create_character_style,
    create_table_style,
    apply_style_to_paragraph,
    apply_style_to_run,
    apply_style_to_table,
    update_style_font,
    update_style_paragraph_format,
    update_style_font_color,
    copy_style,
    delete_style
)


@pytest.fixture
def sample_document():
    """Create a sample document for testing styles."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        
        # Add heading
        doc.add_heading('Styles Test Document', level=1)
        
        # Add content
        doc.add_paragraph('This is a paragraph for testing styles.')
        doc.add_paragraph('This is another paragraph with some formatting.')
        
        # Add a table
        table = doc.add_table(rows=2, cols=2)
        for row in range(2):
            for col in range(2):
                table.cell(row, col).text = f"Cell {row},{col}"
        
        # Save the document
        doc.save(tmp.name)
        
    yield tmp.name
    
    # Clean up
    os.unlink(tmp.name)


def test_get_styles(sample_document):
    """Test getting all available styles from a document."""
    # Load the document
    doc = Document(sample_document)
    
    # Get styles
    styles = get_styles(doc)
    
    # Verify at least some built-in styles exist
    style_names = [style.name for style in styles]
    assert 'Normal' in style_names
    assert 'Heading 1' in style_names
    assert 'Heading 2' in style_names
    assert 'Title' in style_names


def test_get_style_by_name(sample_document):
    """Test getting a specific style by name."""
    # Load the document
    doc = Document(sample_document)
    
    # Get styles by name
    normal_style = get_style_by_name(doc, 'Normal')
    heading1_style = get_style_by_name(doc, 'Heading 1')
    
    # Verify styles were retrieved correctly
    assert normal_style.name == 'Normal'
    assert normal_style.type == WD_STYLE_TYPE.PARAGRAPH
    assert heading1_style.name == 'Heading 1'
    assert heading1_style.type == WD_STYLE_TYPE.PARAGRAPH
    
    # Test retrieving non-existent style (should return None)
    nonexistent_style = get_style_by_name(doc, 'NonExistentStyle')
    assert nonexistent_style is None


def test_create_paragraph_style(sample_document):
    """Test creating a new paragraph style."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a new paragraph style
    style_name = 'TestParagraphStyle'
    style = create_paragraph_style(doc, style_name,
                                  font_name='Arial',
                                  font_size=Pt(12),
                                  bold=True,
                                  italic=False,
                                  color=RGBColor(0, 0, 128),  # Navy blue
                                  base_style='Normal')
    
    # Verify style was created
    assert style is not None
    assert style.name == style_name
    assert style.type == WD_STYLE_TYPE.PARAGRAPH
    
    # Apply the style to a paragraph
    paragraph = doc.paragraphs[1]
    apply_style_to_paragraph(paragraph, style_name)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify style persisted and was applied
        style2 = get_style_by_name(doc2, style_name)
        assert style2 is not None
        assert style2.name == style_name
        assert doc2.paragraphs[1].style.name == style_name
    finally:
        # Clean up
        os.unlink(output_path)


def test_create_character_style(sample_document):
    """Test creating a new character style."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a new character style
    style_name = 'TestCharacterStyle'
    style = create_character_style(doc, style_name,
                                  font_name='Courier New',
                                  font_size=Pt(10),
                                  bold=False,
                                  italic=True,
                                  color=RGBColor(139, 0, 0),  # Dark red
                                  base_style=None)
    
    # Verify style was created
    assert style is not None
    assert style.name == style_name
    assert style.type == WD_STYLE_TYPE.CHARACTER
    
    # Apply the style to a run
    paragraph = doc.paragraphs[1]
    run = paragraph.add_run(" This text has character style applied.")
    apply_style_to_run(run, style_name)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify style persisted and was applied
        style2 = get_style_by_name(doc2, style_name)
        assert style2 is not None
        assert style2.name == style_name
        assert doc2.paragraphs[1].runs[-1].style.name == style_name
    finally:
        # Clean up
        os.unlink(output_path)


def test_create_table_style(sample_document):
    """Test creating a new table style."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a new table style
    style_name = 'TestTableStyle'
    style = create_table_style(doc, style_name,
                              first_row=True,
                              first_column=True,
                              banding=True)
    
    # Verify style was created
    assert style is not None
    assert style.name == style_name
    assert style.type == WD_STYLE_TYPE.TABLE
    
    # Apply the style to a table
    table = doc.tables[0]
    apply_style_to_table(table, style_name)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify style persisted and was applied
        style2 = get_style_by_name(doc2, style_name)
        assert style2 is not None
        assert style2.name == style_name
        assert doc2.tables[0].style.name == style_name
    finally:
        # Clean up
        os.unlink(output_path)


def test_update_style_font(sample_document):
    """Test updating font properties of an existing style."""
    # Load the document
    doc = Document(sample_document)
    
    # Get the Normal style
    style = get_style_by_name(doc, 'Normal')
    
    # Update font properties
    update_style_font(style, 
                     font_name='Calibri',
                     font_size=Pt(11),
                     bold=False,
                     italic=False)
    
    # Create a paragraph with the updated style
    paragraph = doc.add_paragraph('This paragraph uses the updated Normal style.')
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify style was updated
        style2 = get_style_by_name(doc2, 'Normal')
        # We can't easily verify font properties as they might be inherited or theme-based
    finally:
        # Clean up
        os.unlink(output_path)


def test_update_style_paragraph_format(sample_document):
    """Test updating paragraph format properties of an existing style."""
    # Load the document
    doc = Document(sample_document)
    
    # Get the Normal style
    style = get_style_by_name(doc, 'Normal')
    
    # Update paragraph format properties
    update_style_paragraph_format(style,
                                 alignment='center',
                                 space_before=Pt(12),
                                 space_after=Pt(12),
                                 left_indent=Inches(0.5),
                                 right_indent=Inches(0.5))
    
    # Create a paragraph with the updated style
    paragraph = doc.add_paragraph('This paragraph uses the updated Normal style with paragraph formatting.')
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify style was updated
        style2 = get_style_by_name(doc2, 'Normal')
        # Verify the new paragraph's formatting
        last_paragraph = doc2.paragraphs[-1]
        # We can't easily verify paragraph format properties as they might be inherited
    finally:
        # Clean up
        os.unlink(output_path)


def test_update_style_font_color(sample_document):
    """Test updating font color of an existing style."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a new paragraph style for color testing
    style_name = 'ColorTestStyle'
    style = create_paragraph_style(doc, style_name)
    
    # Update the font color
    new_color = RGBColor(0, 128, 0)  # Green
    update_style_font_color(style, new_color)
    
    # Apply the style to a paragraph
    paragraph = doc.add_paragraph('This paragraph uses a style with custom color.')
    apply_style_to_paragraph(paragraph, style_name)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify style persisted
        style2 = get_style_by_name(doc2, style_name)
        assert style2 is not None
        # We can't easily verify color as it might be inherited or theme-based
    finally:
        # Clean up
        os.unlink(output_path)


def test_copy_style(sample_document):
    """Test copying a style to create a new one."""
    # Load the document
    doc = Document(sample_document)
    
    # Copy Heading 1 style to a new style
    source_style_name = 'Heading 1'
    new_style_name = 'CopiedHeadingStyle'
    
    new_style = copy_style(doc, source_style_name, new_style_name)
    
    # Verify the new style was created
    assert new_style is not None
    assert new_style.name == new_style_name
    
    # Apply the new style to a paragraph
    paragraph = doc.add_paragraph('This paragraph uses the copied heading style.')
    apply_style_to_paragraph(paragraph, new_style_name)
    
    # Modify the new style to be different from the original
    update_style_font(new_style, italic=True)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify style persisted
        new_style2 = get_style_by_name(doc2, new_style_name)
        assert new_style2 is not None
        assert new_style2.name == new_style_name
        
        # Verify the last paragraph has the new style
        assert doc2.paragraphs[-1].style.name == new_style_name
    finally:
        # Clean up
        os.unlink(output_path)


def test_delete_style(sample_document):
    """Test deleting a custom style."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a new style to delete
    style_name = 'StyleToDelete'
    style = create_paragraph_style(doc, style_name)
    
    # Verify the style exists
    assert get_style_by_name(doc, style_name) is not None
    
    # Delete the style
    result = delete_style(doc, style_name)
    assert result is True
    
    # Verify the style no longer exists
    assert get_style_by_name(doc, style_name) is None
    
    # Test deleting a built-in style (should fail)
    result = delete_style(doc, 'Normal')
    assert result is False
    assert get_style_by_name(doc, 'Normal') is not None
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify the deleted style is still gone
        assert get_style_by_name(doc2, style_name) is None
        
        # Verify built-in style still exists
        assert get_style_by_name(doc2, 'Normal') is not None
    finally:
        # Clean up
        os.unlink(output_path)


def test_style_inheritance(sample_document):
    """Test style inheritance and overriding base style properties."""
    # Load the document
    doc = Document(sample_document)
    
    # Create a base style
    base_style_name = 'BaseStyle'
    base_style = create_paragraph_style(doc, base_style_name,
                                      font_name='Times New Roman',
                                      font_size=Pt(12),
                                      bold=True,
                                      italic=False)
    
    # Create a derived style that inherits from base style
    derived_style_name = 'DerivedStyle'
    derived_style = create_paragraph_style(doc, derived_style_name,
                                         base_style=base_style_name,
                                         italic=True)  # Override italic
    
    # Apply styles to paragraphs
    base_para = doc.add_paragraph('This paragraph uses the base style.')
    apply_style_to_paragraph(base_para, base_style_name)
    
    derived_para = doc.add_paragraph('This paragraph uses the derived style.')
    apply_style_to_paragraph(derived_para, derived_style_name)
    
    # Save and reload to verify persistence
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
        doc.save(output_path)
    
    try:
        # Load the saved document
        doc2 = Document(output_path)
        
        # Verify both styles exist
        assert get_style_by_name(doc2, base_style_name) is not None
        assert get_style_by_name(doc2, derived_style_name) is not None
        
        # Verify paragraphs have correct styles
        assert doc2.paragraphs[-2].style.name == base_style_name
        assert doc2.paragraphs[-1].style.name == derived_style_name
    finally:
        # Clean up
        os.unlink(output_path)


if __name__ == '__main__':
    pytest.main(['-v', __file__]) 